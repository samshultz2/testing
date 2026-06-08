"""
Main routes for dashboard and general pages
"""
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, session
from datetime import date, timedelta
from models import (
    db, Student, ParentContact, AcademicSession, Term, StudentEnrollment, 
    ClassArmAssignment, Attendance, Week, ClassArm, Subject, SchoolClass
)
from utils.access_control import (
    login_required, admin_required, is_admin, is_teacher, get_accessible_class_ids
)
from utils.audit import log_action
from utils.helpers import RELIGIONS, parse_date, FlashMessages, WAEC_SUBJECTS, STREAMS, STREAM_WAEC_SUBJECTS
from sqlalchemy import extract, func
from urllib.parse import urlparse


def _safe_next(target, fallback):
    """Return target only if it is a same-site relative path, else fallback."""
    if target:
        parsed = urlparse(target)
        if not parsed.scheme and not parsed.netloc and target.startswith('/'):
            return target
    return fallback

main_bp = Blueprint('main', __name__)


@main_bp.route('/set-branch')
@login_required
def set_view_branch():
    """Central users switch the branch they're viewing ('all' clears it)."""
    from utils.branch_scope import is_central, VIEW_KEY
    if is_central():
        raw = request.args.get('branch_id')
        if raw in (None, '', 'all'):
            session.pop(VIEW_KEY, None)
        else:
            try:
                session[VIEW_KEY] = int(raw)
            except (TypeError, ValueError):
                session.pop(VIEW_KEY, None)
    return redirect(request.referrer or url_for('main.dashboard'))


@main_bp.route('/')
@login_required
def dashboard():
    """Main dashboard page with comprehensive statistics"""
    from utils.branch_scope import scope_query
    # Basic student statistics (branch-scoped for branch users)
    total_students = scope_query(Student.query.filter_by(is_active=True), Student).count()
    male_students = scope_query(Student.query.filter_by(is_active=True, gender='Male'), Student).count()
    female_students = scope_query(Student.query.filter_by(is_active=True, gender='Female'), Student).count()

    # Get active session and term
    active_session = AcademicSession.query.filter_by(is_active=True).first()
    active_term = Term.query.filter_by(is_active=True).first()

    # Get recent students
    recent_students = Student.query.filter_by(is_active=True).order_by(
        Student.created_at.desc()
    ).limit(5).all()

    # Enrollments and class statistics
    active_enrollments = 0
    class_stats = []
    total_classes = 0
    
    if active_term:
        # Get enrollment count
        active_enrollments = StudentEnrollment.query.join(ClassArmAssignment).filter(
            ClassArmAssignment.term_id == active_term.id,
            StudentEnrollment.is_active == True
        ).count()
        
        # Get class-wise enrollment stats
        assignments = ClassArmAssignment.query.filter_by(term_id=active_term.id).all()
        total_classes = len(assignments)
        
        for assignment in assignments:
            enrollment_count = StudentEnrollment.query.filter_by(
                class_arm_assignment_id=assignment.id,
                is_active=True
            ).count()
            
            male_count = StudentEnrollment.query.join(Student).filter(
                StudentEnrollment.class_arm_assignment_id == assignment.id,
                StudentEnrollment.is_active == True,
                Student.gender == 'Male'
            ).count()
            
            female_count = enrollment_count - male_count
            
            class_stats.append({
                'name': assignment.display_name,
                'total': enrollment_count,
                'male': male_count,
                'female': female_count
            })
        
        # Sort by name
        class_stats.sort(key=lambda x: x['name'])

    # Attendance statistics for current term
    attendance_stats = {
        'today_present': 0,
        'today_absent': 0,
        'today_percentage': 0,
        'week_average': 0,
        'term_average': 0
    }
    
    if active_term:
        today = date.today()
        
        # Get current week
        current_week = Week.query.filter(
            Week.term_id == active_term.id,
            Week.start_date <= today,
            Week.end_date >= today
        ).first()
        
        if current_week:
            # Today's attendance
            today_records = Attendance.query.filter_by(
                week_id=current_week.id,
                date=today
            ).all()
            
            attendance_stats['today_present'] = sum(
                1 for a in today_records if a.morning_present or a.afternoon_present
            )
            attendance_stats['today_absent'] = sum(
                1 for a in today_records if not a.morning_present and not a.afternoon_present
            )
            
            total_today = len(today_records)
            if total_today > 0:
                present_marks = sum(
                    (1 if a.morning_present else 0) + (1 if a.afternoon_present else 0)
                    for a in today_records
                )
                attendance_stats['today_percentage'] = round(
                    (present_marks / (total_today * 2)) * 100, 1
                )
        
        # Get all weeks for term average
        term_weeks = Week.query.filter_by(term_id=active_term.id).all()
        if term_weeks:
            term_records = Attendance.query.filter(
                Attendance.week_id.in_([w.id for w in term_weeks])
            ).all()
            
            if term_records:
                total_marks = sum(
                    (1 if a.morning_present else 0) + (1 if a.afternoon_present else 0)
                    for a in term_records
                )
                attendance_stats['term_average'] = round(
                    (total_marks / (len(term_records) * 2)) * 100, 1
                )

    # Birthdays
    today = date.today()
    today_month = today.month
    today_day = today.day

    # Birthdays today
    birthdays_today = Student.query.filter(
        Student.is_active == True,
        Student.date_of_birth != None,
        extract('month', Student.date_of_birth) == today_month,
        extract('day', Student.date_of_birth) == today_day
    ).all()

    # Birthdays this week (next 7 days, excluding today)
    birthdays_week = []
    for i in range(1, 7):
        check_date = today + timedelta(days=i)
        students = Student.query.filter(
            Student.is_active == True,
            Student.date_of_birth != None,
            extract('month', Student.date_of_birth) == check_date.month,
            extract('day', Student.date_of_birth) == check_date.day
        ).all()
        for s in students:
            s.birthday_date = check_date
        birthdays_week.extend(students)

    # Age distribution
    age_distribution = {'0-10': 0, '11-13': 0, '14-16': 0, '17-19': 0, '20+': 0}
    students_with_dob = Student.query.filter(
        Student.is_active == True,
        Student.date_of_birth != None
    ).all()
    
    for student in students_with_dob:
        age = student.age
        if age <= 10:
            age_distribution['0-10'] += 1
        elif age <= 13:
            age_distribution['11-13'] += 1
        elif age <= 16:
            age_distribution['14-16'] += 1
        elif age <= 19:
            age_distribution['17-19'] += 1
        else:
            age_distribution['20+'] += 1

    # Religion distribution
    religion_distribution = db.session.query(
        Student.religion, func.count(Student.id)
    ).filter(
        Student.is_active == True,
        Student.religion != None
    ).group_by(Student.religion).all()
    
    religion_stats = {r[0]: r[1] for r in religion_distribution if r[0]}

    # Get total subjects and classes configured
    total_subjects = Subject.query.filter_by(is_active=True).count()
    total_school_classes = SchoolClass.query.filter_by(is_active=True).count()

    # --- Richer dashboard widgets ---
    from models import WAECResult, JAMBResult, AuditLog
    from models.mock_jamb import MockJAMBExam

    # Stream distribution (active students)
    stream_rows = db.session.query(Student.stream, func.count(Student.id)).filter(
        Student.is_active == True).group_by(Student.stream).all()
    stream_dist = {(s or 'Unset'): c for s, c in stream_rows}

    graduates_count = Student.query.filter_by(is_active=True, is_graduated=True).count()

    # JAMB snapshot (latest year)
    jamb_snapshot = None
    jy = db.session.query(JAMBResult.exam_year).order_by(JAMBResult.exam_year.desc()).first()
    if jy:
        scores = [r.total_score for r in JAMBResult.query.filter_by(exam_year=jy[0]).all()]
        if scores:
            top = JAMBResult.query.filter_by(exam_year=jy[0]).order_by(
                JAMBResult.total_score.desc()).limit(5).all()
            jamb_snapshot = {
                'year': jy[0], 'count': len(scores),
                'mean': round(sum(scores) / len(scores), 1),
                'max': max(scores),
                'above_200': sum(1 for s in scores if s >= 200),
                'above_200_pct': round(sum(1 for s in scores if s >= 200) / len(scores) * 100),
                'distribution': {
                    '0-149': sum(1 for s in scores if s < 150),
                    '150-199': sum(1 for s in scores if 150 <= s < 200),
                    '200-249': sum(1 for s in scores if 200 <= s < 250),
                    '250-299': sum(1 for s in scores if 250 <= s < 300),
                    '300+': sum(1 for s in scores if s >= 300),
                },
                'top': [{'name': r.student.full_name, 'score': r.total_score} for r in top],
            }

    # WAEC snapshot (latest year)
    waec_snapshot = None
    wy = db.session.query(WAECResult.exam_year).order_by(WAECResult.exam_year.desc()).first()
    if wy:
        rows = WAECResult.query.filter_by(exam_year=wy[0]).all()
        if rows:
            credit = {'A1', 'B2', 'B3', 'C4', 'C5', 'C6'}
            passes = sum(1 for r in rows if r.grade in credit)
            waec_snapshot = {
                'year': wy[0], 'entries': len(rows),
                'students': len({r.student_id for r in rows}),
                'pass_rate': round(passes / len(rows) * 100, 1),
            }

    # Mock JAMB latest exam average
    mock_snapshot = None
    last_mock = MockJAMBExam.query.order_by(MockJAMBExam.exam_date.desc()).first()
    if last_mock:
        ms = [r.total_score for r in last_mock.results.all()]
        if ms:
            mock_snapshot = {'name': last_mock.display_name, 'count': len(ms),
                             'mean': round(sum(ms) / len(ms), 1), 'max': max(ms)}

    # Attendance trend — average % for the last 8 weeks of the active term
    attendance_trend = []
    if active_term:
        weeks = Week.query.filter_by(term_id=active_term.id).order_by(Week.start_date).all()[-8:]
        for w in weeks:
            recs = Attendance.query.filter_by(week_id=w.id).all()
            if recs:
                marks = sum((1 if a.morning_present else 0) + (1 if a.afternoon_present else 0) for a in recs)
                pct = round(marks / (len(recs) * 2) * 100, 1)
            else:
                pct = 0
            attendance_trend.append({'label': getattr(w, 'week_number', None) or w.start_date.strftime('%d %b'),
                                     'pct': pct})

    # Recent admin activity (admins only)
    recent_activity = []
    if is_admin():
        recent_activity = AuditLog.query.order_by(AuditLog.created_at.desc()).limit(6).all()

    # Active announcements for the dashboard banner.
    announcements = []
    try:
        from models import Announcement
        for a in (Announcement.query.order_by(Announcement.is_pinned.desc(),
                  Announcement.created_at.desc()).limit(15).all()):
            if a.is_active:
                announcements.append(a)
        announcements = announcements[:5]
    except Exception:
        announcements = []

    return render_template('dashboard.html',
        announcements=announcements,
        total_students=total_students,
        male_students=male_students,
        female_students=female_students,
        active_session=active_session,
        active_term=active_term,
        recent_students=recent_students,
        active_enrollments=active_enrollments,
        total_classes=total_classes,
        class_stats=class_stats,
        attendance_stats=attendance_stats,
        birthdays_today=birthdays_today,
        birthdays_week=birthdays_week,
        age_distribution=age_distribution,
        religion_stats=religion_stats,
        total_subjects=total_subjects,
        total_school_classes=total_school_classes,
        stream_dist=stream_dist,
        graduates_count=graduates_count,
        jamb_snapshot=jamb_snapshot,
        waec_snapshot=waec_snapshot,
        mock_snapshot=mock_snapshot,
        attendance_trend=attendance_trend,
        recent_activity=recent_activity
    )


# ============================================================================
# STUDENT ROUTES
# ============================================================================

@main_bp.route('/students')
@login_required
def students_list():
    """List all students with filtering and pagination"""
    page = request.args.get('page', 1, type=int)
    per_page = 20
    is_ajax = request.args.get('ajax') == '1'

    # Get filter parameters
    search = request.args.get('search', '')
    gender = request.args.get('gender', '')
    religion = request.args.get('religion', '')
    stream = request.args.get('stream', '')
    subject = request.args.get('subject', '')
    class_id = request.args.get('class_id', '', type=int) or None
    arm_id = request.args.get('arm_id', '', type=int) or None
    sort_by = request.args.get('sort', 'surname')
    order = request.args.get('order', 'asc')

    # Build query (branch + section/stream scoped)
    from utils.branch_scope import scope_query
    from utils.org_scope import scope_students
    query = scope_query(Student.query.filter_by(is_active=True), Student)
    query = scope_students(query)

    # Teachers are further limited to their assigned classes. Other non-admin
    # roles (principal, HOD, bursar…) rely on branch/section/stream scoping above.
    active_term = Term.query.filter_by(is_active=True).first()
    if is_teacher() and active_term:
        accessible_class_ids = get_accessible_class_ids()
        if accessible_class_ids:
            query = query.join(
                StudentEnrollment,
                Student.id == StudentEnrollment.student_id
            ).join(
                ClassArmAssignment,
                StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id
            ).filter(
                ClassArmAssignment.term_id == active_term.id,
                ClassArmAssignment.id.in_(accessible_class_ids)
            )
        else:
            # No accessible classes, return empty
            query = query.filter(Student.id == -1)

    # Apply search filter
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            db.or_(
                Student.first_name.ilike(search_term),
                Student.surname.ilike(search_term),
                Student.middle_name.ilike(search_term),
                Student.student_id.ilike(search_term)
            )
        )

    # Apply gender filter
    if gender:
        query = query.filter(Student.gender == gender)

    # Apply religion filter
    if religion:
        query = query.filter(Student.religion == religion)

    # Apply stream filter
    if stream:
        query = query.filter(Student.stream == stream)

    # Apply WAEC-subject filter (SSS3 students only). Uses a subquery so it
    # doesn't interfere with the class/arm/teacher joins above.
    if subject:
        sss3_q = (db.session.query(StudentEnrollment.student_id)
                  .join(ClassArmAssignment,
                        StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id)
                  .join(SchoolClass, ClassArmAssignment.class_id == SchoolClass.id)
                  .filter(StudentEnrollment.is_active == True, SchoolClass.name == 'SSS3'))
        if active_term:
            sss3_q = sss3_q.filter(ClassArmAssignment.term_id == active_term.id)
        # Token-bounded match within the comma-joined waec_subjects string.
        query = query.filter(
            Student.id.in_(sss3_q),
            db.or_(Student.waec_subjects == subject,
                   Student.waec_subjects.ilike(f'{subject}, %'),
                   Student.waec_subjects.ilike(f'%, {subject}, %'),
                   Student.waec_subjects.ilike(f'%, {subject}')))

    # Apply class/arm filter - need to join with enrollments (only if admin since teacher already filtered)
    if (class_id or arm_id) and is_admin():
        if active_term:
            # Join with enrollments for current term
            query = query.join(
                StudentEnrollment,
                Student.id == StudentEnrollment.student_id
            ).join(
                ClassArmAssignment,
                StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id
            ).filter(
                ClassArmAssignment.term_id == active_term.id
            )
            
            if class_id:
                query = query.filter(ClassArmAssignment.class_id == class_id)
            if arm_id:
                query = query.filter(ClassArmAssignment.arm_id == arm_id)

    # Apply sorting
    sort_column = getattr(Student, sort_by, Student.surname)
    if order == 'desc':
        query = query.order_by(sort_column.desc())
    else:
        query = query.order_by(sort_column.asc())

    # Paginate
    students = query.paginate(page=page, per_page=per_page, error_out=False)
    
    # Add current class info to each student (single query for the whole page)
    active_term = Term.query.filter_by(is_active=True).first()
    class_map = {}
    page_ids = [s.id for s in students.items]
    if active_term and page_ids:
        enrollments = StudentEnrollment.query.join(ClassArmAssignment).filter(
            StudentEnrollment.student_id.in_(page_ids),
            ClassArmAssignment.term_id == active_term.id
        ).all()
        for enrollment in enrollments:
            class_map.setdefault(enrollment.student_id, enrollment.class_arm_assignment.display_name)
    for student in students.items:
        student.current_class = class_map.get(student.id)

    # Get filter options
    classes = SchoolClass.query.filter_by(is_active=True).order_by(SchoolClass.level).all()
    arms = ClassArm.query.filter_by(is_active=True).order_by(ClassArm.name).all()

    # If AJAX request, return only the content partial
    if is_ajax:
        return render_template('students/_list_content.html',
            students=students,
            search=search,
            gender=gender,
            religion=religion,
            stream=stream,
            subject=subject,
            class_id=class_id,
            arm_id=arm_id
        )

    return render_template('students/list.html',
        students=students,
        search=search,
        gender=gender,
        religion=religion,
        stream=stream,
        subject=subject,
        class_id=class_id,
        arm_id=arm_id,
        sort_by=sort_by,
        order=order,
        religions=RELIGIONS,
        classes=classes,
        arms=arms,
        streams=STREAMS,
        subject_options=WAEC_SUBJECTS
    )


@main_bp.route('/students/add', methods=['GET', 'POST'])
@login_required
def add_student():
    """Add a new student"""
    if request.method == 'POST':
        try:
            # Create student
            student = Student(
                student_id=Student.generate_student_id(),
                first_name=request.form.get('first_name', '').strip(),
                middle_name=request.form.get('middle_name', '').strip() or None,
                surname=request.form.get('surname', '').strip(),
                gender=request.form.get('gender'),
                date_of_birth=parse_date(request.form.get('date_of_birth')),
                religion=request.form.get('religion'),
                home_address=request.form.get('home_address', '').strip() or None,
                hobbies=request.form.get('hobbies', '').strip() or None,
                waec_subjects=', '.join(request.form.getlist('waec_subjects[]')) or None,
                jamb_subjects=', '.join(request.form.getlist('jamb_subjects[]')) or None,
                stream=request.form.get('stream') or None,
                jamb_target=request.form.get('jamb_target', type=int)
            )
            # Stamp the student with the creator's (or chosen) branch.
            from utils.branch_scope import branch_for_new
            student.branch_id = branch_for_new(request.form.get('branch_id', type=int))

            db.session.add(student)
            db.session.flush()

            # Add parent contacts
            phone_numbers = request.form.getlist('phone_number[]')
            relationships = request.form.getlist('relationship[]')
            contact_names = request.form.getlist('contact_name[]')

            for i, phone in enumerate(phone_numbers):
                if phone.strip():
                    contact = ParentContact(
                        student_id=student.id,
                        phone_number=phone.strip(),
                        relationship=relationships[i] if i < len(relationships) else 'Guardian',
                        name=contact_names[i] if i < len(contact_names) else None,
                        is_primary=(i == 0)
                    )
                    db.session.add(contact)

            db.session.commit()
            flash(FlashMessages.STUDENT_CREATED, 'success')
            return redirect(url_for('main.view_student', student_id=student.id))

        except Exception as e:
            db.session.rollback()
            flash(f'Error creating student: {str(e)}', 'error')

    return render_template('students/add.html', religions=RELIGIONS, waec_subjects=WAEC_SUBJECTS,
                           streams=STREAMS, stream_waec=STREAM_WAEC_SUBJECTS)


@main_bp.route('/students/<int:student_id>')
@login_required
def view_student(student_id):
    """View student details"""
    student = Student.query.get_or_404(student_id)
    from utils.branch_scope import can_access_branch
    if not can_access_branch(student.branch_id):
        flash('That student belongs to another branch.', 'error')
        return redirect(url_for('main.students_list'))

    # Get enrollments
    enrollments = student.enrollments.join(ClassArmAssignment).order_by(
        ClassArmAssignment.term_id.desc()
    ).all()

    # Get WAEC results
    waec_results = student.waec_results.all()

    # Get JAMB results
    jamb_results = student.jamb_results.all()

    return render_template('students/view.html',
        student=student,
        enrollments=enrollments,
        waec_results=waec_results,
        jamb_results=jamb_results
    )


@main_bp.route('/students/<int:student_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_student(student_id):
    """Edit student details"""
    student = Student.query.get_or_404(student_id)

    if request.method == 'POST':
        try:
            # When the full edit form is submitted, fields absent from the POST
            # mean "cleared"; for any other (partial/programmatic) POST we only
            # touch fields that are actually present, so nothing gets blanked
            # accidentally.
            complete = request.form.get('form_complete') == '1'
            form = request.form

            def has(key):
                return complete or key in form

            if has('first_name'):
                student.first_name = form.get('first_name', '').strip()
            if has('surname'):
                student.surname = form.get('surname', '').strip()
            if has('middle_name'):
                student.middle_name = form.get('middle_name', '').strip() or None
            if has('gender'):
                student.gender = form.get('gender')
            if has('date_of_birth'):
                student.date_of_birth = parse_date(form.get('date_of_birth'))
            if has('religion'):
                student.religion = form.get('religion')
            if has('home_address'):
                student.home_address = form.get('home_address', '').strip() or None
            if has('hobbies'):
                student.hobbies = form.get('hobbies', '').strip() or None
            if has('stream'):
                student.stream = form.get('stream') or None
            if has('jamb_target'):
                student.jamb_target = form.get('jamb_target', type=int)
            if complete or 'waec_subjects[]' in form:
                student.waec_subjects = ', '.join(form.getlist('waec_subjects[]')) or None
            if complete or 'jamb_subjects[]' in form:
                student.jamb_subjects = ', '.join(form.getlist('jamb_subjects[]')) or None

            # Update contacts only when the contacts section was submitted
            if not (complete or 'phone_number[]' in form):
                db.session.commit()
                flash(FlashMessages.STUDENT_UPDATED, 'success')
                return redirect(_safe_next(form.get('return_to'),
                                           url_for('main.view_student', student_id=student.id)))

            ParentContact.query.filter_by(student_id=student.id).delete()

            phone_numbers = request.form.getlist('phone_number[]')
            relationships = request.form.getlist('relationship[]')
            contact_names = request.form.getlist('contact_name[]')

            for i, phone in enumerate(phone_numbers):
                if phone.strip():
                    contact = ParentContact(
                        student_id=student.id,
                        phone_number=phone.strip(),
                        relationship=relationships[i] if i < len(relationships) else 'Guardian',
                        name=contact_names[i] if i < len(contact_names) else None,
                        is_primary=(i == 0)
                    )
                    db.session.add(contact)

            db.session.commit()
            flash(FlashMessages.STUDENT_UPDATED, 'success')
            return redirect(_safe_next(
                request.form.get('return_to'),
                url_for('main.view_student', student_id=student.id)
            ))

        except Exception as e:
            db.session.rollback()
            flash(f'Error updating student: {str(e)}', 'error')

    # Remember where the user came from so we can return there after saving.
    return_to = _safe_next(
        request.form.get('return_to') or request.args.get('return_to') or request.referrer, '')
    return render_template('students/edit.html', student=student, religions=RELIGIONS, waec_subjects=WAEC_SUBJECTS,
                           streams=STREAMS, stream_waec=STREAM_WAEC_SUBJECTS, return_to=return_to)


@main_bp.route('/students/<int:student_id>/delete', methods=['POST'])
@admin_required
def delete_student(student_id):
    """Delete a student (soft delete)"""
    student = Student.query.get_or_404(student_id)

    try:
        student.is_active = False
        db.session.commit()
        log_action('delete_student', f'{student.full_name} ({student.student_id})')
        flash(FlashMessages.STUDENT_DELETED, 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting student: {str(e)}', 'error')

    return redirect(url_for('main.students_list'))


@main_bp.route('/students/bulk-stream', methods=['POST'])
@admin_required
def bulk_set_stream():
    """Set the stream/track for several students at once."""
    stream = request.form.get('stream') or None
    student_ids = request.form.getlist('student_ids')

    if stream is not None and stream not in STREAMS:
        return jsonify({'error': 'Invalid stream'}), 400
    if not student_ids:
        return jsonify({'error': 'No students selected'}), 400

    try:
        ids = [int(i) for i in student_ids]
    except (TypeError, ValueError):
        return jsonify({'error': 'Invalid student ids'}), 400

    updated = Student.query.filter(Student.id.in_(ids)).update(
        {Student.stream: stream}, synchronize_session=False
    )
    db.session.commit()

    label = stream if stream else 'cleared'
    log_action('bulk_set_stream', f'{updated} students -> {label}')
    flash(f'Stream set to {label} for {updated} student(s).', 'success')
    return jsonify({'updated': updated, 'stream': stream})


@main_bp.route('/students/bulk-add-subject', methods=['POST'])
@admin_required
def bulk_add_subject():
    """Add a WAEC subject to selected students' enrolled subjects (SSS3 only)."""
    subject = (request.form.get('subject') or '').strip()
    student_ids = request.form.getlist('student_ids')
    if subject not in WAEC_SUBJECTS:
        return jsonify({'error': 'Invalid subject'}), 400
    if not student_ids:
        return jsonify({'error': 'No students selected'}), 400
    try:
        ids = [int(i) for i in student_ids]
    except (TypeError, ValueError):
        return jsonify({'error': 'Invalid student ids'}), 400

    # Which of the selected students are SSS3 (current/active term)?
    active_term = Term.query.filter_by(is_active=True).first()
    sss3_q = (db.session.query(StudentEnrollment.student_id)
              .join(ClassArmAssignment,
                    StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id)
              .join(SchoolClass, ClassArmAssignment.class_id == SchoolClass.id)
              .filter(StudentEnrollment.is_active == True,
                      SchoolClass.name == 'SSS3',
                      StudentEnrollment.student_id.in_(ids)))
    if active_term:
        sss3_q = sss3_q.filter(ClassArmAssignment.term_id == active_term.id)
    sss3_ids = {r[0] for r in sss3_q.all()}

    updated = skipped = 0
    for student in Student.query.filter(Student.id.in_(ids)).all():
        if student.id not in sss3_ids:
            skipped += 1
            continue
        subs = student.waec_subject_list
        if subject in subs:
            skipped += 1
            continue
        subs.append(subject)
        student.waec_subjects = ', '.join(subs)
        updated += 1
    db.session.commit()
    log_action('bulk_add_subject', f'{subject} -> {updated} SSS3 students')
    flash(f'Added "{subject}" to {updated} SSS3 student(s).', 'success')
    return jsonify({'updated': updated, 'skipped': skipped, 'subject': subject})


@main_bp.route('/students/apply-stream-waec', methods=['POST'])
@admin_required
def apply_stream_waec():
    """Fill WAEC subjects from each student's stream where not already set."""
    updated = 0
    for student in Student.query.filter_by(is_active=True).all():
        defaults = STREAM_WAEC_SUBJECTS.get(student.stream)
        if defaults and not student.waec_subject_list:
            student.waec_subjects = ', '.join(defaults)
            updated += 1
    db.session.commit()
    flash(f'WAEC subjects filled from stream for {updated} student(s).', 'success')
    return redirect(request.referrer or url_for('main.students_list'))


@main_bp.route('/students/trash')
@admin_required
def students_trash():
    """List soft-deleted students with restore / permanent-delete options."""
    students = Student.query.filter_by(is_active=False).order_by(Student.surname, Student.first_name).all()
    return render_template('students/trash.html', students=students)


@main_bp.route('/students/<int:student_id>/restore', methods=['POST'])
@admin_required
def restore_student(student_id):
    student = Student.query.get_or_404(student_id)
    student.is_active = True
    db.session.commit()
    log_action('restore_student', f'{student.full_name} ({student.student_id})')
    flash(f'{student.full_name} restored.', 'success')
    return redirect(url_for('main.students_trash'))


@main_bp.route('/students/<int:student_id>/purge', methods=['POST'])
@admin_required
def purge_student(student_id):
    """Permanently delete a soft-deleted student and their related records."""
    student = Student.query.get_or_404(student_id)
    if student.is_active:
        flash('Only deleted students can be permanently removed.', 'error')
        return redirect(url_for('main.students_trash'))
    name = student.full_name
    sid = student.student_id
    try:
        db.session.delete(student)
        db.session.commit()
        log_action('purge_student', f'{name} ({sid})')
        flash(f'{name} permanently deleted.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'error')
    return redirect(url_for('main.students_trash'))


@main_bp.route('/search')
@login_required
def global_search():
    """Search across the main entities and group the results."""
    q = (request.args.get('q') or '').strip()
    groups = []
    if len(q) >= 2:
        like = f'%{q}%'

        def add(title, icon, items):
            if items:
                groups.append({'title': title, 'icon': icon, 'rows': items})

        students = (Student.query.filter_by(is_active=True)
                    .filter(db.or_(Student.first_name.ilike(like), Student.surname.ilike(like),
                                   Student.student_id.ilike(like)))
                    .order_by(Student.surname).limit(12).all())
        add('Students', 'fa-user-graduate', [
            {'label': s.full_name, 'sub': s.student_id,
             'url': url_for('main.view_student', student_id=s.id)} for s in students])

        try:
            from models import StaffMember
            staff = (StaffMember.query.filter_by(is_active=True)
                     .filter(db.or_(StaffMember.first_name.ilike(like), StaffMember.surname.ilike(like),
                                    StaffMember.staff_id.ilike(like), StaffMember.phone.ilike(like)))
                     .limit(10).all())
            add('Staff', 'fa-id-badge', [
                {'label': s.full_name, 'sub': s.designation or s.staff_id,
                 'url': url_for('hr.staff_detail', staff_id=s.id)} for s in staff])
        except Exception:
            pass

        try:
            from models import Applicant
            apps = (Applicant.query.filter(db.or_(Applicant.first_name.ilike(like),
                    Applicant.surname.ilike(like), Applicant.application_no.ilike(like)))
                    .limit(10).all())
            add('Applicants', 'fa-clipboard-user', [
                {'label': a.full_name, 'sub': f'{a.application_no} · {a.status}',
                 'url': url_for('admissions.applicant_detail', applicant_id=a.id)} for a in apps])
        except Exception:
            pass

        try:
            from models import Book
            books = (Book.query.filter_by(is_active=True)
                     .filter(db.or_(Book.title.ilike(like), Book.author.ilike(like), Book.isbn.ilike(like)))
                     .limit(10).all())
            add('Library books', 'fa-book', [
                {'label': b.title, 'sub': b.author or '', 'url': url_for('library.books', q=b.title)}
                for b in books])
        except Exception:
            pass

        try:
            from models import FeePayment
            pays = FeePayment.query.filter(FeePayment.receipt_no.ilike(like)).limit(8).all()
            add('Fee receipts', 'fa-receipt', [
                {'label': p.receipt_no, 'sub': (p.student.full_name if p.student else ''),
                 'url': url_for('finance.receipt', payment_id=p.id)} for p in pays])
        except Exception:
            pass

        try:
            from models import CBTExam
            exams = CBTExam.query.filter(CBTExam.title.ilike(like)).limit(8).all()
            add('CBT exams', 'fa-laptop-code', [
                {'label': e.title, 'sub': (e.subject.name if e.subject else ''),
                 'url': url_for('cbt.exam_detail', exam_id=e.id)} for e in exams])
        except Exception:
            pass

    return render_template('search.html', q=q, groups=groups,
                           count=sum(len(g['rows']) for g in groups))


@main_bp.route('/audit')
@admin_required
def audit_log():
    """View the audit trail of administrative actions (with filters)."""
    from models import AuditLog
    from datetime import datetime as _dt
    page = request.args.get('page', 1, type=int)
    q = (request.args.get('q') or '').strip()
    action = (request.args.get('action') or '').strip()
    user = (request.args.get('user') or '').strip()
    from_s = request.args.get('from')
    to_s = request.args.get('to')

    query = AuditLog.query
    if q:
        like = f'%{q}%'
        query = query.filter(db.or_(AuditLog.action.ilike(like),
                                    AuditLog.detail.ilike(like),
                                    AuditLog.user.ilike(like)))
    if action:
        query = query.filter(AuditLog.action == action)
    if user:
        query = query.filter(AuditLog.user.ilike(f'%{user}%'))
    try:
        if from_s:
            query = query.filter(AuditLog.created_at >= _dt.strptime(from_s, '%Y-%m-%d'))
        if to_s:
            d = _dt.strptime(to_s, '%Y-%m-%d')
            query = query.filter(AuditLog.created_at < d.replace(hour=23, minute=59, second=59))
    except ValueError:
        pass

    logs = query.order_by(AuditLog.created_at.desc()).paginate(page=page, per_page=50, error_out=False)
    actions = [a[0] for a in db.session.query(AuditLog.action).distinct().order_by(AuditLog.action).all()]
    return render_template('audit.html', logs=logs, actions=actions,
        q=q, action=action, user=user, from_s=from_s or '', to_s=to_s or '')


# ============================================================================
# API ENDPOINTS FOR AJAX
# ============================================================================

@main_bp.route('/api/students/search')
@login_required
def api_search_students():
    """API endpoint for student search (AJAX)"""
    query = request.args.get('q', '')

    if len(query) < 2:
        return jsonify([])

    students = Student.query.filter(
        Student.is_active == True,
        db.or_(
            Student.first_name.ilike(f'%{query}%'),
            Student.surname.ilike(f'%{query}%'),
            Student.student_id.ilike(f'%{query}%')
        )
    ).limit(10).all()

    return jsonify([{
        'id': s.id,
        'student_id': s.student_id,
        'name': s.full_name,
        'gender': s.gender
    } for s in students])


@main_bp.route('/api/dashboard/stats')
@login_required
def api_dashboard_stats():
    """API endpoint for dashboard statistics"""
    from utils.calculations import get_attendance_statistics

    active_term = Term.query.filter_by(is_active=True).first()

    stats = {
        'total_students': Student.query.filter_by(is_active=True).count(),
        'male_students': Student.query.filter_by(is_active=True, gender='Male').count(),
        'female_students': Student.query.filter_by(is_active=True, gender='Female').count(),
    }

    if active_term:
        attendance_stats = get_attendance_statistics(active_term.id)
        stats.update(attendance_stats)

    return jsonify(stats)


# ============================================================================
# STUDENT EXPORT FUNCTIONALITY
# ============================================================================

@main_bp.route('/students/export')
@login_required
def export_students_data():
    """Export selected students data to various formats with field selection"""
    import json
    
    format_type = request.args.get('format', 'excel')
    fields_json = request.args.get('fields', '[]')
    student_ids_json = request.args.get('student_ids', '[]')
    
    try:
        fields = json.loads(fields_json)
        student_ids = json.loads(student_ids_json)
    except Exception:
        fields = ['student_id', 'surname', 'first_name', 'gender', 'current_class']
        student_ids = []
    
    if not fields:
        flash('No fields selected for export.', 'error')
        return redirect(url_for('main.students_list'))
    
    # Build query
    if student_ids:
        # Export selected students
        query = Student.query.filter(Student.id.in_(student_ids))
    else:
        # Export all students matching current filters
        query = Student.query.filter_by(is_active=True)
        
        # Apply filters
        search = request.args.get('search', '')
        gender = request.args.get('gender', '')
        religion = request.args.get('religion', '')
        class_id = request.args.get('class_id', type=int)
        arm_id = request.args.get('arm_id', type=int)
        
        if search:
            search_term = f"%{search}%"
            query = query.filter(
                db.or_(
                    Student.first_name.ilike(search_term),
                    Student.surname.ilike(search_term),
                    Student.middle_name.ilike(search_term),
                    Student.student_id.ilike(search_term)
                )
            )
        if gender:
            query = query.filter(Student.gender == gender)
        if religion:
            query = query.filter(Student.religion == religion)
        
        if class_id or arm_id:
            active_term = Term.query.filter_by(is_active=True).first()
            if active_term:
                query = query.join(
                    StudentEnrollment, Student.id == StudentEnrollment.student_id
                ).join(
                    ClassArmAssignment, StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id
                ).filter(ClassArmAssignment.term_id == active_term.id)
                
                if class_id:
                    query = query.filter(ClassArmAssignment.class_id == class_id)
                if arm_id:
                    query = query.filter(ClassArmAssignment.arm_id == arm_id)
    
    students = query.order_by(Student.surname).all()
    
    if not students:
        flash('No students to export.', 'error')
        return redirect(url_for('main.students_list'))
    
    # Get current class and parent info for each student
    active_term = Term.query.filter_by(is_active=True).first()
    student_data = []
    
    for student in students:
        data = {}
        
        # Basic fields
        if 'student_id' in fields:
            data['Student ID'] = student.student_id
        if 'surname' in fields:
            data['Surname'] = student.surname
        if 'first_name' in fields:
            data['First Name'] = student.first_name
        if 'middle_name' in fields:
            data['Middle Name'] = student.middle_name or ''
        if 'gender' in fields:
            data['Gender'] = student.gender
        if 'date_of_birth' in fields:
            data['Date of Birth'] = student.date_of_birth.strftime('%Y-%m-%d') if student.date_of_birth else ''
        if 'age' in fields:
            data['Age'] = student.age or ''
        if 'religion' in fields:
            data['Religion'] = student.religion or ''
        if 'home_address' in fields:
            data['Home Address'] = student.home_address or ''
        if 'hobbies' in fields:
            data['Hobbies'] = student.hobbies or ''
        
        # Current class
        if 'current_class' in fields:
            current_class = ''
            if active_term:
                enrollment = StudentEnrollment.query.join(
                    ClassArmAssignment
                ).filter(
                    StudentEnrollment.student_id == student.id,
                    ClassArmAssignment.term_id == active_term.id
                ).first()
                if enrollment:
                    current_class = enrollment.class_arm_assignment.display_name
            data['Class'] = current_class
        
        # Parent phone
        if 'parent_phone' in fields:
            parent = student.parent_contacts.first()
            data['Parent Phone'] = parent.phone if parent else ''
        
        student_data.append(data)
    
    # Get ordered field names for export
    field_order = ['Student ID', 'Surname', 'First Name', 'Middle Name', 'Gender', 
                   'Class', 'Date of Birth', 'Age', 'Religion', 'Home Address', 
                   'Hobbies', 'Parent Phone']
    export_fields = [f for f in field_order if f in student_data[0]] if student_data else []
    
    if format_type == 'excel':
        return export_students_excel(student_data, export_fields)
    elif format_type == 'word':
        return export_students_word(student_data, export_fields)
    elif format_type == 'pdf':
        return export_students_pdf(student_data, export_fields)
    elif format_type == 'image':
        return export_students_image(student_data, export_fields)
    else:
        flash('Invalid export format.', 'error')
        return redirect(url_for('main.students_list'))


def export_students_excel(student_data, fields):
    """Export students to Excel format with selected fields - wraps text and adjusts row heights"""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    from io import BytesIO
    from flask import Response
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Students"
    
    # Styles
    header_font = Font(bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    data_font = Font(size=10)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='top', wrap_text=True)
    
    # Column width mapping based on field type
    col_width_map = {
        'Student ID': 14,
        'Surname': 15,
        'First Name': 15,
        'Middle Name': 15,
        'Gender': 10,
        'Class': 14,
        'Date of Birth': 14,
        'Age': 8,
        'Religion': 12,
        'Home Address': 35,
        'Hobbies': 30,
        'Parent Phone': 15,
    }
    
    # Title row
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(fields) + 1)
    title_cell = ws['A1']
    title_cell.value = 'STUDENTS LIST'
    title_cell.font = Font(bold=True, size=14)
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 25
    
    # Empty row
    ws.row_dimensions[2].height = 10
    
    # Header row (row 3)
    header_row = 3
    ws.cell(row=header_row, column=1, value='S/N')
    ws.cell(row=header_row, column=1).font = header_font
    ws.cell(row=header_row, column=1).fill = header_fill
    ws.cell(row=header_row, column=1).border = border
    ws.cell(row=header_row, column=1).alignment = center_align
    ws.column_dimensions['A'].width = 6
    
    for col, field in enumerate(fields, 2):
        cell = ws.cell(row=header_row, column=col, value=field)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = center_align
        
        # Set column width
        col_letter = get_column_letter(col)
        ws.column_dimensions[col_letter].width = col_width_map.get(field, 15)
    
    ws.row_dimensions[header_row].height = 20
    
    # Data rows
    for idx, student in enumerate(student_data, 1):
        row = idx + 3
        
        # S/N column
        sn_cell = ws.cell(row=row, column=1, value=idx)
        sn_cell.border = border
        sn_cell.alignment = Alignment(horizontal='center', vertical='top')
        sn_cell.font = data_font
        
        # Track max lines needed for this row
        max_lines = 1
        
        for col, field in enumerate(fields, 2):
            value = str(student.get(field, '') or '')
            cell = ws.cell(row=row, column=col, value=value)
            cell.border = border
            cell.alignment = left_align
            cell.font = data_font
            
            # Estimate lines needed based on column width and text length
            col_width = col_width_map.get(field, 15)
            chars_per_line = int(col_width * 1.2)  # Approximate chars that fit
            if chars_per_line > 0 and len(value) > 0:
                lines_needed = max(1, -(-len(value) // chars_per_line))  # Ceiling division
                # Also count actual newlines in the text
                lines_needed = max(lines_needed, value.count('\n') + 1)
                max_lines = max(max_lines, lines_needed)
        
        # Set row height based on content (15 points per line, minimum 20)
        row_height = max(20, max_lines * 15)
        ws.row_dimensions[row].height = row_height
    
    # Footer row
    footer_row = len(student_data) + 4
    ws.cell(row=footer_row, column=1, value=f'Total: {len(student_data)} students')
    ws.cell(row=footer_row, column=1).font = Font(bold=True, size=10)
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return Response(
        output.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename=students_export.xlsx'}
    )


def export_students_word(student_data, fields):
    """Export students to Word format with selected fields - wraps text and adjusts row heights"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from io import BytesIO
    from flask import Response
    
    doc = Document()
    
    # Set page to landscape for more space
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # Title
    title = doc.add_heading('STUDENTS LIST', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Column width mapping (in inches)
    col_width_map = {
        'S/N': 0.5,
        'Student ID': 1.0,
        'Surname': 1.2,
        'First Name': 1.2,
        'Middle Name': 1.2,
        'Gender': 0.7,
        'Class': 1.0,
        'Date of Birth': 1.0,
        'Age': 0.5,
        'Religion': 0.9,
        'Home Address': 2.0,
        'Hobbies': 1.8,
        'Parent Phone': 1.1,
    }
    
    # Create table with S/N + selected fields
    all_headers = ['S/N'] + fields
    table = doc.add_table(rows=1, cols=len(all_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for i, header in enumerate(all_headers):
        width = col_width_map.get(header, 1.0)
        table.columns[i].width = Inches(width)
    
    # Style header row
    header_row = table.rows[0]
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    header_row.height = Pt(25)
    
    for i, header in enumerate(all_headers):
        cell = header_row.cells[i]
        cell.text = header
        
        # Bold and center header text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        
        # Set header background color (blue)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = None  # Reset to let white show
        # Set white text
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for idx, student in enumerate(student_data, 1):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        
        # Calculate row height based on content
        max_lines = 1
        row_data = [str(idx)]
        
        for field in fields:
            value = str(student.get(field, '') or '')
            row_data.append(value)
            
            # Estimate lines needed
            col_width = col_width_map.get(field, 1.0)
            chars_per_line = int(col_width * 12)  # Approximate
            if chars_per_line > 0 and len(value) > 0:
                words = value.split()
                current_line_len = 0
                lines = 1
                for word in words:
                    if current_line_len + len(word) + 1 > chars_per_line:
                        lines += 1
                        current_line_len = len(word)
                    else:
                        current_line_len += len(word) + 1
                max_lines = max(max_lines, lines)
        
        # Set minimum row height based on content
        row.height = Pt(max(20, max_lines * 14))
        
        # Fill cells
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = value
            
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(9)
            
            # Alternate row coloring
            if idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'Total: {len(student_data)} students')
    footer.runs[0].bold = True
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return Response(
        output.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename=students_export.docx'}
    )


def export_students_pdf(student_data, fields):
    """Export students to PDF format with selected fields - wraps text and adjusts row heights"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from io import BytesIO
    from flask import Response
    
    output = BytesIO()
    
    # Use landscape for more space
    page_size = landscape(A4)
    doc = SimpleDocTemplate(output, pagesize=page_size, leftMargin=10*mm, rightMargin=10*mm,
                           topMargin=10*mm, bottomMargin=10*mm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Create custom styles for cells
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontSize=8,
        leading=10,  # Line height
        alignment=0,  # Left align
    )
    
    header_style = ParagraphStyle(
        'HeaderStyle',
        parent=styles['Normal'],
        fontSize=9,
        leading=11,
        alignment=1,  # Center align
        textColor=colors.white,
        fontName='Helvetica-Bold',
    )
    
    # Title
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, alignment=1, spaceAfter=10)
    elements.append(Paragraph('STUDENTS LIST', title_style))
    elements.append(Spacer(1, 5*mm))
    
    # Column width mapping (in mm)
    col_width_map = {
        'S/N': 10*mm,
        'Student ID': 22*mm,
        'Surname': 28*mm,
        'First Name': 28*mm,
        'Middle Name': 25*mm,
        'Gender': 15*mm,
        'Class': 25*mm,
        'Date of Birth': 22*mm,
        'Age': 12*mm,
        'Religion': 20*mm,
        'Home Address': 45*mm,
        'Hobbies': 40*mm,
        'Parent Phone': 25*mm,
    }
    
    # Build column widths list
    all_headers = ['S/N'] + fields
    col_widths = [col_width_map.get(h, 25*mm) for h in all_headers]
    
    # Adjust widths to fit page if needed
    total_width = sum(col_widths)
    available_width = page_size[0] - 20*mm
    if total_width > available_width:
        scale = available_width / total_width
        col_widths = [w * scale for w in col_widths]
    
    # Build table data with Paragraphs for text wrapping
    table_data = []
    
    # Header row
    header_row = [Paragraph(h, header_style) for h in all_headers]
    table_data.append(header_row)
    
    # Data rows
    for idx, student in enumerate(student_data, 1):
        row = [Paragraph(str(idx), cell_style)]
        for field in fields:
            value = str(student.get(field, '') or '')
            row.append(Paragraph(value, cell_style))
        table_data.append(row)
    
    # Create table - let ReportLab calculate row heights automatically
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    
    # Style the table
    style_commands = [
        # Header styling
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        
        # Data styling
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        
        # Alignment
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Header centered
        ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # S/N column centered
        ('ALIGN', (1, 1), (-1, -1), 'LEFT'),   # Data left aligned
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),   # Top align all cells
        
        # Padding
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        
        # Grid
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#4472C4')),
        
        # Header bottom border
        ('LINEBELOW', (0, 0), (-1, 0), 1.5, colors.HexColor('#4472C4')),
    ]
    
    # Add alternating row colors
    for i in range(1, len(table_data)):
        if i % 2 == 0:
            style_commands.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor('#F5F5F5')))
    
    table.setStyle(TableStyle(style_commands))
    elements.append(table)
    
    # Footer
    elements.append(Spacer(1, 5*mm))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=10, fontName='Helvetica-Bold')
    elements.append(Paragraph(f'Total: {len(student_data)} students', footer_style))
    
    doc.build(elements)
    output.seek(0)
    
    return Response(
        output.getvalue(),
        mimetype='application/pdf',
        headers={'Content-Disposition': 'attachment; filename=students_export.pdf'}
    )


def export_students_image(student_data, fields):
    """Render the students list to a modern, high-resolution PNG."""
    from PIL import Image, ImageDraw, ImageFont
    from io import BytesIO
    from flask import Response
    from datetime import datetime

    S = 3  # supersampling for a crisp, high-DPI image

    def fnt(size, bold=False):
        path = "/usr/share/fonts/truetype/dejavu/DejaVuSans%s.ttf" % ("-Bold" if bold else "")
        try:
            return ImageFont.truetype(path, size * S)
        except Exception:
            return ImageFont.load_default()

    body = fnt(12)
    body_b = fnt(12, True)
    head_f = fnt(12, True)
    title_f = fnt(22, True)
    sub_f = fnt(11)
    foot_f = fnt(10)

    # Brand colours
    GREEN = (13, 106, 78)
    GREEN_DK = (6, 78, 54)
    HEAD = (10, 86, 64)
    TEXT = (31, 41, 55)
    MUTED = (107, 114, 128)
    ZEBRA = (249, 250, 251)
    LINE = (229, 231, 235)
    WHITE = (255, 255, 255)

    try:
        from models import SchoolSettings
        school_name = SchoolSettings.get('school_name', 'PosyHub') or 'PosyHub'
    except Exception:
        school_name = 'PosyHub'

    tmp = ImageDraw.Draw(Image.new('RGB', (1, 1)))

    if not student_data:
        img = Image.new('RGB', (560 * S, 160 * S), WHITE)
        d = ImageDraw.Draw(img)
        d.text((40 * S, 60 * S), "No students to export", fill=TEXT, font=title_f)
        out = BytesIO(); img.save(out, format='PNG'); out.seek(0)
        return Response(out.getvalue(), mimetype='image/png',
                        headers={'Content-Disposition': 'attachment; filename=students_export.png'})

    pad = 26 * S
    cpx = 12 * S
    cpy = 9 * S
    header_h = 42 * S
    min_row_h = 38 * S
    line_bbox = tmp.textbbox((0, 0), "Ay", font=body)
    line_h = (line_bbox[3] - line_bbox[1]) + 6 * S

    all_headers = ['S/N'] + fields
    width_map = {'S/N': 56, 'Student ID': 130, 'Surname': 150, 'First Name': 150,
                 'Middle Name': 150, 'Gender': 86, 'Age': 70, 'Class': 130,
                 'Religion': 120, 'Date of Birth': 130, 'Stream': 120,
                 'Home Address': 230, 'Hobbies': 210, 'Parent Phone': 140}
    col_widths = [width_map.get(h, 150) * S for h in all_headers]

    # Wrap + measure rows
    wrapped_data, row_heights = [], []
    for idx, student in enumerate(student_data):
        row, max_lines = [], 1
        for i, header in enumerate(all_headers):
            value = str(idx + 1) if i == 0 else str(student.get(fields[i - 1], '') or '')
            lines = wrap_text(tmp, value, col_widths[i] - cpx * 2, body)
            row.append(lines)
            max_lines = max(max_lines, len(lines))
        wrapped_data.append(row)
        row_heights.append(max(min_row_h, max_lines * line_h + cpy * 2))

    table_w = sum(col_widths)
    band_h = 88 * S
    total_w = table_w + pad * 2
    total_h = band_h + header_h + sum(row_heights) + 56 * S

    img = Image.new('RGB', (int(total_w), int(total_h)), WHITE)
    d = ImageDraw.Draw(img)

    # Header band with a vertical gradient
    for yy in range(band_h):
        t = yy / band_h
        col = tuple(int(GREEN[k] + (GREEN_DK[k] - GREEN[k]) * t) for k in range(3))
        d.line([(0, yy), (total_w, yy)], fill=col)
    d.text((pad, 20 * S), school_name, fill=WHITE, font=title_f)
    d.text((pad, 56 * S),
           "Students List  ·  %d student%s  ·  %s" % (
               len(student_data), '' if len(student_data) == 1 else 's',
               datetime.now().strftime('%d %b %Y')),
           fill=(220, 240, 232), font=sub_f)

    # Column header row
    y = band_h
    d.rectangle([0, y, total_w, y + header_h], fill=HEAD)
    x = pad
    for i, header in enumerate(all_headers):
        bb = d.textbbox((0, 0), header, font=head_f)
        d.text((x + cpx, y + (header_h - (bb[3] - bb[1])) // 2 - bb[1]), header, fill=WHITE, font=head_f)
        x += col_widths[i]
    y += header_h

    # Data rows (zebra + subtle separators)
    for idx, (row, rh) in enumerate(zip(wrapped_data, row_heights)):
        if idx % 2 == 1:
            d.rectangle([pad, y, pad + table_w, y + rh], fill=ZEBRA)
        x = pad
        for i, lines in enumerate(row):
            ty = y + cpy
            f = body_b if i == 0 else body
            fill = GREEN if i == 0 else TEXT
            for ln in lines:
                d.text((x + cpx, ty), ln, fill=fill, font=f)
                ty += line_h
            x += col_widths[i]
        d.line([(pad, y + rh), (pad + table_w, y + rh)], fill=LINE, width=max(1, S // 2))
        y += rh

    # Footer
    d.text((pad, y + 16 * S),
           "Generated by %s · %s" % (school_name, datetime.now().strftime('%d %b %Y %H:%M')),
           fill=MUTED, font=foot_f)

    out = BytesIO()
    img.save(out, format='PNG')
    out.seek(0)
    return Response(out.getvalue(), mimetype='image/png',
                    headers={'Content-Disposition': 'attachment; filename=students_export.png'})


def wrap_text(draw, text, max_width, font):
    """Wrap text to fit within max_width, returns list of lines"""
    if not text:
        return ['']
    
    text = str(text).strip()
    if not text:
        return ['']
    
    # Check if text already fits on one line
    bbox = draw.textbbox((0, 0), text, font=font)
    if bbox[2] - bbox[0] <= max_width:
        return [text]
    
    # Split into words
    words = text.split()
    if not words:
        return ['']
    
    lines = []
    current_line = []
    
    for word in words:
        # Check if adding this word exceeds the width
        test_line = ' '.join(current_line + [word])
        bbox = draw.textbbox((0, 0), test_line, font=font)
        
        if bbox[2] - bbox[0] <= max_width:
            current_line.append(word)
        else:
            # If current line has words, save it
            if current_line:
                lines.append(' '.join(current_line))
                current_line = [word]
            else:
                # Word is too long, need to break it
                lines.append(break_long_word(draw, word, max_width, font))
                current_line = []
    
    # Don't forget the last line
    if current_line:
        lines.append(' '.join(current_line))
    
    return lines if lines else ['']


def break_long_word(draw, word, max_width, font):
    """Break a long word that doesn't fit, adding ellipsis if needed"""
    for i in range(len(word), 0, -1):
        truncated = word[:i]
        bbox = draw.textbbox((0, 0), truncated, font=font)
        if bbox[2] - bbox[0] <= max_width:
            if i < len(word):
                # Try to fit with ellipsis
                with_ellipsis = word[:max(1, i-2)] + '..'
                bbox2 = draw.textbbox((0, 0), with_ellipsis, font=font)
                if bbox2[2] - bbox2[0] <= max_width:
                    return with_ellipsis
            return truncated
    return word[:1] if word else ''
