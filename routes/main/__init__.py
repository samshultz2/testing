"""
Main routes for dashboard and general pages
"""
import re
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, session, abort
from utils.helpers import get_active_term, get_active_session, safe_redirect
from utils.web_exports import xlsx_response, pdf_response
from datetime import date, timedelta
from models import (
    db, Student, ParentContact, StudentEnrollment, ClassArmAssignment, Attendance, 
    Week, ClassArm, Subject, SchoolClass
)
from utils.access_control import (
    login_required, admin_required, central_admin_required, is_admin, is_teacher,
    get_accessible_class_ids, is_sss3_form_teacher, page_can_write
)
from utils.audit import log_action
from utils.helpers import RELIGIONS, parse_date, FlashMessages, WAEC_SUBJECTS, STREAMS, STREAM_WAEC_SUBJECTS
from utils.search import like_term, escape_like
from sqlalchemy import extract, func, nullslast
from sqlalchemy.orm import joinedload
from urllib.parse import urlparse


def _safe_next(target, fallback):
    """Return target only if it is a same-site relative path, else fallback."""
    if target:
        parsed = urlparse(target)
        if not parsed.scheme and not parsed.netloc and target.startswith('/'):
            return target
    return fallback

main_bp = Blueprint('main', __name__)






# Selectable dashboard widgets: (key, label, category, default-on).
DASHBOARD_WIDGETS = [
    ('kpi', 'Student KPIs', 'Students', True),
    ('charts', 'Gender / Stream / Age charts', 'Students', True),
    ('class_religion', 'Class enrolment & Religion', 'Students', True),
    ('people', 'Birthdays, Recent students, Activity', 'Students', True),
    ('attendance_trend', 'Attendance trend', 'Academics', True),
    ('exams', 'WAEC / JAMB / Mock snapshots', 'Academics', True),
    ('finance', 'Finance — fees this term', 'Finance', True),
    ('sales', 'Sales — today & term', 'Finance', False),
    ('hr', 'Staff / HR', 'Operations', False),
    ('cbt', 'CBT activity', 'Academics', False),
    ('library', 'Library', 'Operations', False),
]
DASHBOARD_DEFAULTS = [k for k, _, _, d in DASHBOARD_WIDGETS if d]


# Each widget needs access to a module to be shown — a user only ever sees the
# cards whose module they're permitted to access, regardless of their saved
# widget preferences.
WIDGET_MODULE = {
    'kpi': 'students', 'charts': 'students', 'class_religion': 'students',
    'people': 'students', 'attendance_trend': 'attendance',
    'exams': 'external_exams', 'finance': 'finance', 'sales': 'sales',
    'hr': 'hr', 'cbt': 'cbt', 'library': 'library',
}


def permitted_widgets():
    """Widget keys the current user is allowed to see (by module permission)."""
    from utils.access_control import can_access_module
    return {k for k in WIDGET_MODULE
            if can_access_module(WIDGET_MODULE[k])}


def enabled_widgets():
    """Widget keys enabled for the user: their preference ∩ what they may see."""
    from utils.access_control import get_current_user
    chosen = None
    try:
        user = get_current_user()
        if user and user.dashboard_widgets is not None:
            chosen = set(user.dashboard_widgets)
    except Exception:
        chosen = None
    if chosen is None:
        sp = session.get('dashboard_prefs')
        chosen = set(sp) if sp is not None else set(DASHBOARD_DEFAULTS)
    # Never show a card the user has no permission for.
    return chosen & permitted_widgets()








def _ser_student_brief(s):
    return {'id': s.id, 'student_id': s.student_id, 'full_name': s.full_name,
            'gender': s.gender, 'url': url_for('main.view_student', student_id=s.id)}




def _dashboard_urls():
    """Server-resolved link targets so the React dashboard never hardcodes paths."""
    return {
        'customize': url_for('main.dashboard_customize'),
        'week_grid': url_for('attendance.week_grid'),
        'bulk_entry': url_for('subjects.bulk_entry'),
        'students_list': url_for('main.students_list'),
        'audit_log': url_for('main.audit_log'),
        'add_student': url_for('main.add_student'),
        'scores_entry': url_for('subjects.scores_entry'),
        'scan_waec': url_for('results.scan_waec'),
        'analytics_hub': url_for('results.analytics_hub'),
        'readiness': url_for('results.readiness'),
        'jamb_list': url_for('results.jamb_list'),
        'waec_list': url_for('results.waec_list'),
        'mock_index': url_for('mock_jamb.index'),
        'classes_list': url_for('academics.classes_list'),
        'weekly_summary': url_for('attendance.weekly_summary'),
        'mark_attendance': url_for('attendance.attendance_app', tab='mark'),
        'record_payment': url_for('finance.record_payment'),
        'defaulters': url_for('finance.defaulters'),
        'finance_overview': url_for('finance.dashboard'),
        'new_sale': url_for('sales.new_sale'),
        'send_message': url_for('comms.compose'),
    }


def _quick_actions():
    """Permission-aware dashboard shortcuts. Each role only sees actions it can
    actually perform — a bursar gets Record Payment / Defaulters, a teacher gets
    Mark Attendance / Enter Scores — so nobody is shown a button that 403s on
    click. Write actions are suppressed for view-only users."""
    from utils.access_control import can_access_module, module_level, is_read_only
    u = _dashboard_urls()
    read_only = is_read_only()

    def can_edit(mod):
        return module_level(mod) == 'edit' and not read_only

    acts = []
    if can_edit('students'):
        acts.append(('Add Student', 'fa-user-plus', u['add_student'], 'btn-primary'))
    if can_edit('attendance'):
        acts.append(('Mark Attendance', 'fa-clipboard-check', u['mark_attendance'], 'btn-success'))
    if can_edit('results'):
        acts.append(('Enter Scores', 'fa-edit', u['scores_entry'], 'btn-info'))
    if can_edit('finance'):
        acts.append(('Record Payment', 'fa-cash-register', u['record_payment'], 'btn-primary'))
    if can_access_module('finance'):
        acts.append(('Defaulters', 'fa-user-clock', u['defaulters'], 'btn-outline'))
    if can_edit('sales'):
        acts.append(('New Sale', 'fa-cart-plus', u['new_sale'], 'btn-secondary'))
    if can_edit('communication'):
        acts.append(('Send Message', 'fa-paper-plane', u['send_message'], 'btn-secondary'))
    if can_access_module('external_exams'):
        acts.append(('Scan Result', 'fa-camera', u['scan_waec'], 'btn-secondary'))
    if can_access_module('results') or can_access_module('external_exams'):
        acts.append(('Analytics', 'fa-chart-pie', u['analytics_hub'], 'btn-outline'))
    if can_access_module('external_exams'):
        acts.append(('Readiness', 'fa-clipboard-check', u['readiness'], 'btn-outline'))
    return [{'label': l, 'icon': i, 'url': url, 'tone': t} for (l, i, url, t) in acts]


# Modules whose presence means the academic dashboard has something to show.
_ACADEMIC_MODULES = ('students', 'attendance', 'results', 'external_exams',
                     'academics', 'admissions', 'cbt', 'timetable', 'promotion', 'events')


def _home_focus():
    """Role-aware landing hint for the React dashboard. A finance-only staffer
    (e.g. a bursar with no academic modules) gets a finance-first home with a
    prominent shortcut into the Finance workspace, instead of an empty academic
    grid. `/` itself still renders for everyone (it's the common home)."""
    from utils.access_control import is_admin, can_access_module
    if is_admin():
        return None
    if can_access_module('finance') and not any(can_access_module(m) for m in _ACADEMIC_MODULES):
        return 'finance'
    return None


def _floats(d):
    """Coerce a dict of numeric values to JSON-safe floats (e.g. Decimals)."""
    return None if d is None else {k: float(v) for k, v in d.items()}


def dashboard_payload():
    """Assemble all dashboard widget data as a JSON-serialisable dict.

    Cross-module widgets are computed only when enabled (preference ∩ module
    permission), so a user never receives data for a module they can't access.
    Student/attendance helpers are branch- or teacher-scoped internally."""
    active_session = get_active_session()
    active_term = get_active_term()
    enabled_set = enabled_widgets()
    permitted_set = permitted_widgets()
    enabled = sorted(enabled_set)
    tscope = _teacher_scope()   # teacher: limit student stats to their classes

    active_enrollments, total_classes, class_stats = _dash_class_stats(active_term, tscope)
    birthdays_today, birthdays_week = _dash_birthdays(tscope)
    recent_students = _student_scope(
        Student.query.filter_by(is_active=True), tscope).order_by(
        Student.created_at.desc()).limit(5).all()
    recent_activity = _dash_recent_activity()
    announcements = _dash_announcements()
    # Which of these the viewer has already acknowledged (for needs_ack notices).
    _acked_ann_ids = set()
    _ann_ack_ids = [a.id for a in announcements if a.needs_ack]
    if _ann_ack_ids:
        from models import AnnouncementAck
        from utils.access_control import get_current_user
        _u = get_current_user()
        if _u is not None:
            _acked_ann_ids = {r[0] for r in db.session.query(AnnouncementAck.announcement_id)
                              .filter(AnnouncementAck.user_id == _u.id,
                                      AnnouncementAck.announcement_id.in_(_ann_ack_ids)).all()}
    sales = _dash_sales(active_term) if 'sales' in enabled else None

    from utils.access_control import can_access_module
    tcs = _teacher_today(active_term, tscope)
    teacher_classes = None if tcs is None else [
        {**c,
         'mark_url': url_for('attendance.attendance_app', term_id=active_term.id,
                             assignment_id=c['id'], tab='mark'),
         'week_url': url_for('attendance.attendance_app', term_id=active_term.id,
                             assignment_id=c['id'], tab='week')}
        for c in tcs]

    user = session.get('user')
    return dict(
        today=date.today().isoformat(),
        user_name=(user.split()[0] if user else ''),
        active_session={'id': active_session.id, 'name': active_session.name} if active_session else None,
        active_term={'id': active_term.id, 'name': active_term.name} if active_term else None,
        enabled=enabled,
        permitted=sorted(permitted_set),
        widget_catalog=[{'key': k, 'label': label, 'group': cat,
                         'permitted': k in permitted_set, 'enabled': k in enabled_set}
                        for k, label, cat, _ in DASHBOARD_WIDGETS],
        announcements=[{'id': a.id, 'title': a.title, 'body': a.body, 'category': a.category,
                        'is_pinned': bool(a.is_pinned), 'needs_ack': bool(a.needs_ack),
                        'acked': a.id in _acked_ann_ids,
                        'ack_url': url_for('comms.ack_announcement', ann_id=a.id),
                        'attachment': _ann_attachment(a.attachment_id)}
                       for a in announcements],
        recent_students=[_ser_student_brief(s) for s in recent_students],
        active_enrollments=active_enrollments,
        total_classes=total_classes,
        class_stats=class_stats,
        attendance_stats=_dash_attendance_stats(active_term, tscope),
        birthdays_today=[{'full_name': s.full_name, 'age': s.age} for s in birthdays_today],
        birthdays_week=[{'full_name': s.full_name,
                         'date_label': s.date_of_birth.strftime('%d %b')} for s in birthdays_week],
        age_distribution=_dash_age_distribution(tscope),
        religion_stats=_dash_religion_stats(tscope),
        total_subjects=Subject.query.filter_by(is_active=True).count(),
        total_school_classes=SchoolClass.query.filter_by(is_active=True).count(),
        stream_dist=_dash_stream_distribution(tscope),
        jamb_snapshot=_dash_jamb_snapshot(tscope) if 'exams' in enabled else None,
        waec_snapshot=_dash_waec_snapshot(tscope) if 'exams' in enabled else None,
        mock_snapshot=_dash_mock_snapshot(tscope) if 'exams' in enabled else None,
        attendance_trend=_dash_attendance_trend(active_term, tscope) if 'attendance_trend' in enabled else [],
        recent_activity=[{'action': a.action, 'detail': a.detail, 'user': a.user,
                          'created_at': a.created_at.strftime('%d %b %H:%M') if a.created_at else ''}
                         for a in recent_activity],
        # Cross-module widgets (computed only when enabled).
        finance_stat=_floats(_dash_finance(active_term)) if 'finance' in enabled else None,
        sales_stat=({**sales, 'today': float(sales['today'])} if sales else None),
        hr_stat=_dash_hr() if 'hr' in enabled else None,
        cbt_stat=_dash_cbt() if 'cbt' in enabled else None,
        library_stat=_dash_library() if 'library' in enabled else None,
        teacher_classes=teacher_classes,
        can_results=can_access_module('results'),
        quick_actions=_quick_actions(),
        home_focus=_home_focus(),
        urls=_dashboard_urls(),
        **_dash_student_counts(tscope)
    )


def _teacher_today(active_term, tscope):
    """For a teacher, their classes in the active term with per-class quick
    actions (mark attendance, mark a week). None for admins / non-teachers."""
    if tscope is None or not active_term:
        return None
    aidset = set(tscope[0])
    if not aidset:
        return []
    rows = []
    for a in (ClassArmAssignment.query
              .filter(ClassArmAssignment.id.in_(aidset),
                      ClassArmAssignment.term_id == active_term.id).all()):
        rows.append({
            'id': a.id,
            'name': a.display_name,
            'count': StudentEnrollment.query.filter_by(
                class_arm_assignment_id=a.id, is_active=True).count(),
        })
    rows.sort(key=lambda r: r['name'])
    return rows






def _dash_finance(active_term):
    """Fees collected/expenses/net for the active term (branch-scoped)."""
    if not active_term:
        return None
    try:
        from utils.branch_scope import scope_query
        from models import FeePayment
        from models.models_finance import Expense
        collected = sum(p.amount for p in scope_query(
            FeePayment.query.filter_by(term_id=active_term.id), FeePayment).all())
        expenses = sum(e.amount for e in scope_query(
            Expense.query.filter_by(term_id=active_term.id), Expense).all())
        return {'collected': collected, 'expenses': expenses, 'net': collected - expenses}
    except Exception:
        return None


def _dash_sales(active_term):
    """Sales totals today and (active term window) — branch-scoped."""
    try:
        from utils.branch_scope import scope_query
        from models import Sale
        today = date.today()
        rows = scope_query(Sale.query, Sale).all()
        today_total = sum(s.total for s in rows if getattr(s, 'created_at', None)
                          and s.created_at.date() == today)
        return {'today': today_total, 'count_today': sum(
            1 for s in rows if getattr(s, 'created_at', None) and s.created_at.date() == today)}
    except Exception:
        return None


def _dash_hr():
    """Active staff headcount."""
    try:
        from utils.branch_scope import scope_query
        from models import StaffMember
        rows = scope_query(StaffMember.query.filter_by(is_active=True), StaffMember).all()
        return {'total': len(rows),
                'teaching': sum(1 for s in rows if (getattr(s, 'staff_type', '') or '').lower().startswith('teach'))}
    except Exception:
        return None


def _dash_cbt():
    """CBT exam + attempt counts (branch-scoped on the attempt's student)."""
    try:
        from models import CBTExam
        from models.models_cbt import CBTAttempt
        published = CBTExam.query.filter_by(is_published=True).count()
        attempts = CBTAttempt.query.count()
        return {'published': published, 'attempts': attempts}
    except Exception:
        return None


def _dash_library():
    """Library books + active/overdue loans."""
    try:
        from models import Book, BookLoan
        books = Book.query.count()
        active = BookLoan.query.filter_by(returned_at=None).count() if hasattr(BookLoan, 'returned_at') else 0
        return {'books': books, 'on_loan': active}
    except Exception:
        return None


def _teacher_scope():
    """For a teacher (non-admin), the class assignments + student ids they may
    see on their dashboard; None for admins/non-teachers (full branch view)."""
    from utils.access_control import is_admin, is_teacher, get_accessible_class_ids
    if is_admin() or not is_teacher():
        return None
    aids = get_accessible_class_ids()
    if not aids:
        return ([], [])
    sids = [r[0] for r in db.session.query(StudentEnrollment.student_id)
            .filter(StudentEnrollment.class_arm_assignment_id.in_(aids),
                    StudentEnrollment.is_active == True).distinct().all()]
    return (aids, sids)


def _dash_student_counts(tscope=None):
    """Active-student headcounts (branch-scoped, or teacher's classes only)."""
    from utils.branch_scope import scope_query
    def sq(query):
        if tscope is not None:
            # restrict to the teacher's own students ([-1] => match nothing)
            return query.filter(Student.id.in_(tscope[1] or [-1]))
        return scope_query(query, Student)
    # Current (non-graduate) students only; graduates are counted separately.
    current = Student.query.filter_by(is_active=True).filter(
        db.or_(Student.is_graduated.is_(False), Student.is_graduated.is_(None)))
    return {
        'total_students': sq(current).count(),
        'male_students': sq(current.filter(Student.gender == 'Male')).count(),
        'female_students': sq(current.filter(Student.gender == 'Female')).count(),
        'graduates_count': sq(Student.query.filter_by(is_active=True, is_graduated=True)).count(),
    }


def _dash_class_stats(active_term, tscope=None):
    """Per-class enrolment breakdown for the active term (teacher: own classes)."""
    from utils.branch_scope import scope_query
    if not active_term:
        return 0, 0, []
    assignments = scope_query(
        ClassArmAssignment.query.filter_by(term_id=active_term.id),
        ClassArmAssignment).all()
    if tscope is not None:
        aidset = set(tscope[0])
        assignments = [a for a in assignments if a.id in aidset]
    active_enrollments = StudentEnrollment.query.filter(
        StudentEnrollment.class_arm_assignment_id.in_([a.id for a in assignments] or [-1]),
        StudentEnrollment.is_active == True).count()
    class_stats = []
    for assignment in assignments:
        enrollment_count = StudentEnrollment.query.filter_by(
            class_arm_assignment_id=assignment.id, is_active=True).count()
        male_count = StudentEnrollment.query.join(Student).filter(
            StudentEnrollment.class_arm_assignment_id == assignment.id,
            StudentEnrollment.is_active == True,
            Student.gender == 'Male').count()
        class_stats.append({
            'name': assignment.display_name,
            'total': enrollment_count,
            'male': male_count,
            'female': enrollment_count - male_count,
        })
    class_stats.sort(key=lambda x: x['name'])
    return active_enrollments, len(assignments), class_stats


def _dash_attendance_stats(active_term, tscope=None):
    """Today + term attendance percentages (branch-scoped, or teacher's classes)."""
    from utils.branch_scope import viewing_branch_id
    stats = {'today_present': 0, 'today_absent': 0, 'today_percentage': 0,
             'week_average': 0, 'term_average': 0}
    if not active_term:
        return stats
    today = date.today()
    branch_enr = None
    if tscope is not None:
        # Restrict to enrolments in the teacher's own class assignments.
        branch_enr = (db.session.query(StudentEnrollment.id)
                      .filter(StudentEnrollment.class_arm_assignment_id.in_(
                          tscope[0] or [-1])))
    else:
        bid = viewing_branch_id()
        if bid is not None:
            branch_enr = (db.session.query(StudentEnrollment.id)
                          .join(ClassArmAssignment,
                                StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id)
                          .filter(ClassArmAssignment.branch_id == bid))

    current_week = Week.query.filter(
        Week.term_id == active_term.id,
        Week.start_date <= today,
        Week.end_date >= today).first()
    if current_week:
        q = Attendance.query.filter_by(week_id=current_week.id, date=today)
        if branch_enr is not None:
            q = q.filter(Attendance.enrollment_id.in_(branch_enr))
        today_records = q.all()
        stats['today_present'] = sum(
            1 for a in today_records if a.morning_present or a.afternoon_present)
        stats['today_absent'] = sum(
            1 for a in today_records if not a.morning_present and not a.afternoon_present)
        total_today = len(today_records)
        if total_today > 0:
            present_marks = sum(
                (1 if a.morning_present else 0) + (1 if a.afternoon_present else 0)
                for a in today_records)
            stats['today_percentage'] = round((present_marks / (total_today * 2)) * 100, 1)

    term_weeks = Week.query.filter_by(term_id=active_term.id).all()
    if term_weeks:
        tq = Attendance.query.filter(Attendance.week_id.in_([w.id for w in term_weeks]))
        if branch_enr is not None:
            tq = tq.filter(Attendance.enrollment_id.in_(branch_enr))
        term_records = tq.all()
        if term_records:
            total_marks = sum(
                (1 if a.morning_present else 0) + (1 if a.afternoon_present else 0)
                for a in term_records)
            stats['term_average'] = round((total_marks / (len(term_records) * 2)) * 100, 1)
    return stats


def _student_scope(query, tscope):
    """Restrict a Student-based query to the teacher's own students; otherwise
    apply the normal branch scope. Used by every student-centric dashboard
    widget so a form teacher only sees their classes' students. Graduated
    students are excluded so the dashboard analytics are about current students."""
    from utils.branch_scope import scope_query
    query = query.filter(db.or_(Student.is_graduated.is_(False),
                                Student.is_graduated.is_(None)))
    if tscope is not None:
        return query.filter(Student.id.in_(tscope[1] or [-1]))
    return scope_query(query, Student)


def _viewer_student_scope(query):
    """Scope a Student query to what the CURRENT user may see: branch always,
    and a form teacher additionally to their own students. Use for any
    student list/search/export/stats so a teacher never sees other classes."""
    from utils.branch_scope import scope_query
    from utils.access_control import teacher_form_student_ids
    q = scope_query(query, Student)
    tids = teacher_form_student_ids()
    if tids is not None:
        q = q.filter(Student.id.in_(tids or [-1]))
    return q


def _dash_birthdays(tscope=None):
    """Students whose birthday is today, and within the next 6 days."""
    def sq(query):
        return _student_scope(query, tscope)
    today = date.today()
    birthdays_today = sq(Student.query.filter(
        Student.is_active == True,
        Student.date_of_birth != None,
        extract('month', Student.date_of_birth) == today.month,
        extract('day', Student.date_of_birth) == today.day)).all()
    birthdays_week = []
    for i in range(1, 7):
        check_date = today + timedelta(days=i)
        students = sq(Student.query.filter(
            Student.is_active == True,
            Student.date_of_birth != None,
            extract('month', Student.date_of_birth) == check_date.month,
            extract('day', Student.date_of_birth) == check_date.day)).all()
        for s in students:
            s.birthday_date = check_date
        birthdays_week.extend(students)
    return birthdays_today, birthdays_week


def _dash_age_distribution(tscope=None):
    """Active students bucketed by age band."""
    dist = {'0-10': 0, '11-13': 0, '14-16': 0, '17-19': 0, '20+': 0}
    students = _student_scope(Student.query.filter(
        Student.is_active == True, Student.date_of_birth != None), tscope).all()
    for student in students:
        age = student.age
        if age <= 10:
            dist['0-10'] += 1
        elif age <= 13:
            dist['11-13'] += 1
        elif age <= 16:
            dist['14-16'] += 1
        elif age <= 19:
            dist['17-19'] += 1
        else:
            dist['20+'] += 1
    return dist


def _dash_religion_stats(tscope=None):
    """Active-student counts grouped by religion."""
    rows = _student_scope(db.session.query(
        Student.religion, func.count(Student.id)).filter(
        Student.is_active == True, Student.religion != None),
        tscope).group_by(Student.religion).all()
    return {r[0]: r[1] for r in rows if r[0]}


def _dash_stream_distribution(tscope=None):
    """Active-student counts grouped by stream (None -> 'Unset')."""
    rows = _student_scope(db.session.query(
        Student.stream, func.count(Student.id)).filter(
        Student.is_active == True), tscope).group_by(Student.stream).all()
    return {(s or 'Unset'): c for s, c in rows}


def _dash_jamb_snapshot(tscope=None):
    """Latest-year JAMB summary (branch- or teacher-scoped), or None."""
    from utils.branch_scope import scope_by_student
    from models import JAMBResult
    jy = db.session.query(JAMBResult.exam_year).order_by(JAMBResult.exam_year.desc()).first()
    if not jy:
        return None
    q = JAMBResult.query.filter_by(exam_year=jy[0]).options(joinedload(JAMBResult.student))
    if tscope is not None:
        q = q.filter(JAMBResult.student_id.in_(tscope[1] or [-1]))
    else:
        q = scope_by_student(q, JAMBResult)
    rows = q.order_by(JAMBResult.total_score.desc()).all()
    scores = [r.total_score for r in rows]
    if not scores:
        return None
    top = rows[:5]
    return {
        'year': jy[0], 'count': len(scores),
        'mean': round(sum(scores) / len(scores), 1),
        'max': max(scores),
        'above_200': sum(1 for s in scores if s >= 200),
        'above_200_pct': round(sum(1 for s in scores if s >= 200) / len(scores) * 100),
        'distribution': {
            '0-149': sum(1 for s in scores if s < 150),
            '150-199': sum(1 for s in scores if 150 <= s < 200),
            '200-249': sum(1 for s in scores if 200 <= s < 250),
            '250-299': sum(1 for s in scores if 250 <= s < 300),
            '300+': sum(1 for s in scores if s >= 300),
        },
        'top': [{'name': r.student.full_name, 'score': r.total_score} for r in top],
    }


def _dash_waec_snapshot(tscope=None):
    """Latest-year WAEC summary (branch- or teacher-scoped), or None."""
    from utils.branch_scope import scope_by_student
    from models import WAECResult
    wy = db.session.query(WAECResult.exam_year).order_by(WAECResult.exam_year.desc()).first()
    if not wy:
        return None
    q = WAECResult.query.filter_by(exam_year=wy[0])
    if tscope is not None:
        q = q.filter(WAECResult.student_id.in_(tscope[1] or [-1]))
    else:
        q = scope_by_student(q, WAECResult)
    rows = q.all()
    if not rows:
        return None
    credit = {'A1', 'B2', 'B3', 'C4', 'C5', 'C6'}
    passes = sum(1 for r in rows if r.grade in credit)
    return {
        'year': wy[0], 'entries': len(rows),
        'students': len({r.student_id for r in rows}),
        'pass_rate': round(passes / len(rows) * 100, 1),
    }


def _dash_mock_snapshot(tscope=None):
    """Latest Mock JAMB exam average (branch- or teacher-scoped), or None."""
    from utils.branch_scope import viewing_branch_id
    from models.mock_jamb import MockJAMBExam, MockJAMBResult
    last_mock = MockJAMBExam.query.order_by(MockJAMBExam.exam_date.desc()).first()
    if not last_mock:
        return None
    res_q = last_mock.results
    if tscope is not None:
        res_q = res_q.filter(MockJAMBResult.student_id.in_(tscope[1] or [-1]))
    else:
        bid = viewing_branch_id()
        if bid is not None:
            res_q = res_q.join(Student, MockJAMBResult.student_id == Student.id) \
                         .filter(Student.branch_id == bid)
    ms = [r.total_score for r in res_q.all()]
    if not ms:
        return None
    return {'name': last_mock.display_name, 'count': len(ms),
            'mean': round(sum(ms) / len(ms), 1), 'max': max(ms)}


def _dash_attendance_trend(active_term, tscope=None):
    """Average attendance % for the last 8 weeks of the active term — branch-
    scoped, or restricted to the teacher's own class assignments."""
    from utils.branch_scope import viewing_branch_id
    if not active_term:
        return []
    branch_enr = None
    if tscope is not None:
        branch_enr = (db.session.query(StudentEnrollment.id)
                      .filter(StudentEnrollment.class_arm_assignment_id.in_(tscope[0] or [-1])))
    else:
        bid = viewing_branch_id()
        if bid is not None:
            branch_enr = (db.session.query(StudentEnrollment.id)
                          .join(ClassArmAssignment,
                                StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id)
                          .filter(ClassArmAssignment.branch_id == bid))
    weeks = Week.query.filter_by(term_id=active_term.id).order_by(Week.start_date).all()[-8:]
    trend = []
    for w in weeks:
        wq = Attendance.query.filter_by(week_id=w.id)
        if branch_enr is not None:
            wq = wq.filter(Attendance.enrollment_id.in_(branch_enr))
        recs = wq.all()
        if recs:
            marks = sum((1 if a.morning_present else 0) + (1 if a.afternoon_present else 0)
                        for a in recs)
            pct = round(marks / (len(recs) * 2) * 100, 1)
        else:
            pct = 0
        trend.append({'label': getattr(w, 'week_number', None) or w.start_date.strftime('%d %b'),
                      'pct': pct})
    return trend


def _dash_recent_activity():
    """Recent audit entries (central admins only — spans the whole system)."""
    from utils.branch_scope import is_central
    from models import AuditLog
    if is_admin() and is_central():
        return AuditLog.query.order_by(AuditLog.created_at.desc()).limit(6).all()
    return []


def _ann_attachment(att_id):
    """Attachment metadata (name + download url) for a dashboard announcement."""
    if not att_id:
        return None
    from models import CommAttachment
    att = db.session.get(CommAttachment, att_id)
    if not att:
        return None
    return {'name': att.original_name, 'size': att.human_size,
            'url': url_for('comms.download_attachment', att_id=att.id)}


def _dash_announcements():
    """Up to 5 active announcements for the dashboard banner."""
    try:
        from models import Announcement
        rows = Announcement.query.order_by(
            Announcement.is_pinned.desc(), Announcement.created_at.desc()).limit(15).all()
        return [a for a in rows if a.is_active][:5]
    except Exception:
        return []


# ============================================================================
# STUDENT ROUTES
# ============================================================================

def _students_query():
    """Active-student query with branch/section/teacher scope plus the standard
    list filters (search/gender/religion/stream/subject/class/arm) and sort —
    all read from request.args. Shared by the list page and the JSON API so
    scoping and filtering can never drift between them."""
    from utils.branch_scope import scope_query
    from utils.org_scope import scope_students
    search = request.args.get('search', '')
    gender = request.args.get('gender', '')
    religion = request.args.get('religion', '')
    stream = request.args.get('stream', '')
    subject = request.args.get('subject', '')
    house = request.args.get('house', '')
    boarding = request.args.get('boarding', '')
    class_id = request.args.get('class_id', '', type=int) or None
    arm_id = request.args.get('arm_id', '', type=int) or None
    sort_by = request.args.get('sort', 'surname')
    order = request.args.get('order', 'asc')

    # Branch + section/stream scope first. Graduated students are kept out of the
    # main list and its analytics — they live in the Graduates section instead.
    query = scope_query(Student.query.filter_by(is_active=True), Student)
    query = query.filter(db.or_(Student.is_graduated.is_(False),
                                Student.is_graduated.is_(None)))
    query = scope_students(query)
    active_term = get_active_term()

    # Teachers see only the students they are FORM teacher of — not every class
    # they merely teach a subject in. Other non-admin roles rely on the
    # branch/section/stream scope above.
    if is_teacher():
        from utils.access_control import teacher_form_student_ids
        form_ids = teacher_form_student_ids()   # set of their form-class student ids
        if form_ids:
            query = query.filter(Student.id.in_(form_ids))
        else:
            query = query.filter(Student.id == -1)   # not a form teacher of anyone

    if search:
        search_term = like_term(search)
        # Broaden search beyond name/ID: a registrar routinely looks a student up
        # by a parent's phone/name/email or by NIN / JAMB registration. Parent
        # contacts live in a child table, matched via a subquery so the OR stays a
        # single scoped query (home address stays out — it's encrypted at rest).
        parent_q = (db.session.query(ParentContact.student_id)
                    .filter(db.or_(
                        ParentContact.name.ilike(search_term, escape='\\'),
                        ParentContact.phone_number.ilike(search_term, escape='\\'),
                        ParentContact.email.ilike(search_term, escape='\\'))))
        query = query.filter(db.or_(
            Student.first_name.ilike(search_term, escape='\\'), Student.surname.ilike(search_term, escape='\\'),
            Student.middle_name.ilike(search_term, escape='\\'), Student.student_id.ilike(search_term, escape='\\'),
            Student.nin.ilike(search_term, escape='\\'),
            Student.jamb_reg_number.ilike(search_term, escape='\\'),
            Student.jamb_profile_code.ilike(search_term, escape='\\'),
            Student.id.in_(parent_q)))
    if gender:
        query = query.filter(Student.gender == gender)
    if religion:
        query = query.filter(Student.religion == religion)
    if stream:
        query = query.filter(Student.stream == stream)
    if house:
        query = query.filter(Student.house == house)
    if boarding:
        query = query.filter(Student.boarding_status == boarding)

    # WAEC-subject filter (SSS3 only), via a subquery so it composes with the joins.
    if subject:
        sss3_q = (db.session.query(StudentEnrollment.student_id)
                  .join(ClassArmAssignment,
                        StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id)
                  .join(SchoolClass, ClassArmAssignment.class_id == SchoolClass.id)
                  .filter(StudentEnrollment.is_active == True, SchoolClass.name == 'SSS3'))
        if active_term:
            sss3_q = sss3_q.filter(ClassArmAssignment.term_id == active_term.id)
        # Match the subject as a whole CSV token. Escape any LIKE wildcards in the
        # user-supplied subject so `%`/`_` can't turn these into match-everything
        # patterns; the surrounding `, ` / `%` separators stay real wildcards.
        _subj = escape_like(subject)
        query = query.filter(
            Student.id.in_(sss3_q),
            db.or_(Student.waec_subjects == subject,
                   Student.waec_subjects.ilike(f'{_subj}, %', escape='\\'),
                   Student.waec_subjects.ilike(f'%, {_subj}, %', escape='\\'),
                   Student.waec_subjects.ilike(f'%, {_subj}', escape='\\')))

    # Class/arm filter (any user), via a subquery so it composes cleanly.
    if class_id or arm_id:
        enr_q = (db.session.query(StudentEnrollment.student_id)
                 .join(ClassArmAssignment,
                       StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id)
                 .filter(StudentEnrollment.is_active == True))
        if active_term:
            enr_q = enr_q.filter(ClassArmAssignment.term_id == active_term.id)
        if class_id:
            enr_q = enr_q.filter(ClassArmAssignment.class_id == class_id)
        if arm_id:
            enr_q = enr_q.filter(ClassArmAssignment.arm_id == arm_id)
        query = query.filter(Student.id.in_(enr_q))

    # Whitelisted sort columns. `age` isn't a DB column (it's a derived
    # property), so map it onto date_of_birth with the direction inverted:
    # oldest student == earliest birth date.
    sort_map = {
        'surname': Student.surname, 'first_name': Student.first_name,
        'student_id': Student.student_id, 'created_at': Student.created_at,
    }
    if sort_by == 'age':
        col = Student.date_of_birth
        # 'desc' means oldest-first -> earliest DOB first -> ascending DOB.
        ordering = col.asc() if order == 'desc' else col.desc()
        return query.order_by(nullslast(ordering), Student.surname.asc())
    sort_column = sort_map.get(sort_by, Student.surname)
    return query.order_by(sort_column.desc() if order == 'desc' else sort_column.asc())


def _page_class_map(items):
    """{student_id: current class display name} for one page of students."""
    active_term = get_active_term()
    class_map = {}
    page_ids = [s.id for s in items]
    if active_term and page_ids:
        enrollments = (StudentEnrollment.query.join(ClassArmAssignment).filter(
            StudentEnrollment.student_id.in_(page_ids),
            ClassArmAssignment.term_id == active_term.id
        ).options(
            joinedload(StudentEnrollment.class_arm_assignment).joinedload(ClassArmAssignment.school_class),
            joinedload(StudentEnrollment.class_arm_assignment).joinedload(ClassArmAssignment.arm),
        ).all())
        for e in enrollments:
            class_map.setdefault(e.student_id, e.class_arm_assignment.display_name)
    return class_map




def _students_payload():
    """The students list as a JSON-serialisable dict — shared by the page
    (embedded for instant render) and /api/students (filter/page changes)."""
    page = request.args.get('page', 1, type=int)
    per_page = min(request.args.get('per_page', 20, type=int) or 20, 100)
    pg = _students_query().paginate(page=page, per_page=per_page, error_out=False)
    class_map = _page_class_map(pg.items)
    # Edit/delete follow the user's WRITE permission on the students module. A
    # form teacher with write access can manage their own class's students (the
    # routes enforce form-class scope); a view-only user gets a read-only list.
    can_manage = page_can_write()
    # Form teachers may add students too: the add route auto-enrols a teacher's
    # new student into their own form class, so it lands in their (scoped) list.
    can_add = can_manage
    students = [{
        'id': s.id,
        'student_id': s.student_id,
        'name': s.full_name,
        'first_name': s.first_name,
        'surname': s.surname,
        'gender': s.gender,
        'religion': s.religion,
        'stream': s.stream,
        'age': s.age,
        'is_graduated': bool(s.is_graduated),
        'current_class': class_map.get(s.id),
        'url': url_for('main.view_student', student_id=s.id),
        'edit_url': url_for('main.edit_student', student_id=s.id),
        'delete_url': url_for('main.delete_student', student_id=s.id),
        'graduate_url': url_for('promotion.unmark_graduate', student_id=s.id) if s.is_graduated
                        else url_for('promotion.mark_graduate', student_id=s.id),
    } for s in pg.items]
    classes = SchoolClass.query.filter_by(is_active=True).order_by(SchoolClass.level).all()
    arms = ClassArm.query.filter_by(is_active=True, is_default=False).order_by(ClassArm.name).all()
    # Distinct pastoral houses in use (branch-scoped), so the House filter only
    # offers values that actually exist — empty for schools that don't use houses.
    from utils.branch_scope import scope_query as _scope_q
    houses = sorted({h for (h,) in _scope_q(
        db.session.query(Student.house).filter(Student.house.isnot(None),
                                               Student.house != ''), Student).distinct().all() if h})
    return {
        'students': students,
        'page': pg.page, 'pages': pg.pages or 1, 'total': pg.total,
        'per_page': per_page, 'has_next': pg.has_next, 'has_prev': pg.has_prev,
        'applied': {
            'search': request.args.get('search', ''), 'gender': request.args.get('gender', ''),
            'religion': request.args.get('religion', ''), 'stream': request.args.get('stream', ''),
            'subject': request.args.get('subject', ''),
            'house': request.args.get('house', ''), 'boarding': request.args.get('boarding', ''),
            'class_id': request.args.get('class_id', '', type=int) or None,
            'arm_id': request.args.get('arm_id', '', type=int) or None,
            'sort': request.args.get('sort', 'surname'), 'order': request.args.get('order', 'asc'),
        },
        'filters': {
            'classes': [{'id': c.id, 'name': c.name} for c in classes],
            'arms': [{'id': a.id, 'name': a.name} for a in arms],
            'religions': list(RELIGIONS),
            'streams': list(STREAMS),
            'subjects': list(WAEC_SUBJECTS),
            'houses': houses,
        },
        'can_manage': can_manage,
        'can_add': can_add,
        'can_admin': is_admin(),   # admin-only bulk tools (add subject / delete)
        # Mass-assign gender/stream: admins and teachers (teachers are scoped to
        # their own form students server-side).
        'can_bulk': is_admin() or is_teacher(),
        # The SSS3-only WAEC subject filter/tools: admins + SSS3 form teachers.
        'can_sss3': is_admin() or is_sss3_form_teacher(),
        'add_url': url_for('main.add_student'),
        'import_url': url_for('main.import_students'),
        'enrolment': _student_form_options(with_enrolment=True)['enrolment'] if can_manage else None,
        'export_url': url_for('main.export_students_data'),
        'trash_url': url_for('main.students_trash'),
        'waec_by_stream_url': url_for('main.apply_stream_waec'),
        'bulk_delete_url': url_for('main.bulk_delete_students'),
        'bulk_stream_url': url_for('main.bulk_set_stream'),
        'bulk_gender_url': url_for('main.bulk_set_gender'),
        'bulk_subject_url': url_for('main.bulk_add_subject'),
    }




def _wants_json():
    """The React forms POST via fetch with this header; reply JSON to them and
    keep the classic flash+redirect behaviour for plain form submits."""
    return request.headers.get('X-Requested-With') == 'fetch' or request.is_json


RELATIONSHIPS = ['Father', 'Mother', 'Guardian', 'Sibling', 'Other']


def _student_form_options(with_enrolment=False):
    """Shared option lists for the add/edit forms; the enrolment block (add
    only) is class-scoped to the current user with their form class default."""
    opts = {
        'religions': RELIGIONS, 'streams': STREAMS, 'waec_subjects': WAEC_SUBJECTS,
        'stream_waec': STREAM_WAEC_SUBJECTS, 'relationships': RELATIONSHIPS,
        'genders': ['Male', 'Female'],
    }
    if with_enrolment:
        from utils.access_control import get_teacher_profile, filter_classes_for_user
        active_term = get_active_term()
        enrol = None
        if active_term:
            caas = ClassArmAssignment.query.filter_by(term_id=active_term.id).all()
            class_options = filter_classes_for_user(caas)
            default_id = None
            teacher = get_teacher_profile()
            if teacher and not is_admin():
                form_ids = teacher.form_class_ids or set()
                default_id = next((c.id for c in class_options if c.id in form_ids), None)
            enrol = {
                'term_label': active_term.full_name,
                'has_classes': bool(class_options),
                'default_id': default_id,
                'classes': [{'id': c.id, 'label': c.display_name} for c in class_options],
            }
        opts['enrolment'] = enrol
    return opts


def _blank_contact():
    return {'name': '', 'phone_number': '', 'email': '', 'relationship': 'Father'}




# =============================================================================
# PASTE IMPORT  —  create many students from copy-pasted text
# =============================================================================



def _student_or_redirect(student_id):
    """Load a student the current user may view, or return (None, response).
    Branch-scoped + a form teacher is limited to their own students."""
    student = db.get_or_404(Student, student_id)
    from utils.branch_scope import can_access_branch
    from utils.access_control import teacher_form_student_ids
    if not can_access_branch(student.branch_id):
        return None, ('branch', 'That student belongs to another branch.')
    tids = teacher_form_student_ids()
    if tids is not None and student.id not in tids:
        return None, ('teacher', 'You can only view your own students.')
    return student, None


def _fmt_date(d):
    return d.strftime('%d %b %Y') if d else None


def _student_attendance_summary(sid):
    """Compact attendance snapshot for the profile page: overall %, present /
    absent / late days, low-attendance warning, latest term, and a deep link to
    the full attendance profile. Returns None when the student has no attendance
    data or the attendance module can't be reached (never breaks the profile)."""
    try:
        from utils.attendance_profile import build_student_profile
        prof = build_student_profile(sid)
    except Exception:
        return None
    if not prof or not prof.get('terms'):
        return None
    overall = prof.get('overall') or {}
    latest = (prof.get('terms') or [None])[0]
    return {
        'percentage': overall.get('percentage', 0),
        'present_days': overall.get('full_days', 0),
        'late_days': overall.get('late_days', 0),
        'absent_days': overall.get('absent_days', 0),
        'terms': overall.get('terms', 0),
        'threshold': prof.get('threshold'),
        'warning': bool(prof.get('warning')),
        'latest_term': (latest or {}).get('term') if latest else None,
        'latest_percentage': (latest or {}).get('percentage') if latest else None,
        'url': url_for('attendance.attendance_app', tab='student', student_id=sid),
    }


def _student_comms_history(sid, limit=8):
    """Recent messages addressed to this student's parents (per-recipient), so
    the profile shows what the school has communicated. Read-only; returns an
    empty history if the comms module isn't present. Body is trimmed to a
    snippet — the full campaign lives in Communication."""
    try:
        from models.models_comms import Message, MessageRecipient
    except Exception:
        return {'count': 0, 'items': []}
    try:
        q = (db.session.query(MessageRecipient, Message)
             .join(Message, MessageRecipient.message_id == Message.id)
             .filter(MessageRecipient.student_id == sid)
             .order_by(MessageRecipient.created_at.desc()))
        total = q.count()
        rows = q.limit(limit).all()
    except Exception:
        return {'count': 0, 'items': []}
    items = []
    for rec, msg in rows:
        body = (rec.body or msg.body or '').strip()
        items.append({
            'title': msg.title or (msg.audience_label or 'Message'),
            'channel': msg.channel, 'status': rec.status,
            'snippet': (body[:120] + '…') if len(body) > 120 else body,
            'date': _fmt_date(rec.created_at.date() if rec.created_at else None),
            'sent': bool(rec.sent_at), 'read': bool(rec.read_at),
        })
    return {'count': total, 'items': items}


def _student_view_payload(student):
    """Everything the student detail page shows, JSON-serialisable."""
    enrollments = student.enrollments.join(ClassArmAssignment).order_by(
        ClassArmAssignment.term_id.desc()).all()
    waec_results = student.waec_results.all()
    jamb_results = student.jamb_results.order_by(None).all() if hasattr(student.jamb_results, 'order_by') else student.jamb_results.all()
    latest_jamb = student.jamb_results.first()
    contacts = student.parent_contacts.all()
    discipline = sorted(student.discipline_records.all(), key=lambda r: (r.date or date.min), reverse=True)
    clinic = sorted(student.clinic_visits.all(), key=lambda v: (v.date or date.min), reverse=True)
    # Edit/welfare controls follow write permission (the route already enforces
    # branch + form-class scope), so a form teacher can manage their own student.
    can_manage = page_can_write()
    sid = student.id
    return {
        'student': {
            'id': sid, 'student_id': student.student_id, 'full_name': student.full_name,
            'first_name': student.first_name, 'gender': student.gender,
            'date_of_birth': _fmt_date(student.date_of_birth), 'age': student.age,
            'religion': student.religion, 'home_address': student.home_address,
            'hobbies': student.hobbies, 'stream': student.stream,
            'house': student.house, 'boarding_status': student.boarding_status,
            'waec_subjects': student.waec_subject_list or [],
            'jamb_subjects': student.jamb_subject_list or [],
            'is_graduated': bool(student.is_graduated),
        },
        'identity': ({'nin': student.nin, 'jamb_reg_number': student.jamb_reg_number,
                      'jamb_profile_code': student.jamb_profile_code}
                     if student.has_identity else None),
        'medical': ({'blood_group': student.blood_group, 'genotype': student.genotype,
                     'allergies': student.allergies, 'medical_conditions': student.medical_conditions,
                     'disabilities': student.disabilities, 'medications': student.medications,
                     'medical_notes': student.medical_notes, 'emergency_medical': student.emergency_medical}
                    if student.has_medical else None),
        'contacts': [{'name': c.name, 'is_primary': bool(c.is_primary),
                      'phone_number': c.phone_number, 'relationship': c.relationship} for c in contacts],
        'enrollments': [{'term': e.class_arm_assignment.term.name,
                         'class': e.class_arm_assignment.school_class.name,
                         'arm': ('' if e.class_arm_assignment.arm.is_default
                                 else e.class_arm_assignment.arm.name)} for e in enrollments],
        'waec': {'count': len(waec_results),
                 'add_url': url_for('results.add_waec') + f'?student_id={sid}',
                 'view_url': url_for('results.view_waec_student', student_id=sid)},
        'jamb': {'latest': ({'year': latest_jamb.exam_year, 'score': latest_jamb.total_score}
                            if latest_jamb else None),
                 'add_url': url_for('results.add_jamb') + f'?student_id={sid}'},
        'discipline': [{'id': r.id, 'date': _fmt_date(r.date), 'category': r.category,
                        'severity': r.severity, 'description': r.description,
                        'action_taken': r.action_taken, 'reported_by': r.reported_by,
                        'delete_url': url_for('welfare.delete_discipline', record_id=r.id)} for r in discipline],
        'clinic': [{'id': v.id, 'date': _fmt_date(v.date), 'complaint': v.complaint,
                    'treatment': v.treatment, 'parent_notified': bool(v.parent_notified),
                    'attended_by': v.attended_by,
                    'delete_url': url_for('welfare.delete_clinic', visit_id=v.id)} for v in clinic],
        'attendance': _student_attendance_summary(sid),
        'communications': _student_comms_history(sid),
        'today': date.today().isoformat(),
        'can_manage': can_manage,
        'urls': {
            'list': url_for('main.students_list'),
            'self': url_for('main.view_student', student_id=sid),
            'edit': url_for('main.edit_student', student_id=sid),
            'graduate': url_for('promotion.unmark_graduate', student_id=sid) if student.is_graduated
                        else url_for('promotion.mark_graduate', student_id=sid),
            'exam_report': url_for('results.student_report', student_id=sid),
            'predictions': url_for('results.student_predictions', student_id=sid),
            'report_card': url_for('subjects.student_report_card', student_id=sid),
            'discipline_add': url_for('welfare.add_discipline', student_id=sid),
            'clinic_add': url_for('welfare.add_clinic', student_id=sid),
        },
        'discipline_categories': ['Lateness', 'Truancy', 'Uniform', 'Fighting', 'Rudeness',
                                  'Bullying', 'Property damage', 'Exam misconduct', 'Other'],
    }










# Optional pastoral / identity / medical fields, shared by add + edit so the two
# code paths can never drift. Plain-text fields are stripped; identity numbers
# stay searchable; medical free-text lands in encrypted columns.
_OPTIONAL_STUDENT_FIELDS = (
    'house', 'boarding_status', 'nin', 'jamb_reg_number', 'jamb_profile_code',
    'blood_group', 'genotype', 'allergies', 'medical_conditions', 'disabilities',
    'medications', 'medical_notes', 'emergency_medical',
)


def _apply_optional_student_fields(student, form, has=None):
    """Copy the optional identity/medical/pastoral fields from a form onto a
    student. ``has(key)`` decides whether a given field is present in this POST
    (so a partial edit never blanks fields it didn't submit); default: always."""
    from utils.security import strip_tags
    if has is None:
        has = lambda k: True   # noqa: E731
    for f in _OPTIONAL_STUDENT_FIELDS:
        if has(f):
            val = (form.get(f) or '').strip()
            setattr(student, f, strip_tags(val) or None)


def _manageable_student_ids(ids):
    """Restrict a list of student ids to those the current user may bulk-edit:
    branch-scoped, and a teacher only their own form-class students. Prevents a
    crafted id list from touching another branch's (or class's) students."""
    from utils.branch_scope import scope_query
    from utils.access_control import teacher_form_student_ids
    allowed = {s.id for s in scope_query(
        Student.query.filter(Student.id.in_(ids)), Student).all()}
    tids = teacher_form_student_ids()
    if tids is not None:
        allowed &= tids
    return [i for i in ids if i in allowed]










def _int_ids(raw_ids):
    """Coerce a list of form values to a clean list of ints (drops bad values)."""
    out = []
    for i in raw_ids:
        try:
            out.append(int(i))
        except (TypeError, ValueError):
            continue
    return out






def _trash_payload():
    """Soft-deleted students the current user may act on — branch- and
    teacher-scoped (a form teacher only sees their own students' trash)."""
    from utils.branch_scope import scope_query
    from utils.access_control import teacher_form_student_ids, page_can_write
    q = scope_query(Student.query.filter_by(is_active=False), Student)
    tids = teacher_form_student_ids()
    if tids is not None:
        q = q.filter(Student.id.in_(tids or [-1]))
    students = q.order_by(Student.surname, Student.first_name).all()
    return {
        'students': [{
            'id': s.id, 'full_name': s.full_name, 'student_id': s.student_id,
            'gender': s.gender, 'stream': s.stream,
            'restore_url': url_for('main.restore_student', student_id=s.id),
            'purge_url': url_for('main.purge_student', student_id=s.id),
        } for s in students],
        'can_manage': page_can_write(),
        'urls': {
            'list': url_for('main.students_list'),
            'bulk_restore': url_for('main.bulk_restore_students'),
            'bulk_purge': url_for('main.bulk_purge_students'),
        },
    }






def _bulk_no_selection():
    if _wants_json():
        return jsonify({'ok': False, 'error': 'No students selected.'}), 400
    flash('No students selected.', 'error')
    return redirect(url_for('main.students_trash'))


def _trash_scope(query):
    """Limit a soft-deleted-student query to the user's branch + form-teacher
    scope, so bulk restore/purge can never touch out-of-scope records."""
    from utils.branch_scope import scope_query
    from utils.access_control import teacher_form_student_ids
    query = scope_query(query, Student)
    tids = teacher_form_student_ids()
    if tids is not None:
        query = query.filter(Student.id.in_(tids or [-1]))
    return query
























# ============================================================================
# API ENDPOINTS FOR AJAX
# ============================================================================





# ============================================================================
# STUDENT EXPORT FUNCTIONALITY
# ============================================================================



def export_students_excel(student_data, fields):
    """Export students to Excel format with selected fields - wraps text and adjusts row heights"""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    from io import BytesIO
    from flask import Response
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Students"
    
    # Styles
    header_font = Font(bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    data_font = Font(size=10)
    border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='top', wrap_text=True)
    
    # Column width mapping based on field type
    col_width_map = {
        'Student ID': 14,
        'Surname': 15,
        'First Name': 15,
        'Middle Name': 15,
        'Gender': 10,
        'Class': 14,
        'Date of Birth': 14,
        'Age': 8,
        'Religion': 12,
        'Home Address': 35,
        'Hobbies': 30,
        'Parent Phone': 15,
    }
    
    # Title row
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(fields) + 1)
    title_cell = ws['A1']
    title_cell.value = 'STUDENTS LIST'
    title_cell.font = Font(bold=True, size=14)
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 25
    
    # Empty row
    ws.row_dimensions[2].height = 10
    
    # Header row (row 3)
    header_row = 3
    ws.cell(row=header_row, column=1, value='S/N')
    ws.cell(row=header_row, column=1).font = header_font
    ws.cell(row=header_row, column=1).fill = header_fill
    ws.cell(row=header_row, column=1).border = border
    ws.cell(row=header_row, column=1).alignment = center_align
    ws.column_dimensions['A'].width = 6
    
    for col, field in enumerate(fields, 2):
        cell = ws.cell(row=header_row, column=col, value=field)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = center_align
        
        # Set column width
        col_letter = get_column_letter(col)
        ws.column_dimensions[col_letter].width = col_width_map.get(field, 15)
    
    ws.row_dimensions[header_row].height = 20
    
    # Data rows
    for idx, student in enumerate(student_data, 1):
        row = idx + 3
        
        # S/N column
        sn_cell = ws.cell(row=row, column=1, value=idx)
        sn_cell.border = border
        sn_cell.alignment = Alignment(horizontal='center', vertical='top')
        sn_cell.font = data_font
        
        # Track max lines needed for this row
        max_lines = 1
        
        for col, field in enumerate(fields, 2):
            value = str(student.get(field, '') or '')
            cell = ws.cell(row=row, column=col, value=value)
            cell.border = border
            cell.alignment = left_align
            cell.font = data_font
            
            # Estimate lines needed based on column width and text length
            col_width = col_width_map.get(field, 15)
            chars_per_line = int(col_width * 1.2)  # Approximate chars that fit
            if chars_per_line > 0 and len(value) > 0:
                lines_needed = max(1, -(-len(value) // chars_per_line))  # Ceiling division
                # Also count actual newlines in the text
                lines_needed = max(lines_needed, value.count('\n') + 1)
                max_lines = max(max_lines, lines_needed)
        
        # Set row height based on content (15 points per line, minimum 20)
        row_height = max(20, max_lines * 15)
        ws.row_dimensions[row].height = row_height
    
    # Footer row
    footer_row = len(student_data) + 4
    ws.cell(row=footer_row, column=1, value=f'Total: {len(student_data)} students')
    ws.cell(row=footer_row, column=1).font = Font(bold=True, size=10)
    
    return xlsx_response(wb, 'students_export.xlsx')


def export_students_word(student_data, fields):
    """Export students to Word format with selected fields - wraps text and adjusts row heights"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from io import BytesIO
    from flask import Response
    
    doc = Document()
    
    # Set page to landscape for more space
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # Title
    title = doc.add_heading('STUDENTS LIST', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Column width mapping (in inches)
    col_width_map = {
        'S/N': 0.5,
        'Student ID': 1.0,
        'Surname': 1.2,
        'First Name': 1.2,
        'Middle Name': 1.2,
        'Gender': 0.7,
        'Class': 1.0,
        'Date of Birth': 1.0,
        'Age': 0.5,
        'Religion': 0.9,
        'Home Address': 2.0,
        'Hobbies': 1.8,
        'Parent Phone': 1.1,
    }
    
    # Create table with S/N + selected fields
    all_headers = ['S/N'] + fields
    table = doc.add_table(rows=1, cols=len(all_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for i, header in enumerate(all_headers):
        width = col_width_map.get(header, 1.0)
        table.columns[i].width = Inches(width)
    
    # Style header row
    header_row = table.rows[0]
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    header_row.height = Pt(25)
    
    for i, header in enumerate(all_headers):
        cell = header_row.cells[i]
        cell.text = header
        
        # Bold and center header text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        
        # Set header background color (blue)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = None  # Reset to let white show
        # Set white text
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for idx, student in enumerate(student_data, 1):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        
        # Calculate row height based on content
        max_lines = 1
        row_data = [str(idx)]
        
        for field in fields:
            value = str(student.get(field, '') or '')
            row_data.append(value)
            
            # Estimate lines needed
            col_width = col_width_map.get(field, 1.0)
            chars_per_line = int(col_width * 12)  # Approximate
            if chars_per_line > 0 and len(value) > 0:
                words = value.split()
                current_line_len = 0
                lines = 1
                for word in words:
                    if current_line_len + len(word) + 1 > chars_per_line:
                        lines += 1
                        current_line_len = len(word)
                    else:
                        current_line_len += len(word) + 1
                max_lines = max(max_lines, lines)
        
        # Set minimum row height based on content
        row.height = Pt(max(20, max_lines * 14))
        
        # Fill cells
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = value
            
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(9)
            
            # Alternate row coloring
            if idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'Total: {len(student_data)} students')
    footer.runs[0].bold = True
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return Response(
        output.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename=students_export.docx'}
    )


def export_students_pdf(student_data, fields):
    """Export students to PDF format with selected fields - wraps text and adjusts row heights"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from io import BytesIO
    from flask import Response
    
    output = BytesIO()
    
    # Use landscape for more space
    page_size = landscape(A4)
    doc = SimpleDocTemplate(output, pagesize=page_size, leftMargin=10*mm, rightMargin=10*mm,
                           topMargin=10*mm, bottomMargin=10*mm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Create custom styles for cells
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontSize=8,
        leading=10,  # Line height
        alignment=0,  # Left align
    )
    
    header_style = ParagraphStyle(
        'HeaderStyle',
        parent=styles['Normal'],
        fontSize=9,
        leading=11,
        alignment=1,  # Center align
        textColor=colors.white,
        fontName='Helvetica-Bold',
    )
    
    # Title
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, alignment=1, spaceAfter=10)
    elements.append(Paragraph('STUDENTS LIST', title_style))
    elements.append(Spacer(1, 5*mm))
    
    # Column width mapping (in mm)
    col_width_map = {
        'S/N': 10*mm,
        'Student ID': 22*mm,
        'Surname': 28*mm,
        'First Name': 28*mm,
        'Middle Name': 25*mm,
        'Gender': 15*mm,
        'Class': 25*mm,
        'Date of Birth': 22*mm,
        'Age': 12*mm,
        'Religion': 20*mm,
        'Home Address': 45*mm,
        'Hobbies': 40*mm,
        'Parent Phone': 25*mm,
    }
    
    # Build column widths list
    all_headers = ['S/N'] + fields
    col_widths = [col_width_map.get(h, 25*mm) for h in all_headers]
    
    # Adjust widths to fit page if needed
    total_width = sum(col_widths)
    available_width = page_size[0] - 20*mm
    if total_width > available_width:
        scale = available_width / total_width
        col_widths = [w * scale for w in col_widths]
    
    # Build table data with Paragraphs for text wrapping
    table_data = []
    
    # Header row
    header_row = [Paragraph(h, header_style) for h in all_headers]
    table_data.append(header_row)
    
    # Data rows
    for idx, student in enumerate(student_data, 1):
        row = [Paragraph(str(idx), cell_style)]
        for field in fields:
            value = str(student.get(field, '') or '')
            row.append(Paragraph(value, cell_style))
        table_data.append(row)
    
    # Create table - let ReportLab calculate row heights automatically
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    
    # Style the table
    style_commands = [
        # Header styling
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        
        # Data styling
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        
        # Alignment
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Header centered
        ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # S/N column centered
        ('ALIGN', (1, 1), (-1, -1), 'LEFT'),   # Data left aligned
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),   # Top align all cells
        
        # Padding
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        
        # Grid
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#4472C4')),
        
        # Header bottom border
        ('LINEBELOW', (0, 0), (-1, 0), 1.5, colors.HexColor('#4472C4')),
    ]
    
    # Add alternating row colors
    for i in range(1, len(table_data)):
        if i % 2 == 0:
            style_commands.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor('#F5F5F5')))
    
    table.setStyle(TableStyle(style_commands))
    elements.append(table)
    
    # Footer
    elements.append(Spacer(1, 5*mm))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=10, fontName='Helvetica-Bold')
    elements.append(Paragraph(f'Total: {len(student_data)} students', footer_style))
    
    doc.build(elements)
    return pdf_response(output, 'students_export.pdf', inline=False)


def export_students_image(student_data, fields):
    """Render the students list to a modern, high-resolution PNG."""
    from PIL import Image, ImageDraw, ImageFont
    from io import BytesIO
    from flask import Response
    from datetime import datetime

    S = 3  # supersampling for a crisp, high-DPI image

    def fnt(size, bold=False):
        path = "/usr/share/fonts/truetype/dejavu/DejaVuSans%s.ttf" % ("-Bold" if bold else "")
        try:
            return ImageFont.truetype(path, size * S)
        except Exception:
            return ImageFont.load_default()

    body = fnt(12)
    body_b = fnt(12, True)
    head_f = fnt(12, True)
    title_f = fnt(22, True)
    sub_f = fnt(11)
    foot_f = fnt(10)

    # Brand colours
    GREEN = (13, 106, 78)
    GREEN_DK = (6, 78, 54)
    HEAD = (10, 86, 64)
    TEXT = (31, 41, 55)
    MUTED = (107, 114, 128)
    ZEBRA = (249, 250, 251)
    LINE = (229, 231, 235)
    WHITE = (255, 255, 255)

    try:
        from models import SchoolSettings
        school_name = SchoolSettings.get('school_name', 'PosyHub') or 'PosyHub'
    except Exception:
        school_name = 'PosyHub'

    tmp = ImageDraw.Draw(Image.new('RGB', (1, 1)))

    if not student_data:
        img = Image.new('RGB', (560 * S, 160 * S), WHITE)
        d = ImageDraw.Draw(img)
        d.text((40 * S, 60 * S), "No students to export", fill=TEXT, font=title_f)
        out = BytesIO(); img.save(out, format='PNG'); out.seek(0)
        return Response(out.getvalue(), mimetype='image/png',
                        headers={'Content-Disposition': 'attachment; filename=students_export.png'})

    pad = 26 * S
    cpx = 12 * S
    cpy = 9 * S
    header_h = 42 * S
    min_row_h = 38 * S
    line_bbox = tmp.textbbox((0, 0), "Ay", font=body)
    line_h = (line_bbox[3] - line_bbox[1]) + 6 * S

    all_headers = ['S/N'] + fields
    width_map = {'S/N': 56, 'Student ID': 130, 'Surname': 150, 'First Name': 150,
                 'Middle Name': 150, 'Gender': 86, 'Age': 70, 'Class': 130,
                 'Religion': 120, 'Date of Birth': 130, 'Stream': 120,
                 'Home Address': 230, 'Hobbies': 210, 'Parent Phone': 140}
    col_widths = [width_map.get(h, 150) * S for h in all_headers]

    # Wrap + measure rows
    wrapped_data, row_heights = [], []
    for idx, student in enumerate(student_data):
        row, max_lines = [], 1
        for i, header in enumerate(all_headers):
            value = str(idx + 1) if i == 0 else str(student.get(fields[i - 1], '') or '')
            lines = wrap_text(tmp, value, col_widths[i] - cpx * 2, body)
            row.append(lines)
            max_lines = max(max_lines, len(lines))
        wrapped_data.append(row)
        row_heights.append(max(min_row_h, max_lines * line_h + cpy * 2))

    table_w = sum(col_widths)
    band_h = 88 * S
    total_w = table_w + pad * 2
    total_h = band_h + header_h + sum(row_heights) + 56 * S

    img = Image.new('RGB', (int(total_w), int(total_h)), WHITE)
    d = ImageDraw.Draw(img)

    # Header band with a vertical gradient
    for yy in range(band_h):
        t = yy / band_h
        col = tuple(int(GREEN[k] + (GREEN_DK[k] - GREEN[k]) * t) for k in range(3))
        d.line([(0, yy), (total_w, yy)], fill=col)
    d.text((pad, 20 * S), school_name, fill=WHITE, font=title_f)
    d.text((pad, 56 * S),
           "Students List  ·  %d student%s  ·  %s" % (
               len(student_data), '' if len(student_data) == 1 else 's',
               datetime.now().strftime('%d %b %Y')),
           fill=(220, 240, 232), font=sub_f)

    # Column header row
    y = band_h
    d.rectangle([0, y, total_w, y + header_h], fill=HEAD)
    x = pad
    for i, header in enumerate(all_headers):
        bb = d.textbbox((0, 0), header, font=head_f)
        d.text((x + cpx, y + (header_h - (bb[3] - bb[1])) // 2 - bb[1]), header, fill=WHITE, font=head_f)
        x += col_widths[i]
    y += header_h

    # Data rows (zebra + subtle separators)
    for idx, (row, rh) in enumerate(zip(wrapped_data, row_heights)):
        if idx % 2 == 1:
            d.rectangle([pad, y, pad + table_w, y + rh], fill=ZEBRA)
        x = pad
        for i, lines in enumerate(row):
            ty = y + cpy
            f = body_b if i == 0 else body
            fill = GREEN if i == 0 else TEXT
            for ln in lines:
                d.text((x + cpx, ty), ln, fill=fill, font=f)
                ty += line_h
            x += col_widths[i]
        d.line([(pad, y + rh), (pad + table_w, y + rh)], fill=LINE, width=max(1, S // 2))
        y += rh

    # Footer
    d.text((pad, y + 16 * S),
           "Generated by %s · %s" % (school_name, datetime.now().strftime('%d %b %Y %H:%M')),
           fill=MUTED, font=foot_f)

    out = BytesIO()
    img.save(out, format='PNG')
    out.seek(0)
    return Response(out.getvalue(), mimetype='image/png',
                    headers={'Content-Disposition': 'attachment; filename=students_export.png'})


def wrap_text(draw, text, max_width, font):
    """Wrap text to fit within max_width, returns list of lines"""
    if not text:
        return ['']
    
    text = str(text).strip()
    if not text:
        return ['']
    
    # Check if text already fits on one line
    bbox = draw.textbbox((0, 0), text, font=font)
    if bbox[2] - bbox[0] <= max_width:
        return [text]
    
    # Split into words
    words = text.split()
    if not words:
        return ['']
    
    lines = []
    current_line = []
    
    for word in words:
        # Check if adding this word exceeds the width
        test_line = ' '.join(current_line + [word])
        bbox = draw.textbbox((0, 0), test_line, font=font)
        
        if bbox[2] - bbox[0] <= max_width:
            current_line.append(word)
        else:
            # If current line has words, save it
            if current_line:
                lines.append(' '.join(current_line))
                current_line = [word]
            else:
                # Word is too long, need to break it
                lines.append(break_long_word(draw, word, max_width, font))
                current_line = []
    
    # Don't forget the last line
    if current_line:
        lines.append(' '.join(current_line))
    
    return lines if lines else ['']


def break_long_word(draw, word, max_width, font):
    """Break a long word that doesn't fit, adding ellipsis if needed"""
    for i in range(len(word), 0, -1):
        truncated = word[:i]
        bbox = draw.textbbox((0, 0), truncated, font=font)
        if bbox[2] - bbox[0] <= max_width:
            if i < len(word):
                # Try to fit with ellipsis
                with_ellipsis = word[:max(1, i-2)] + '..'
                bbox2 = draw.textbbox((0, 0), with_ellipsis, font=font)
                if bbox2[2] - bbox2[0] <= max_width:
                    return with_ellipsis
            return truncated
    return word[:1] if word else ''

# Auto-export every module-level name (incl. the underscore helpers) so the
# route submodules below can `from routes.main import *` and see them.
__all__ = [_n for _n in dir() if not _n.startswith('__')]

# Importing the submodules registers their routes on main_bp.
from . import branches, dashboard, students, misc  # noqa: E402,F401
