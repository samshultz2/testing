"""
CBT / Online Tests.

Two surfaces:
  * Admin/teacher (``cbt_bp`` at /cbt): build exams + questions, set the per-exam
    access password, publish, view results, and manage student portal passwords.
  * Student portal (``cbt_portal_bp`` at /exam): a separate login (student ID +
    portal password); shows exams active for the student's class *today*; each
    exam is gated by its own access password; answers are auto-graded.
"""
from datetime import datetime, timedelta
from utils.helpers import get_active_term
from functools import wraps
import io
import os
import random
import secrets

from flask import (Blueprint, render_template, request, redirect, url_for,
                   flash, session, jsonify, Response, current_app)
from werkzeug.utils import secure_filename
from utils.web_exports import xlsx_response, pdf_response
from sqlalchemy import func
from sqlalchemy.orm import contains_eager

from models import (db, CBTExam, CBTQuestion, CBTAttempt, CBTAnswer, CBTViolation,
                    QuestionBank, CBTLoginEvent, CBTDeviceSession, Subject, SchoolClass,
                    ClassArm, Term, Student, StudentEnrollment, ClassArmAssignment,
                    SchoolSettings)
from utils.access_control import login_required, admin_required, is_admin
from utils import timeutil
from utils.useragent import parse_user_agent

cbt_bp = Blueprint('cbt', __name__, url_prefix='/cbt')
cbt_portal_bp = Blueprint('cbt_portal', __name__, url_prefix='/exam')

PORTAL_KEY = 'cbt_student_id'
MAX_VIOLATIONS = 3   # leave-page events allowed before the test auto-submits
# After the timer expires, how long a returning student's device still gets the
# chance to sync its locally-saved answers and submit them, before the server
# force-grades the attempt from what it already has (network-outage recovery).
OFFLINE_GRACE_SECONDS = 30 * 60


def _d(value, default=None):
    try:
        return datetime.strptime(value, '%Y-%m-%d').date()
    except (ValueError, TypeError):
        return default


_IMG_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg'}


def _save_question_image(file):
    """Save an uploaded question figure under static/uploads/cbt; return its URL."""
    if not file or not file.filename:
        return None
    ext = os.path.splitext(secure_filename(file.filename))[1].lower()
    if ext not in _IMG_EXTS:
        return None
    name = secrets.token_hex(8) + ext
    folder = os.path.join(current_app.root_path, 'static', 'uploads', 'cbt')
    os.makedirs(folder, exist_ok=True)
    file.save(os.path.join(folder, name))
    return url_for('static', filename='uploads/cbt/' + name)


def _active_term():
    return get_active_term()


def _student_placement(student_id, term):
    if not term:
        return None, None
    enr = (StudentEnrollment.query
           .join(ClassArmAssignment,
                 StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id)
           .filter(StudentEnrollment.student_id == student_id,
                   StudentEnrollment.is_active == True,
                   ClassArmAssignment.term_id == term.id).first())
    if not enr:
        return None, None
    asg = enr.class_arm_assignment
    return asg.class_id, asg.arm_id


# ============================================================================
# ADMIN — DASHBOARD + EXAM MANAGEMENT
# ============================================================================

@cbt_bp.route('/')
@login_required
def dashboard():
    terms = Term.query.order_by(Term.id.desc()).all()
    active = _active_term()
    # Term-scoped (default active term); 'all' shows history across terms.
    raw = request.args.get('term_id')
    if raw == 'all':
        term_id = None
    else:
        term_id = request.args.get('term_id', type=int) or (active.id if active else None)
    q = CBTExam.query
    from utils.branch_scope import scope_query
    q = scope_query(q, CBTExam)
    if term_id:
        q = q.filter(CBTExam.term_id == term_id)
    exams = q.order_by(CBTExam.exam_date.desc(), CBTExam.id.desc()).all()
    total = len(exams)
    published = sum(1 for e in exams if e.is_published)
    today_count = sum(1 for e in exams if e.exam_date == timeutil.today() and e.is_published)
    exam_ids = [e.id for e in exams]
    attempts = (CBTAttempt.query.filter(CBTAttempt.status == 'Submitted',
                CBTAttempt.exam_id.in_(exam_ids)).count() if exam_ids else 0)
    return render_template('cbt/dashboard.html', exams=exams, total=total,
        published=published, today_count=today_count, attempts=attempts,
        terms=terms, term_id=term_id, show_all=(raw == 'all'))


@cbt_bp.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if request.method == 'POST':
        if not is_admin():
            flash('Admins only.', 'error')
            return redirect(url_for('cbt.settings'))
        SchoolSettings.set('cbt_supervisor_pin', (request.form.get('supervisor_pin') or '').strip(),
                           'string', 'CBT supervisor override PIN')
        flash('CBT settings saved.', 'success')
        return redirect(url_for('cbt.settings'))
    return render_template('cbt/settings.html',
        supervisor_pin=SchoolSettings.get('cbt_supervisor_pin', ''))


@cbt_bp.route('/lab-setup')
@login_required
def lab_setup():
    """Invigilator guide: launch lab laptops in locked kiosk mode."""
    portal_url = url_for('cbt_portal.login', _external=True)
    return render_template('cbt/lab_setup.html', portal_url=portal_url)


def _exam_choices():
    return {
        'subjects': Subject.query.filter_by(is_active=True).order_by(Subject.name).all(),
        'classes': SchoolClass.query.filter_by(is_active=True).order_by(SchoolClass.level).all(),
        'arms': ClassArm.query.filter_by(is_active=True).order_by(ClassArm.name).all(),
        'terms': Term.query.order_by(Term.id.desc()).all(),
    }


def _read_exam(e):
    e.title = (request.form.get('title') or '').strip()
    e.subject_id = request.form.get('subject_id', type=int) or None
    e.class_id = request.form.get('class_id', type=int) or None
    e.arm_id = request.form.get('arm_id', type=int) or None
    e.term_id = request.form.get('term_id', type=int) or None
    e.exam_date = _d(request.form.get('exam_date')) or timeutil.today()
    e.start_time = (request.form.get('start_time') or '').strip() or None
    e.end_time = (request.form.get('end_time') or '').strip() or None
    e.duration_minutes = request.form.get('duration_minutes', type=int) or 30
    e.max_score = request.form.get('max_score', type=float)
    e.access_password = (request.form.get('access_password') or '').strip()
    e.instructions = (request.form.get('instructions') or '').strip() or None
    e.shuffle = bool(request.form.get('shuffle'))
    e.strict_mode = bool(request.form.get('strict_mode'))
    vl = request.form.get('violation_limit', type=int)
    e.violation_limit = vl if vl is not None and vl >= 0 else 3


@cbt_bp.route('/exams/add', methods=['GET', 'POST'])
@login_required
def add_exam():
    if request.method == 'POST':
        if not (request.form.get('title') and request.form.get('access_password')):
            flash('Title and an access password are required.', 'error')
            return redirect(url_for('cbt.add_exam'))
        from flask import session as _s
        e = CBTExam(created_by=_s.get('username') or 'Admin')
        _read_exam(e)
        from utils.branch_scope import branch_for_new
        e.branch_id = branch_for_new()
        db.session.add(e)
        db.session.commit()
        from utils.audit import log_action
        log_action('cbt.exam_create', target=e)
        flash('Exam created — now add questions.', 'success')
        return redirect(url_for('cbt.exam_detail', exam_id=e.id))
    active = _active_term()
    return render_template('cbt/exam_form.html', exam=None, active_term=active, **_exam_choices())


@cbt_bp.route('/exams/<int:exam_id>')
@login_required
def exam_detail(exam_id):
    e = db.get_or_404(CBTExam, exam_id)
    questions = e.questions.order_by(CBTQuestion.order, CBTQuestion.id).all()
    return render_template('cbt/exam_detail.html', e=e, questions=questions)


@cbt_bp.route('/exams/<int:exam_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_exam(exam_id):
    e = db.get_or_404(CBTExam, exam_id)
    if request.method == 'POST':
        _read_exam(e)
        db.session.commit()
        flash('Exam updated.', 'success')
        return redirect(url_for('cbt.exam_detail', exam_id=e.id))
    return render_template('cbt/exam_form.html', exam=e, active_term=_active_term(), **_exam_choices())


@cbt_bp.route('/exams/<int:exam_id>/publish', methods=['POST'])
@login_required
def toggle_publish(exam_id):
    e = db.get_or_404(CBTExam, exam_id)
    if not e.is_published and e.question_count == 0:
        flash('Add at least one question before publishing.', 'error')
        return redirect(url_for('cbt.exam_detail', exam_id=e.id))
    e.is_published = not e.is_published
    db.session.commit()
    from utils.audit import log_action
    log_action('cbt.exam_publish' if e.is_published else 'cbt.exam_unpublish', target=e)
    flash('Exam published.' if e.is_published else 'Exam unpublished.', 'success')
    return redirect(url_for('cbt.exam_detail', exam_id=e.id))


@cbt_bp.route('/exams/<int:exam_id>/delete', methods=['POST'])
@admin_required
def delete_exam(exam_id):
    e = db.get_or_404(CBTExam, exam_id)
    db.session.delete(e)
    db.session.commit()
    flash('Exam deleted.', 'success')
    return redirect(url_for('cbt.dashboard'))


@cbt_bp.route('/exams/<int:exam_id>/questions/add', methods=['POST'])
@login_required
def add_question(exam_id):
    e = db.get_or_404(CBTExam, exam_id)
    text = (request.form.get('question_text') or '').strip()
    correct = (request.form.get('correct_option') or '').strip().upper()
    if not text or correct not in ('A', 'B', 'C', 'D'):
        flash('Question text and a correct option (A–D) are required.', 'error')
        return redirect(url_for('cbt.exam_detail', exam_id=exam_id))
    nextord = (db.session.query(func.coalesce(func.max(CBTQuestion.order), 0))
               .filter(CBTQuestion.exam_id == exam_id).scalar()) + 1
    image_url = _save_question_image(request.files.get('image'))
    db.session.add(CBTQuestion(
        exam_id=exam_id, question_text=text, image_url=image_url,
        option_a=(request.form.get('option_a') or '').strip(),
        option_b=(request.form.get('option_b') or '').strip(),
        option_c=(request.form.get('option_c') or '').strip(),
        option_d=(request.form.get('option_d') or '').strip(),
        correct_option=correct, marks=request.form.get('marks', type=float) or 1,
        order=nextord))
    db.session.commit()
    flash('Question added.', 'success')
    return redirect(url_for('cbt.exam_detail', exam_id=exam_id))


@cbt_bp.route('/exams/<int:exam_id>/import-file', methods=['POST'])
@login_required
def import_questions_file(exam_id):
    """Bulk-add questions to an exam from an uploaded Excel/CSV file."""
    from utils import cbt_import
    e = db.get_or_404(CBTExam, exam_id)
    f = request.files.get('file')
    if not f or not f.filename:
        flash('Choose an Excel or CSV file.', 'error')
        return redirect(url_for('cbt.exam_detail', exam_id=exam_id))
    try:
        parsed = cbt_import.parse_file(f.read(), f.filename)
    except Exception as ex:
        flash(f'Could not read the file: {ex}', 'error')
        return redirect(url_for('cbt.exam_detail', exam_id=exam_id))
    if not parsed:
        flash('No valid questions found. Check the columns (Question, A, B, C, D, Correct, Marks).', 'warning')
        return redirect(url_for('cbt.exam_detail', exam_id=exam_id))
    nextord = (db.session.query(func.coalesce(func.max(CBTQuestion.order), 0))
               .filter(CBTQuestion.exam_id == exam_id).scalar())
    for q in parsed:
        nextord += 1
        db.session.add(CBTQuestion(exam_id=exam_id, order=nextord, **{k: q[k] for k in
            ('question_text', 'option_a', 'option_b', 'option_c', 'option_d', 'correct_option', 'marks')}))
    db.session.commit()
    flash(f'Imported {len(parsed)} question(s) into the exam.', 'success')
    return redirect(url_for('cbt.exam_detail', exam_id=exam_id))


@cbt_bp.route('/exams/<int:exam_id>/import-bank', methods=['GET', 'POST'])
@login_required
def import_from_bank(exam_id):
    """Pick questions from the bank to copy into an exam."""
    e = db.get_or_404(CBTExam, exam_id)
    if request.method == 'POST':
        ids = request.form.getlist('bank_id', type=int)
        nextord = (db.session.query(func.coalesce(func.max(CBTQuestion.order), 0))
                   .filter(CBTQuestion.exam_id == exam_id).scalar())
        n = 0
        for bq in QuestionBank.query.filter(QuestionBank.id.in_(ids or [-1])).all():
            nextord += 1
            db.session.add(CBTQuestion(exam_id=exam_id, order=nextord,
                question_text=bq.question_text, image_url=bq.image_url,
                option_a=bq.option_a, option_b=bq.option_b,
                option_c=bq.option_c, option_d=bq.option_d,
                correct_option=bq.correct_option, marks=bq.marks))
            n += 1
        db.session.commit()
        flash(f'Added {n} question(s) from the bank.', 'success')
        return redirect(url_for('cbt.exam_detail', exam_id=exam_id))
    subject_id = request.args.get('subject_id', type=int) or e.subject_id
    q = QuestionBank.query.filter_by(is_active=True)
    if subject_id:
        q = q.filter_by(subject_id=subject_id)
    topic = (request.args.get('topic') or '').strip()
    if topic:
        q = q.filter(QuestionBank.topic.ilike(f'%{topic}%'))
    questions = q.order_by(QuestionBank.topic, QuestionBank.id.desc()).all()
    subjects = Subject.query.filter_by(is_active=True).order_by(Subject.name).all()
    return render_template('cbt/import_bank.html', e=e, questions=questions,
                           subjects=subjects, subject_id=subject_id, topic=topic)


@cbt_bp.route('/questions/<int:question_id>/delete', methods=['POST'])
@login_required
def delete_question(question_id):
    q = db.get_or_404(CBTQuestion, question_id)
    exam_id = q.exam_id
    db.session.delete(q)
    db.session.commit()
    flash('Question removed.', 'success')
    return redirect(url_for('cbt.exam_detail', exam_id=exam_id))


# ============================================================================
# QUESTION BANK
# ============================================================================

@cbt_bp.route('/bank')
@login_required
def bank():
    subject_id = request.args.get('subject_id', type=int)
    topic = (request.args.get('topic') or '').strip()
    search = (request.args.get('q') or '').strip()
    q = QuestionBank.query.filter_by(is_active=True)
    if subject_id:
        q = q.filter_by(subject_id=subject_id)
    if topic:
        q = q.filter(QuestionBank.topic.ilike(f'%{topic}%'))
    if search:
        q = q.filter(QuestionBank.question_text.ilike(f'%{search}%'))
    questions = q.order_by(QuestionBank.id.desc()).limit(500).all()
    subjects = Subject.query.filter_by(is_active=True).order_by(Subject.name).all()
    total = QuestionBank.query.filter_by(is_active=True).count()
    return render_template('cbt/bank.html', questions=questions, subjects=subjects,
        subject_id=subject_id, topic=topic, search=search, total=total)


@cbt_bp.route('/bank/add', methods=['POST'])
@login_required
def bank_add():
    text = (request.form.get('question_text') or '').strip()
    correct = (request.form.get('correct_option') or '').strip().upper()
    if not text or correct not in ('A', 'B', 'C', 'D'):
        flash('Question text and a correct option (A–D) are required.', 'error')
        return redirect(url_for('cbt.bank'))
    from flask import session as _s
    db.session.add(QuestionBank(
        question_text=text, correct_option=correct,
        image_url=_save_question_image(request.files.get('image')),
        subject_id=request.form.get('subject_id', type=int) or None,
        topic=(request.form.get('topic') or '').strip() or None,
        difficulty=request.form.get('difficulty') or 'Medium',
        option_a=(request.form.get('option_a') or '').strip(),
        option_b=(request.form.get('option_b') or '').strip(),
        option_c=(request.form.get('option_c') or '').strip(),
        option_d=(request.form.get('option_d') or '').strip(),
        marks=request.form.get('marks', type=float) or 1,
        created_by=_s.get('username') or 'Admin'))
    db.session.commit()
    flash('Question added to the bank.', 'success')
    return redirect(url_for('cbt.bank', subject_id=request.form.get('subject_id') or ''))


@cbt_bp.route('/bank/<int:bank_id>/delete', methods=['POST'])
@login_required
def bank_delete(bank_id):
    bq = db.get_or_404(QuestionBank, bank_id)
    bq.is_active = False
    db.session.commit()
    flash('Question removed from the bank.', 'success')
    return redirect(request.referrer or url_for('cbt.bank'))


@cbt_bp.route('/bank/import', methods=['GET', 'POST'])
@login_required
def bank_import():
    subjects = Subject.query.filter_by(is_active=True).order_by(Subject.name).all()
    if request.method == 'POST':
        from utils import cbt_import
        from flask import session as _s
        f = request.files.get('file')
        subject_id = request.form.get('subject_id', type=int) or None
        if not f or not f.filename:
            flash('Choose an Excel or CSV file.', 'error')
            return redirect(url_for('cbt.bank_import'))
        try:
            parsed = cbt_import.parse_file(f.read(), f.filename)
        except Exception as ex:
            flash(f'Could not read the file: {ex}', 'error')
            return redirect(url_for('cbt.bank_import'))
        if not parsed:
            flash('No valid questions found. Check the column headers.', 'warning')
            return redirect(url_for('cbt.bank_import'))
        for q in parsed:
            db.session.add(QuestionBank(subject_id=subject_id, created_by=_s.get('username') or 'Admin',
                question_text=q['question_text'], option_a=q['option_a'], option_b=q['option_b'],
                option_c=q['option_c'], option_d=q['option_d'], correct_option=q['correct_option'],
                marks=q['marks'], topic=q.get('topic'), difficulty=q.get('difficulty') or 'Medium'))
        db.session.commit()
        flash(f'Imported {len(parsed)} question(s) into the bank.', 'success')
        return redirect(url_for('cbt.bank', subject_id=subject_id or ''))
    return render_template('cbt/bank_import.html', subjects=subjects)


@cbt_bp.route('/bank/template')
@login_required
def bank_template():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    wb = Workbook()
    ws = wb.active
    ws.title = 'Questions'
    headers = ['Question', 'A', 'B', 'C', 'D', 'Correct', 'Marks', 'Topic', 'Difficulty']
    fill = PatternFill(start_color='0d6a4e', end_color='0d6a4e', fill_type='solid')
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.fill = fill
        c.font = Font(bold=True, color='FFFFFF')
    ws.append(['What is 5 + 7?', '10', '12', '13', '11', 'B', 1, 'Arithmetic', 'Easy'])
    ws.append(['Water is made of hydrogen and ___?', 'Oxygen', 'Nitrogen', 'Carbon', 'Helium', 'A', 1, 'Science', 'Easy'])
    return xlsx_response(wb, 'cbt_question_template.xlsx')


@cbt_bp.route('/exams/<int:exam_id>/results')
@login_required
def results(exam_id):
    e = db.get_or_404(CBTExam, exam_id)
    attempts = (e.attempts.join(Student).order_by(Student.surname, Student.first_name).all())
    submitted = [a for a in attempts if a.status == 'Submitted']
    avg = round(sum(a.score for a in submitted) / len(submitted), 1) if submitted else 0

    # Per-question analytics: % of submitted attempts that got each question right.
    questions = e.questions.order_by(CBTQuestion.order, CBTQuestion.id).all()
    sub_ids = [a.id for a in submitted]
    analysis = []
    for q in questions:
        correct = 0
        if sub_ids:
            correct = CBTAnswer.query.filter(
                CBTAnswer.question_id == q.id,
                CBTAnswer.attempt_id.in_(sub_ids),
                CBTAnswer.is_correct == True).count()
        pct = round(correct / len(submitted) * 100, 1) if submitted else 0
        analysis.append({'q': q, 'correct': correct, 'pct': pct})
    return render_template('cbt/results.html', e=e, attempts=attempts, avg=avg,
                           analysis=analysis, submitted_count=len(submitted))


@cbt_bp.route('/exams/<int:exam_id>/monitor')
@login_required
def monitor(exam_id):
    e = db.get_or_404(CBTExam, exam_id)
    # Central users get a branch filter (CBT is school-wide); branch users are
    # already scoped to their own branch.
    from utils.branch_scope import is_central
    from models import Branch
    branches = (Branch.query.filter_by(is_active=True).order_by(Branch.name).all()
                if is_central() else [])
    return render_template('cbt/monitor.html', e=e, branches=branches)


@cbt_bp.route('/attempts/<int:attempt_id>/force-submit', methods=['POST'])
@admin_required
def force_submit(attempt_id):
    """Invigilator action: end an in-progress attempt and grade it from the
    answers already saved on the server. The recovery tool for a student whose
    device died / went offline and never came back to submit."""
    attempt = db.get_or_404(CBTAttempt, attempt_id)
    if attempt.status == 'Submitted':
        return jsonify({'ok': True, 'already': True})
    exam = db.session.get(CBTExam, attempt.exam_id)
    _finalize(attempt, exam)
    from utils.audit import log_action
    log_action('cbt.force_submit',
               f'attempt {attempt.id} ({attempt.student.full_name if attempt.student else "?"}) '
               f'exam {exam.id}')
    return jsonify({'ok': True, 'score': attempt.score, 'total': attempt.total})


@cbt_bp.route('/exams/<int:exam_id>/monitor/data')
@login_required
def monitor_data(exam_id):
    """Live JSON snapshot of every attempt for the invigilator monitor.

    CBT exams are school-wide (set centrally for every branch), but the device
    fingerprints/attempts can be filtered by branch to make tracking easier.
    Branch users are always restricted to their own branch.
    """
    e = db.get_or_404(CBTExam, exam_id)
    now = timeutil.now()
    qcount = e.question_count
    # Resolve the branch filter: branch users -> their branch; a central user may
    # pick one via the ?branch= param (else sees all branches).
    from utils.branch_scope import viewing_branch_id
    _vbid = viewing_branch_id()
    if _vbid is not None:
        branch_filter = _vbid
    else:
        branch_filter = request.args.get('branch', type=int)
    # answered counts per attempt (only saved answers with a choice)
    answered = dict(db.session.query(CBTAnswer.attempt_id, func.count(CBTAnswer.id))
                    .join(CBTAttempt, CBTAnswer.attempt_id == CBTAttempt.id)
                    .filter(CBTAttempt.exam_id == exam_id,
                            CBTAnswer.selected_option != None)
                    .group_by(CBTAnswer.attempt_id).all())
    # Latest device fingerprint per student for this exam (login or start event).
    # Bound the scan: this exam's events, plus recent global-login events only,
    # newest first, capped — so the monitor never loads the full history table.
    fp_by_student = {}
    _recent_cut = now - timedelta(days=1)
    fps = (CBTLoginEvent.query
           .filter(db.or_(CBTLoginEvent.exam_id == exam_id,
                          db.and_(CBTLoginEvent.exam_id == None,
                                  CBTLoginEvent.created_at >= _recent_cut)))
           .order_by(CBTLoginEvent.created_at.desc())
           .limit(6000).all())
    for fp in fps:
        fp_by_student.setdefault(fp.student_id, fp)
    # Live devices per student (pinged recently) -> flag concurrent sessions.
    active_cut = now - timedelta(seconds=DEVICE_ACTIVE_WINDOW)
    active_by_student = {}
    for ds in (CBTDeviceSession.query
               .filter(CBTDeviceSession.exam_id == exam_id,
                       CBTDeviceSession.last_seen >= active_cut).all()):
        active_by_student.setdefault(ds.student_id, []).append(ds)
    rows = []
    # join Student for filter/order AND map it onto each attempt (no per-row load).
    _attempts_q = e.attempts.join(Student).options(contains_eager(CBTAttempt.student))
    if branch_filter:
        _attempts_q = _attempts_q.filter(Student.branch_id == branch_filter)
    for a in _attempts_q.order_by(Student.surname, Student.first_name).all():
        deadline = _deadline(a, e)
        remaining = int((deadline - now).total_seconds())
        seen_ago = int((now - a.last_seen).total_seconds()) if a.last_seen else None
        paused = bool(a.paused_until and now < a.paused_until)
        active_devs = active_by_student.get(a.student_id, [])
        concurrent = a.status == 'In progress' and len(active_devs) >= 2
        fp = fp_by_student.get(a.student_id)
        device = None
        if fp:
            bits = [fp.device_model or fp.device_type, fp.browser, fp.os]
            device = {
                'model': fp.device_model,
                'type': fp.device_type,
                'is_mobile': bool(fp.is_mobile),
                'browser': fp.browser,
                'os': fp.os,
                'ip': fp.ip_address or a.ip_address,
                'location': fp.location,
                'label': ' · '.join(b for b in bits if b) or '—',
            }
        rows.append({
            'id': a.id,
            'name': a.student.full_name if a.student else '—',
            'sid': a.student.student_id if a.student else '',
            'status': a.status,
            'answered': answered.get(a.id, 0),
            'questions': qcount,
            'violations': a.violations or 0,
            'remaining': max(remaining, 0) if a.status == 'In progress' else 0,
            'seen_ago': seen_ago,
            'online': (a.status == 'In progress' and seen_ago is not None and seen_ago <= 35),
            'paused': paused,
            'score': a.score if a.status == 'Submitted' else None,
            'total': a.total,
            'device': device,
            'concurrent': concurrent,
            'concurrent_devices': [d.label for d in active_devs] if concurrent else [],
        })
    inprog = sum(1 for r in rows if r['status'] == 'In progress')
    return jsonify({'rows': rows, 'in_progress': inprog,
                    'submitted': sum(1 for r in rows if r['status'] == 'Submitted'),
                    'concurrent': sum(1 for r in rows if r['concurrent']),
                    'server_time': now.strftime('%H:%M:%S')})


def _review_rows(exam, attempt):
    """Per-question review: each question with the student's pick + correctness."""
    ans = {a.question_id: a for a in attempt.answers}
    rows = []
    for q in exam.questions.order_by(CBTQuestion.order, CBTQuestion.id).all():
        a = ans.get(q.id)
        rows.append({'q': q,
                     'selected': a.selected_option if a else None,
                     'is_correct': bool(a and a.is_correct)})
    return rows


@cbt_bp.route('/attempts/<int:attempt_id>/review')
@login_required
def attempt_review(attempt_id):
    attempt = db.get_or_404(CBTAttempt, attempt_id)
    rows = _review_rows(attempt.exam, attempt)
    violations = attempt.violation_log.order_by(CBTViolation.created_at).all()
    total_away = sum(v.away_seconds or 0 for v in violations)
    # Device fingerprints recorded for this student around this exam.
    logins = (CBTLoginEvent.query
              .filter(CBTLoginEvent.student_id == attempt.student_id)
              .filter((CBTLoginEvent.exam_id == attempt.exam_id)
                      | (CBTLoginEvent.exam_id == None))
              .order_by(CBTLoginEvent.created_at.desc()).limit(10).all())
    return render_template('cbt/attempt_review.html', attempt=attempt,
                           exam=attempt.exam, rows=rows, violations=violations,
                           total_away=total_away, logins=logins)



def _exam_sheet(ws, exam):
    """Write one exam's results (alphabetical by name) into an openpyxl sheet."""
    from openpyxl.styles import Font, PatternFill
    head_fill = PatternFill(start_color='0d6a4e', end_color='0d6a4e', fill_type='solid')
    head_font = Font(bold=True, color='FFFFFF')
    title = f'{exam.subject.name if exam.subject else "Exam"} — {exam.title}'
    ws.merge_cells('A1:E1')
    ws['A1'] = title
    ws['A1'].font = Font(bold=True, size=13)
    ws.merge_cells('A2:E2')
    ws['A2'] = (f'{exam.school_class.name if exam.school_class else ""} · '
                f'{exam.exam_date.strftime("%d %b %Y") if exam.exam_date else ""} · '
                f'Total {exam.total_marks} marks')
    ws['A2'].font = Font(italic=True, color='666666')
    headers = ['S/N', 'Student ID', 'Name', 'Raw', f'Score (/{exam.total_marks:g})', '%']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = head_fill
        cell.font = head_font
    # Submitted attempts, alphabetical by surname then first name.
    attempts = (exam.attempts.filter_by(status='Submitted')
                .join(Student).order_by(Student.surname, Student.first_name).all())
    r = 5
    for i, a in enumerate(attempts, 1):
        ws.cell(row=r, column=1, value=i)
        ws.cell(row=r, column=2, value=a.student.student_id if a.student else '')
        ws.cell(row=r, column=3, value=a.student.full_name if a.student else '')
        ws.cell(row=r, column=4, value=f'{a.raw_score:g}/{a.raw_total:g}' if a.raw_total else a.score)
        ws.cell(row=r, column=5, value=a.score)
        ws.cell(row=r, column=6, value=a.percentage)
        r += 1
    widths = [6, 16, 32, 12, 12, 8]
    for i, wd in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = wd
    return len(attempts)


def _safe_sheet_title(text):
    bad = set('[]:*?/\\')
    return ''.join(c for c in text if c not in bad)[:31] or 'Sheet'


@cbt_bp.route('/exams/<int:exam_id>/results/export')
@login_required
def results_export(exam_id):
    from openpyxl import Workbook
    e = db.get_or_404(CBTExam, exam_id)
    wb = Workbook()
    ws = wb.active
    ws.title = _safe_sheet_title(e.subject.name if e.subject else 'Results')
    _exam_sheet(ws, e)
    fname = _safe_sheet_title(f'{e.subject.name if e.subject else "exam"}_{e.title}') + '.xlsx'
    return xlsx_response(wb, fname)


@cbt_bp.route('/results/export-all')
@login_required
def results_export_all():
    """Workbook with one alphabetical results sheet per exam (grouped by subject)."""
    from openpyxl import Workbook
    wb = Workbook()
    wb.remove(wb.active)
    exams = (CBTExam.query.join(Subject, isouter=True)
             .order_by(Subject.name, CBTExam.exam_date).all())
    exams = [e for e in exams if e.attempts.filter_by(status='Submitted').count()]
    if not exams:
        ws = wb.create_sheet('No results')
        ws['A1'] = 'No submitted exam results yet.'
    seen = {}
    for e in exams:
        base = _safe_sheet_title(f'{e.subject.name if e.subject else "Exam"}-{e.title}')
        title, n = base, seen.get(base, 0)
        if n:
            title = _safe_sheet_title(f'{base} {n+1}')
        seen[base] = n + 1
        _exam_sheet(wb.create_sheet(title), e)
    return xlsx_response(wb, 'cbt_all_results.xlsx')


# ============================================================================
# ADMIN — STUDENT PORTAL PASSWORDS
# ============================================================================

@cbt_bp.route('/passwords', methods=['GET', 'POST'])
@login_required
def passwords():
    term = _active_term()
    class_id = request.values.get('class_id', type=int)
    arm_id = request.values.get('arm_id', type=int)
    classes = SchoolClass.query.filter_by(is_active=True).order_by(SchoolClass.level).all()
    arms = ClassArm.query.filter_by(is_active=True).order_by(ClassArm.name).all()

    if request.method == 'POST':
        action = request.form.get('action')
        ids = request.form.getlist('student_id', type=int)
        students = Student.query.filter(Student.id.in_(ids or [-1])).all()
        if action == 'set_individual':
            sid = request.form.get('one_student_id', type=int)
            pw = (request.form.get('one_password') or '').strip()
            s = db.session.get(Student, sid) if sid else None
            if s and pw:
                s.set_portal_password(pw)
                db.session.commit()
                flash(f'Password set for {s.full_name}.', 'success')
        elif action == 'set_same' and students:
            pw = (request.form.get('bulk_password') or '').strip()
            if pw:
                for s in students:
                    s.set_portal_password(pw)
                db.session.commit()
                flash(f'Set the same password for {len(students)} student(s).', 'success')
        elif action == 'generate' and students:
            generated = []
            for s in students:
                pw = secrets.token_hex(3)   # 6-char one-time password
                s.set_portal_password(pw)
                generated.append((s, pw))
            db.session.commit()
            return render_template('cbt/passwords.html', classes=classes, arms=arms, term=term,
                class_id=class_id, arm_id=arm_id,
                students=_password_roster(class_id, arm_id, term), generated=generated)
        return redirect(url_for('cbt.passwords', class_id=class_id or '', arm_id=arm_id or ''))

    return render_template('cbt/passwords.html', classes=classes, arms=arms, term=term,
        class_id=class_id, arm_id=arm_id,
        students=_password_roster(class_id, arm_id, term), generated=None)


def _password_roster(class_id, arm_id, term):
    if class_id and term:
        q = (StudentEnrollment.query
             .join(ClassArmAssignment,
                   StudentEnrollment.class_arm_assignment_id == ClassArmAssignment.id)
             .filter(StudentEnrollment.is_active == True,
                     ClassArmAssignment.term_id == term.id,
                     ClassArmAssignment.class_id == class_id))
        if arm_id:
            q = q.filter(ClassArmAssignment.arm_id == arm_id)
        ids = [e.student_id for e in q.all()]
        return Student.query.filter(Student.id.in_(ids or [-1])).order_by(Student.surname, Student.first_name).all()
    return Student.query.filter_by(is_active=True).order_by(Student.surname).limit(300).all()


def _roster_label(class_id, arm_id, classes, arms):
    cls = next((c for c in classes if c.id == class_id), None)
    arm = next((a for a in arms if a.id == arm_id), None)
    return ' '.join(p for p in [cls.name if cls else 'All students', arm.name if arm else ''] if p)


@cbt_bp.route('/passwords/export')
@login_required
def passwords_export():
    """Export portal passwords for a class arm as Excel / Word / PDF."""
    term = _active_term()
    class_id = request.args.get('class_id', type=int)
    arm_id = request.args.get('arm_id', type=int)
    fmt = (request.args.get('format') or 'xlsx').lower()
    classes = SchoolClass.query.all()
    arms = ClassArm.query.all()
    label = _roster_label(class_id, arm_id, classes, arms)
    students = _password_roster(class_id, arm_id, term)
    rows = [(i + 1, s.student_id, s.full_name, s.portal_password_plain or '—')
            for i, s in enumerate(students)]
    school = SchoolSettings.get('school_name', 'School')
    safe = ''.join(c for c in label if c.isalnum() or c in ' -_').strip().replace(' ', '_') or 'students'

    if fmt == 'docx':
        return _passwords_docx(school, label, rows, safe)
    if fmt == 'pdf':
        return _passwords_pdf(school, label, rows, safe)
    return _passwords_xlsx(school, label, rows, safe)


def _passwords_xlsx(school, label, rows, safe):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    wb = Workbook()
    ws = wb.active
    ws.title = 'Passwords'
    ws.merge_cells('A1:D1')
    ws['A1'] = f'{school} — Test Portal Passwords ({label})'
    ws['A1'].font = Font(bold=True, size=13)
    fill = PatternFill(start_color='0d6a4e', end_color='0d6a4e', fill_type='solid')
    for col, h in enumerate(['S/N', 'Student ID', 'Name', 'Password'], 1):
        c = ws.cell(row=3, column=col, value=h)
        c.fill = fill
        c.font = Font(bold=True, color='FFFFFF')
    for r, row in enumerate(rows, 4):
        for col, val in enumerate(row, 1):
            ws.cell(row=r, column=col, value=val)
    for i, w in enumerate([6, 18, 34, 16], 1):
        ws.column_dimensions[chr(64 + i)].width = w
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return Response(out.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=passwords_{safe}.xlsx'})


def _passwords_docx(school, label, rows, safe):
    from docx import Document
    doc = Document()
    doc.add_heading(f'{school}', level=1)
    doc.add_heading(f'Test Portal Passwords — {label}', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(['S/N', 'Student ID', 'Name', 'Password']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return Response(out.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=passwords_{safe}.docx'})


def _passwords_pdf(school, label, rows, safe):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm)
    styles = getSampleStyleSheet()
    elems = [Paragraph(f'<b>{school}</b>', styles['Title']),
             Paragraph(f'Test Portal Passwords — {label}', styles['Heading2']),
             Spacer(1, 6 * mm)]
    data = [['S/N', 'Student ID', 'Name', 'Password']] + [[str(c) for c in row] for row in rows]
    t = Table(data, colWidths=[14 * mm, 32 * mm, 78 * mm, 36 * mm], repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0d6a4e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#cccccc')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f3f6f4')]),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (3, 1), (3, -1), 'Courier-Bold'),
    ]))
    elems.append(t)
    doc.build(elems)
    return pdf_response(out, f'passwords_{safe}.pdf', inline=False)


# ============================================================================
# STUDENT PORTAL
# ============================================================================

def cbt_login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get(PORTAL_KEY):
            return redirect(url_for('cbt_portal.login'))
        return f(*args, **kwargs)
    return wrapper


def _current_student():
    sid = session.get(PORTAL_KEY)
    return db.session.get(Student, sid) if sid else None


def _record_login_event(student, event='login', exam_id=None):
    """Capture an IP + parsed-UA device fingerprint for investigations.

    Returns the event id so the client can enrich it (screen/timezone/geo).
    """
    try:
        ua = (request.headers.get('User-Agent') or '')
        info = parse_user_agent(ua)
        ev = CBTLoginEvent(
            student_id=student.id, exam_id=exam_id, event=event,
            ip_address=(request.headers.get('X-Forwarded-For', request.remote_addr) or '')[:60],
            user_agent=ua[:400], browser=info['browser'], os=info['os'],
            device_type=info['device_type'], is_mobile=info['is_mobile'],
            device_model=info.get('model'))
        db.session.add(ev)
        db.session.commit()
        return ev.id
    except Exception:
        db.session.rollback()
        return None


# A device is "still live" if it pinged within this window (≈2 missed pings).
DEVICE_ACTIVE_WINDOW = 50


def _touch_device_session(student, exam_id, client_fields=None):
    """Upsert the live device (per-browser token) the student is active on."""
    token = (request.form.get('client_token') or '').strip()[:64]
    if not student or not token:
        return None
    try:
        sess = CBTDeviceSession.query.filter_by(
            student_id=student.id, client_token=token).first()
        if not sess:
            sess = CBTDeviceSession(student_id=student.id, client_token=token,
                                    first_seen=timeutil.now())
            db.session.add(sess)
        if exam_id:
            sess.exam_id = exam_id
        info = parse_user_agent(request.headers.get('User-Agent') or '')
        sess.ip_address = (request.headers.get('X-Forwarded-For', request.remote_addr) or '')[:60]
        sess.browser = info['browser']
        sess.os = info['os']
        sess.device_type = info['device_type']
        sess.is_mobile = info['is_mobile']
        if info.get('model'):
            sess.device_model = info['model']
        for k in ('device_model', 'latitude', 'longitude'):
            if client_fields and client_fields.get(k) is not None:
                setattr(sess, k, client_fields[k])
        sess.last_seen = timeutil.now()
        db.session.commit()
        return sess
    except Exception:
        db.session.rollback()
        return None


@cbt_portal_bp.route('/login', methods=['GET', 'POST'])
def login():
    if session.get(PORTAL_KEY):
        return redirect(url_for('cbt_portal.home'))
    if request.method == 'POST':
        from utils.security import login_limiter
        rkey = f"cbt_login:{request.remote_addr or 'unknown'}"
        if login_limiter.is_rate_limited(rkey, max_attempts=15, window_minutes=15):
            wait = login_limiter.get_remaining_time(rkey, 15) // 60 + 1
            flash(f'Too many attempts. Try again in about {wait} minute(s).', 'error')
            return render_template('cbt/portal_login.html')
        student_id = (request.form.get('student_id') or '').strip()
        password = request.form.get('password') or ''
        student = Student.query.filter_by(student_id=student_id, is_active=True).first()
        if student and student.check_portal_password(password):
            login_limiter.clear_attempts(rkey)
            session[PORTAL_KEY] = student.id
            ev_id = _record_login_event(student, event='login')
            if ev_id:
                session['cbt_login_event'] = ev_id
            return redirect(url_for('cbt_portal.home'))
        login_limiter.record_attempt(rkey)
        flash('Invalid student ID or password.', 'error')
    return render_template('cbt/portal_login.html')


@cbt_portal_bp.route('/logout')
def logout():
    session.pop(PORTAL_KEY, None)
    return redirect(url_for('cbt_portal.login'))


@cbt_portal_bp.route('/')
@cbt_login_required
def home():
    student = _current_student()
    if not student:
        session.pop(PORTAL_KEY, None)
        return redirect(url_for('cbt_portal.login'))
    term = _active_term()
    class_id, arm_id = _student_placement(student.id, term)
    exams = []
    if class_id:
        q = CBTExam.query.filter(CBTExam.is_published == True,
                                 CBTExam.exam_date == timeutil.today(),
                                 CBTExam.class_id == class_id)
        q = q.filter((CBTExam.arm_id == None) | (CBTExam.arm_id == arm_id))
        # Only this student's branch (a shared class name must not leak across branches).
        if student.branch_id:
            q = q.filter((CBTExam.branch_id == student.branch_id) | (CBTExam.branch_id == None))
        exams = q.order_by(CBTExam.start_time, CBTExam.title).all()
    # attach this student's attempt status + availability window
    rows = []
    for e in exams:
        att = CBTAttempt.query.filter_by(exam_id=e.id, student_id=student.id).first()
        state, msg = e.access_state()
        rows.append({'exam': e, 'attempt': att, 'state': state, 'msg': msg})
    return render_template('cbt/portal_home.html', student=student, rows=rows,
                           has_class=bool(class_id))


def _student_can_access(student, exam):
    """The exam must be for the student's current class/arm."""
    class_id, arm_id = _student_placement(student.id, _active_term())
    if not class_id or exam.class_id != class_id:
        return False
    if exam.arm_id and exam.arm_id != arm_id:
        return False
    # Branch isolation: an exam stamped to a branch is only for that branch.
    if exam.branch_id and student.branch_id and exam.branch_id != student.branch_id:
        return False
    return True


@cbt_portal_bp.route('/<int:exam_id>/start', methods=['GET', 'POST'])
@cbt_login_required
def start(exam_id):
    student = _current_student()
    exam = db.get_or_404(CBTExam, exam_id)
    if not exam.is_published or not _student_can_access(student, exam):
        flash('This exam is not available for your class.', 'error')
        return redirect(url_for('cbt_portal.home'))
    existing = CBTAttempt.query.filter_by(exam_id=exam.id, student_id=student.id).first()
    if existing and existing.status == 'Submitted':
        return redirect(url_for('cbt_portal.result', exam_id=exam.id))
    state, msg = exam.access_state()
    if state != 'open' and not (existing and existing.status == 'In progress'):
        flash(f'This test is not open right now — {msg}.', 'error')
        return redirect(url_for('cbt_portal.home'))
    if request.method == 'POST':
        pw = (request.form.get('access_password') or '').strip()
        if pw != (exam.access_password or ''):
            flash('Incorrect exam password.', 'error')
            return render_template('cbt/portal_start.html', exam=exam, student=student)
        if not existing:
            existing = CBTAttempt(exam_id=exam.id, student_id=student.id,
                                  started_at=timeutil.now(), total=exam.total_marks,
                                  last_seen=timeutil.now(),
                                  ip_address=(request.headers.get('X-Forwarded-For', request.remote_addr) or '')[:50],
                                  user_agent=(request.headers.get('User-Agent') or '')[:255])
            db.session.add(existing)
            db.session.commit()
            _record_login_event(student, event='start', exam_id=exam.id)
        return redirect(url_for('cbt_portal.take', exam_id=exam.id))
    return render_template('cbt/portal_start.html', exam=exam, student=student)


def _deadline(attempt, exam):
    return attempt.started_at + timedelta(minutes=exam.duration_minutes or 30)


def _finalize(attempt, exam):
    """Grade from saved answers, apply the exam's scaling, mark Submitted."""
    raw_score = 0.0
    raw_total = exam.raw_total
    saved = {a.question_id: a for a in attempt.answers}
    for q in exam.questions.all():
        ans = saved.get(q.id)
        correct = bool(ans and ans.selected_option == q.correct_option)
        if ans:
            ans.is_correct = correct
        if correct:
            raw_score += q.marks or 0
    attempt.raw_score = raw_score
    attempt.raw_total = raw_total
    attempt.total = exam.total_marks
    attempt.score = exam.scale(raw_score)
    attempt.status = 'Submitted'
    attempt.submitted_at = timeutil.now()
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        from flask import current_app
        current_app.logger.exception(
            'CBT grading failed to commit for attempt %s (exam %s)',
            getattr(attempt, 'id', '?'), getattr(exam, 'id', '?'))
        raise   # surface the failure instead of leaving a half-graded attempt


@cbt_portal_bp.route('/<int:exam_id>/take')
@cbt_login_required
def take(exam_id):
    student = _current_student()
    exam = db.get_or_404(CBTExam, exam_id)
    if not _student_can_access(student, exam):
        flash('This exam is not available for your class.', 'error')
        return redirect(url_for('cbt_portal.home'))
    attempt = CBTAttempt.query.filter_by(exam_id=exam.id, student_id=student.id).first()
    if not attempt:
        return redirect(url_for('cbt_portal.start', exam_id=exam.id))
    if attempt.status == 'Submitted':
        return redirect(url_for('cbt_portal.result', exam_id=exam.id))
    # Timer already expired while away. Give a grace window first: if the
    # student went offline mid-exam, their latest answers may exist only on
    # their device — rendering the page (at 00:00) lets the browser restore
    # them and submit, instead of grading from stale server-side answers.
    remaining = int((_deadline(attempt, exam) - timeutil.now()).total_seconds())
    if remaining <= -OFFLINE_GRACE_SECONDS:
        _finalize(attempt, exam)
        flash('Time was up — your test was submitted automatically.', 'info')
        return redirect(url_for('cbt_portal.result', exam_id=exam.id))
    remaining = max(remaining, 0)
    questions = exam.questions.order_by(CBTQuestion.order, CBTQuestion.id).all()
    if exam.shuffle:
        random.Random(attempt.id).shuffle(questions)   # stable per attempt
    # Build per-question option lists, shuffled per attempt. Displayed labels
    # always read A, B, C, D top-to-bottom (less confusing), but the *content*
    # is shuffled and each radio carries the option's ORIGINAL letter as its
    # value — so the backend always knows the real option chosen, and grading
    # (selected_option vs correct_option) stays correct.
    qview = []
    for q in questions:
        opts = [(letter, text) for letter, text in q.options if text]
        if exam.shuffle:
            random.Random(attempt.id * 1009 + q.id).shuffle(opts)
        options = [{'label': 'ABCD'[i], 'value': orig, 'text': text}
                   for i, (orig, text) in enumerate(opts)]
        qview.append({'q': q, 'options': options})
    saved = {a.question_id: a.selected_option for a in attempt.answers}
    has_pin = bool((SchoolSettings.get('cbt_supervisor_pin', '') or '').strip())
    return render_template('cbt/portal_take.html', exam=exam, student=student,
        qview=qview, attempt=attempt, remaining=remaining, saved=saved,
        max_violations=(exam.violation_limit if exam.violation_limit is not None else 3),
        strict=bool(exam.strict_mode), has_supervisor_pin=has_pin)


# Event types that count toward the auto-submit limit (leaving the exam).
_COUNTING_TYPES = {'tab_switch', 'fullscreen_exit'}


@cbt_portal_bp.route('/<int:exam_id>/flag', methods=['POST'])
@cbt_login_required
def flag(exam_id):
    """Log an anti-malpractice event with its type, away-time and any detail."""
    student = _current_student()
    attempt = CBTAttempt.query.filter_by(exam_id=exam_id, student_id=student.id).first()
    if not attempt or attempt.status != 'In progress':
        return jsonify({'ok': False}), 400
    exam = db.session.get(CBTExam, exam_id)
    vtype = (request.form.get('type') or 'tab_switch').strip()[:30]
    away = request.form.get('away', type=int) or 0
    detail = (request.form.get('detail') or '').strip()[:500] or None

    # During a supervisor-authorised pause, log but never count.
    paused = bool(attempt.paused_until and timeutil.now() < attempt.paused_until)
    counts = (vtype in _COUNTING_TYPES) and not paused

    db.session.add(CBTViolation(attempt_id=attempt.id, vtype=vtype, detail=detail,
                                away_seconds=max(away, 0), counted=counts))
    if counts:
        attempt.violations = (attempt.violations or 0) + 1
    attempt.last_seen = timeutil.now()
    db.session.commit()

    limit = exam.violation_limit if exam.violation_limit is not None else 3
    over = counts and bool(limit) and attempt.violations >= limit
    if over:
        _finalize(attempt, exam)   # forced submit
    return jsonify({'ok': True, 'violations': attempt.violations, 'max': limit,
                    'counted': counts, 'paused': paused, 'autosubmit': over})


@cbt_portal_bp.route('/<int:exam_id>/supervisor-unlock', methods=['POST'])
@cbt_login_required
def supervisor_unlock(exam_id):
    """A supervisor enters the PIN to pause lockdown (e.g. to fix the laptop)."""
    student = _current_student()
    attempt = CBTAttempt.query.filter_by(exam_id=exam_id, student_id=student.id).first()
    if not attempt or attempt.status != 'In progress':
        return jsonify({'ok': False}), 400
    pin = (request.form.get('pin') or '').strip()
    real = (SchoolSettings.get('cbt_supervisor_pin', '') or '').strip()
    if not real or pin != real:
        return jsonify({'ok': False, 'error': 'Wrong supervisor PIN'}), 403
    minutes = request.form.get('minutes', type=int) or 3
    minutes = max(1, min(minutes, 15))
    from datetime import timedelta
    attempt.paused_until = timeutil.now() + timedelta(minutes=minutes)
    db.session.add(CBTViolation(attempt_id=attempt.id, vtype='resume',
                                detail=f'Supervisor unlock for {minutes} min', counted=False))
    db.session.commit()
    return jsonify({'ok': True, 'minutes': minutes})


@cbt_portal_bp.route('/<int:exam_id>/ping', methods=['POST'])
@cbt_login_required
def ping(exam_id):
    """Heartbeat so admins can spot a disconnected/abandoned attempt."""
    student = _current_student()
    attempt = CBTAttempt.query.filter_by(exam_id=exam_id, student_id=student.id).first()
    if attempt and attempt.status == 'In progress':
        # Coalesce: only persist last_seen if it's gone stale. Clients ping
        # often (~every 30s); writing on every ping is the main heartbeat write
        # load at scale. A <25s-old last_seen is fresh enough for monitoring.
        now = timeutil.now()
        if not attempt.last_seen or (now - attempt.last_seen).total_seconds() >= 25:
            attempt.last_seen = now
            db.session.commit()
            _touch_device_session(student, exam_id)
    return jsonify({'ok': True})


@cbt_portal_bp.route('/fingerprint', methods=['POST'])
@cbt_login_required
def fingerprint():
    """Enrich the latest login event with client-side device details (best effort)."""
    student = _current_student()
    if not student:
        return jsonify({'ok': False}), 400
    ev = None
    ev_id = session.get('cbt_login_event')
    if ev_id:
        ev = CBTLoginEvent.query.filter_by(id=ev_id, student_id=student.id).first()
    if not ev:
        ev = (CBTLoginEvent.query.filter_by(student_id=student.id)
              .order_by(CBTLoginEvent.id.desc()).first())
    if not ev:
        return jsonify({'ok': False}), 400

    def _s(key, n):
        return (request.form.get(key) or '').strip()[:n] or None
    ev.screen = _s('screen', 20)
    ev.viewport = _s('viewport', 20)
    ev.timezone = _s('timezone', 60)
    ev.language = _s('language', 40)
    ev.platform = _s('platform', 60)
    # UA Client Hints model (more accurate than the UA string) wins if provided.
    model = _s('model', 80)
    if model:
        ev.device_model = model
    cf = {}
    for col, key in (('latitude', 'lat'), ('longitude', 'lon'), ('geo_accuracy', 'acc')):
        try:
            val = request.form.get(key)
            if val not in (None, ''):
                setattr(ev, col, float(val))
                cf[col] = float(val)
        except (ValueError, TypeError):
            pass
    if model:
        cf['device_model'] = model
    db.session.commit()
    # Also register this as a live device for concurrent-session detection.
    _touch_device_session(student, ev.exam_id, client_fields=cf)
    return jsonify({'ok': True})


@cbt_portal_bp.route('/<int:exam_id>/answer', methods=['POST'])
@cbt_login_required
def answer(exam_id):
    """Autosave a single answer so a refresh / disconnect never loses progress."""
    student = _current_student()
    attempt = CBTAttempt.query.filter_by(exam_id=exam_id, student_id=student.id).first()
    if not attempt or attempt.status == 'Submitted':
        return jsonify({'ok': False}), 400
    qid = request.form.get('question_id', type=int)
    sel = (request.form.get('option') or '').strip().upper() or None
    q = CBTQuestion.query.filter_by(id=qid, exam_id=exam_id).first()
    if not q:
        return jsonify({'ok': False}), 400
    ans = CBTAnswer.query.filter_by(attempt_id=attempt.id, question_id=qid).first()
    if not ans:
        ans = CBTAnswer(attempt_id=attempt.id, question_id=qid)
        db.session.add(ans)
    ans.selected_option = sel
    ans.is_correct = (sel == q.correct_option)
    db.session.commit()
    return jsonify({'ok': True})


@cbt_portal_bp.route('/<int:exam_id>/answers', methods=['POST'])
@cbt_login_required
def answers_batch(exam_id):
    """Autosave several answers in ONE request/commit (the scalable autosave path).

    The client batches changes and flushes here periodically; the final submit
    still posts every answer from the form, so a missed batch never loses data.
    Accepts form fields named ``q_<question_id>=<option>``.
    """
    student = _current_student()
    attempt = CBTAttempt.query.filter_by(exam_id=exam_id, student_id=student.id).first()
    if not attempt or attempt.status == 'Submitted':
        return jsonify({'ok': False}), 400
    correct = {q.id: q.correct_option
               for q in CBTQuestion.query.filter_by(exam_id=exam_id).all()}
    existing = {a.question_id: a for a in attempt.answers}
    changed = 0
    for key, val in request.form.items():
        if not key.startswith('q_'):
            continue
        try:
            qid = int(key[2:])
        except ValueError:
            continue
        if qid not in correct:
            continue
        sel = (val or '').strip().upper() or None
        ans = existing.get(qid)
        if ans is None:
            ans = CBTAnswer(attempt_id=attempt.id, question_id=qid)
            db.session.add(ans)
            existing[qid] = ans
        ans.selected_option = sel
        ans.is_correct = (sel == correct[qid])
        changed += 1
    if changed:
        db.session.commit()
    return jsonify({'ok': True, 'saved': changed})


@cbt_portal_bp.route('/<int:exam_id>/submit', methods=['POST'])
@cbt_login_required
def submit(exam_id):
    student = _current_student()
    exam = db.get_or_404(CBTExam, exam_id)
    attempt = CBTAttempt.query.filter_by(exam_id=exam.id, student_id=student.id).first()
    if not attempt:
        return redirect(url_for('cbt_portal.home'))
    if attempt.status == 'Submitted':
        return redirect(url_for('cbt_portal.result', exam_id=exam.id))
    # Persist any answers posted with the final submit (covers JS-off / last picks).
    for q in exam.questions.all():
        if f'q_{q.id}' not in request.form:
            continue
        sel = (request.form.get(f'q_{q.id}') or '').strip().upper() or None
        ans = CBTAnswer.query.filter_by(attempt_id=attempt.id, question_id=q.id).first()
        if not ans:
            ans = CBTAnswer(attempt_id=attempt.id, question_id=q.id)
            db.session.add(ans)
        ans.selected_option = sel
    db.session.flush()
    _finalize(attempt, exam)
    return redirect(url_for('cbt_portal.result', exam_id=exam.id))


@cbt_portal_bp.route('/<int:exam_id>/result')
@cbt_login_required
def result(exam_id):
    student = _current_student()
    exam = db.get_or_404(CBTExam, exam_id)
    attempt = CBTAttempt.query.filter_by(exam_id=exam.id, student_id=student.id).first()
    if not attempt or attempt.status != 'Submitted':
        return redirect(url_for('cbt_portal.home'))
    rows = _review_rows(exam, attempt)
    return render_template('cbt/portal_result.html', exam=exam, student=student,
                           attempt=attempt, rows=rows)
