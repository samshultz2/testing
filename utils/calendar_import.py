"""
Scan a school calendar (Word .docx or an image/PDF) into draft calendar events.

Word docs are parsed by table structure (Week | Activity | Dates), pairing each
activity with its date. Images/PDFs go through Tesseract OCR and are parsed
line-by-line. Parsing is best-effort — results are always shown in an editable
review grid before anything is saved.
"""
import io
import re
from datetime import date


_FULL_RANGE = re.compile(r'(\d{1,2})/(\d{1,2})/(\d{2,4})\s*[–\-—]\s*(\d{1,2})/(\d{1,2})/(\d{2,4})')
_SHORT_RANGE = re.compile(r'(\d{1,2})\s*[–\-—]\s*(\d{1,2})/(\d{1,2})/(\d{2,4})')
_SINGLE = re.compile(r'(\d{1,2})/(\d{1,2})/(\d{2,4})')

_SKIP = {'activity', 'activities', 'date/week', 'dates', 'date', 'week',
         'one', 'two', 'three', 'four', 'five', 'six', 'seven', 'eight',
         'nine', 'ten', 'eleven', 'twelve', 'thirteen', 'fourteen'}

_CATEGORY_RULES = [
    ('Holiday', ('public holiday', 'mid-term break', 'mid term break', 'break',
                 "workers' day", 'eid', 'sallah', 'christmas', 'easter', 'new year')),
    ('Exam', ('exam', 'bece', 'waec', 'neco', 'jamb', 'practical', 'cbt', ' test')),
    ('Activity', ('competition', 'spelling bee', 'debate', 'quiz', 'fellowship',
                  'prayer', 'fasting', 'visitation', 'excursion', 'sport',
                  'inter-house', 'party', 'graduation', 'speech', 'cultural')),
    ('Meeting', ('meeting', 'pta', 'assembly', 'orientation')),
]


def _year(y):
    y = int(y)
    return 2000 + y if y < 100 else y


def _mk(d, m, y):
    try:
        return date(_year(y), int(m), int(d))
    except ValueError:
        return None


def parse_dates(text):
    """Return (start_date, end_date_or_None) from a token, or (None, None)."""
    if not text:
        return None, None
    m = _FULL_RANGE.search(text)
    if m:
        s = _mk(m.group(1), m.group(2), m.group(3))
        e = _mk(m.group(4), m.group(5), m.group(6))
        return s, (e if e and s and e != s else None)
    m = _SHORT_RANGE.search(text)      # "02 - 06/06/2025"  /  "26 – 30/05/26"
    if m:
        e = _mk(m.group(2), m.group(3), m.group(4))
        s = _mk(m.group(1), m.group(3), m.group(4))
        return s, (e if e and s and e != s else None)
    m = _SINGLE.search(text)
    if m:
        return _mk(m.group(1), m.group(2), m.group(3)), None
    return None, None


def _infer_category(text):
    low = (text or '').lower()
    for cat, words in _CATEGORY_RULES:
        if any(w in low for w in words):
            return cat
    return 'General'


def _clean_title(text):
    # strip any trailing/inline date tokens and tidy whitespace
    t = _SINGLE.sub('', text)
    t = re.sub(r'\s*[–\-—]\s*$', '', t).strip(' .,-–—')
    return re.sub(r'\s{2,}', ' ', t).strip()


def _is_skippable(text):
    low = re.sub(r'[^a-z/ ]', '', (text or '').lower()).strip()
    return not low or low in _SKIP


def parse_docx(file_bytes):
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    events = []
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            # Week cell -> the week's start date (fallback for undated activities).
            wk_start, _ = parse_dates(cells[0].text)
            # IMPORTANT: keep blank paragraphs so activity[i] lines up with date[i]
            # (the dates column pads with empty paragraphs for undated activities).
            acts = [p.text.strip() for p in cells[1].paragraphs]
            dates = [p.text.strip() for p in cells[-1].paragraphs] if len(cells) >= 3 else []
            if not any(acts) or _is_skippable(' '.join(a for a in acts if a)):
                continue
            for i, act in enumerate(acts):
                if not act or _is_skippable(act):
                    continue
                dstr = dates[i] if i < len(dates) else ''
                s, e = parse_dates(dstr)
                if not s:
                    s, e = parse_dates(act)        # date inline in the activity text
                if not s:
                    s = wk_start                   # undated -> week's start date
                if not s:
                    continue
                events.append({'start_date': s, 'end_date': e,
                               'title': _clean_title(act) or act.strip(),
                               'category': _infer_category(act)})
    if not events:
        events = parse_text('\n'.join(p.text for p in doc.paragraphs))
    return _dedupe(events)


def parse_text(text):
    """Parse OCR'd / plain text line by line: any line with a date + words."""
    events = []
    for raw in (text or '').splitlines():
        line = raw.strip()
        if not line or _is_skippable(line):
            continue
        s, e = parse_dates(line)
        if not s:
            continue
        title = _clean_title(line)
        if len(title) < 3:
            continue
        events.append({'start_date': s, 'end_date': e,
                       'title': title, 'category': _infer_category(line)})
    return _dedupe(events)


def parse_image(file_bytes, filename=''):
    from utils.waec_ocr import extract_text, extract_text_from_pdf
    if filename.lower().endswith('.pdf'):
        text = extract_text_from_pdf(file_bytes)
    else:
        text = extract_text(file_bytes)
    return parse_text(text)


def _dedupe(events):
    seen = set()
    out = []
    for ev in events:
        key = (ev['start_date'], ev['title'].lower())
        if key in seen:
            continue
        seen.add(key)
        out.append(ev)
    out.sort(key=lambda e: (e['start_date'] or date.max, e['title']))
    return out
