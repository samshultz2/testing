"""Broadsheet & score-sheet exports in multiple formats.

Two products:

1. A **filled broadsheet** (subject totals per student, ranked) rendered as PDF,
   Word (.docx) or a high-resolution PNG — alongside the existing Excel export.
2. A **blank score-entry sheet** (A4, one subject) with the school's own
   assessment columns (any number of CAs, optional Holiday Assignment, optional
   Practical/Mid-term, CBT, PBT/Theory, Exam Total, General Total) and the class
   roster pre-printed in separate First / Middle / Surname columns — ready to
   photocopy and fill by hand.

Columns are derived from the tenant's active ``AssessmentType`` rows, so a school
that runs 2 CAs and no Holiday Assignment gets exactly those columns.
"""
import io

from models import (db, Term, ClassArmAssignment, ClassSubject, Subject, Student,
                    StudentEnrollment, StudentScore, AssessmentType, SchoolSettings, GradeScale)
from utils.assessments import is_midterm, is_theory, is_cbt
from utils.numfmt import fmt_num as _n


# --------------------------------------------------------------------------- #
# Assessment-column layout (school-configurable)
# --------------------------------------------------------------------------- #

def _is_ca(at):
    code = (at.short_name or at.name or '').upper().replace(' ', '')
    return code.startswith('CA') or 'CONTINUOUS' in code


def _is_ha(at):
    code = (at.short_name or '').upper()
    name = (at.name or '').upper()
    return code in ('HA', 'HW') or 'HOLIDAY' in name or 'ASSIGNMENT' in name


def score_columns():
    """Ordered blank-sheet score columns from the active assessment types.

    Returns a list of dicts ``{'label', 'max', 'kind'}`` where ``kind`` is one of
    ca / ha / peme / cbt / theory / other, deduplicated by short name and ordered
    as CAs → HA → Practical/Mid-term → CBT → Theory → anything else."""
    types = AssessmentType.query.filter_by(is_active=True).order_by(
        AssessmentType.order, AssessmentType.id).all()
    seen, uniq = set(), []
    for at in types:
        key = (at.short_name or at.name or '').upper()
        if key and key not in seen:
            seen.add(key)
            uniq.append(at)

    def label_for(at, kind):
        if kind == 'cbt':
            return 'CBT'
        if kind == 'theory':
            return 'PBT/THEORY'
        if kind == 'peme':
            return 'P.E/M.E'
        if kind == 'ha':
            return at.short_name or 'HA'
        return at.short_name or at.name

    buckets = {'ca': [], 'ha': [], 'peme': [], 'cbt': [], 'theory': [], 'other': []}
    for at in uniq:
        if _is_ca(at):
            kind = 'ca'
        elif _is_ha(at):
            kind = 'ha'
        elif is_midterm(at):
            kind = 'peme'
        elif is_cbt(at):
            kind = 'cbt'
        elif is_theory(at):
            kind = 'theory'
        else:
            kind = 'other'
        buckets[kind].append({'label': label_for(at, kind), 'max': at.max_score, 'kind': kind})
    ordered = (buckets['ca'] + buckets['ha'] + buckets['peme'] + buckets['other']
               + buckets['cbt'] + buckets['theory'])
    return ordered


# --------------------------------------------------------------------------- #
# Shared filled-broadsheet model
# --------------------------------------------------------------------------- #

def _passes_filter(row, field, min_val):
    """Whether a broadsheet row meets a ``field >= min_val`` threshold. ``field``
    is 'average', 'total' or a class-subject id (int/str). Missing cells fail."""
    if field == 'average':
        v = row['average']
    elif field == 'total':
        v = row['total']
    else:
        try:
            v = row['subjects'].get(int(field))
        except (TypeError, ValueError):
            v = None
    return v is not None and v >= min_val


def build_model(term_id, assignment_id, min_score=None, filter_field=None):
    """Shared broadsheet data: ranked students × subject totals. Returns a dict
    with ``term``, ``assignment``, ``subjects`` and ``rows`` (already ranked).

    When ``min_score`` is given, rows are filtered to those scoring at or above it
    in ``filter_field`` ('average' default, 'total' or a subject id) — applied
    AFTER ranking so each shown student keeps their true arm position."""
    term = db.session.get(Term, term_id)
    asg = db.session.get(ClassArmAssignment, assignment_id)
    if not (term and asg):
        return None
    subjects = ClassSubject.query.filter_by(
        term_id=term_id, class_id=asg.class_id, is_active=True
    ).filter((ClassSubject.arm_id == None) | (ClassSubject.arm_id == asg.arm_id)  # noqa: E711
             ).join(Subject).order_by(Subject.name).all()
    enrollments = (StudentEnrollment.query
                   .filter_by(class_arm_assignment_id=assignment_id, is_active=True)
                   .join(Student).order_by(Student.surname, Student.first_name).all())
    pass_mark = SchoolSettings.get('pass_mark', 50)

    cs_ids = [cs.id for cs in subjects]
    sids = [e.student_id for e in enrollments]
    totals = {}                        # (student_id, cs_id) -> summed score
    if cs_ids and sids:
        for s in StudentScore.query.filter(
                StudentScore.student_id.in_(sids),
                StudentScore.class_subject_id.in_(cs_ids)).all():
            totals[(s.student_id, s.class_subject_id)] = totals.get(
                (s.student_id, s.class_subject_id), 0) + (s.score or 0)

    rows = []
    for e in enrollments:
        subj_totals = {cs.id: totals.get((e.student_id, cs.id), 0) for cs in subjects}
        grand = sum(subj_totals.values())
        avg = round(grand / len(subjects), 2) if subjects else 0
        passed = sum(1 for v in subj_totals.values() if v >= pass_mark)
        failed = sum(1 for v in subj_totals.values() if 0 < v < pass_mark)
        rows.append({'student': e.student, 'subjects': subj_totals, 'total': round(grand, 1),
                     'average': avg, 'passed': passed, 'failed': failed})
    rows.sort(key=lambda r: r['average'], reverse=True)
    for i, r in enumerate(rows, 1):
        r['position'] = i
    if min_score is not None:
        field = filter_field or 'average'
        rows = [r for r in rows if _passes_filter(r, field, min_score)]
    return {'term': term, 'assignment': asg, 'subjects': subjects, 'rows': rows,
            'pass_mark': pass_mark}


def _school_name():
    try:
        from utils.school import school_profile
        return school_profile().get('name') or 'School'
    except Exception:
        return 'School'


def _fname(assignment, term, ext):
    base = (assignment.display_name or 'broadsheet').replace(' ', '_')
    return f"broadsheet_{base}_{term.name.replace(' ', '_')}.{ext}"


# --------------------------------------------------------------------------- #
# Filled broadsheet — PDF (reportlab, landscape A4)
# --------------------------------------------------------------------------- #

def _theme():
    from reportlab.lib.colors import HexColor
    return HexColor('#0D6A4E'), HexColor('#C9A227'), HexColor('#F4F7F5'), HexColor('#14211C')


def _neutral():
    """A professional, brand-neutral slate palette for the generic table exports
    (combination / broadsheet downloads) — no green. Returns
    (header, accent, zebra, ink)."""
    from reportlab.lib.colors import HexColor
    return HexColor('#334155'), HexColor('#64748B'), HexColor('#F1F5F9'), HexColor('#0F172A')

# Neutral RGB tuples for the Pillow (image) exports.
_NEUTRAL_RGB = {
    'header': (51, 65, 85),      # slate-700  (title/heading text)
    'text': (15, 23, 42),        # slate-900
    'muted': (100, 116, 139),    # slate-500
    'zebra': (245, 247, 250),    # very light row stripe
    'line': (213, 219, 227),     # light gridline
    'white': (255, 255, 255),
    # Print-friendly table header: a light fill with dark text + a rule, instead
    # of a heavy dark band (saves toner and reads cleanly in black & white).
    'head_bg': (233, 237, 242),
    'head_fg': (15, 23, 42),
    'rule': (148, 163, 184),
}


def broadsheet_pdf(term_id, assignment_id, mono=True, min_score=None, filter_field=None):
    """Filled broadsheet as a PDF.

    ``mono=True`` (the downloadable PDF) is plain black-on-white — no colour or
    shading, clean to print / photocopy. ``mono=False`` keeps the branded design
    (green header, zebra rows) and is what the HD-image export renders from."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer)
    from utils.web_exports import pdf_escape

    m = build_model(term_id, assignment_id, min_score=min_score, filter_field=filter_field)
    if not m:
        return None
    primary, accent, light, ink = _theme()
    title_color = colors.black if mono else primary
    sub_color = colors.black if mono else colors.HexColor('#6B7A74')
    styles = getSampleStyleSheet()
    h = ParagraphStyle('h', parent=styles['Title'], fontSize=15, textColor=title_color, spaceAfter=2)
    sub = ParagraphStyle('sub', parent=styles['Normal'], fontSize=9.5, textColor=sub_color)
    cell = ParagraphStyle('c', parent=styles['Normal'], fontSize=7.5, leading=9)

    subjects = m['subjects']
    head = ['Pos', 'Student'] + [(cs.subject.short_name or cs.subject.name[:6]) for cs in subjects] \
        + ['Total', 'Avg', 'Grade']
    data = [head]
    for r in m['rows']:
        row = [str(r['position']), Paragraph(pdf_escape(r['student'].full_name), cell)]
        for cs in subjects:
            v = r['subjects'].get(cs.id, 0)
            row.append(_n(v) if v else '–')
        row += [_n(r['total']), _n(r['average']),
                GradeScale.get_grade(r['average']) if r['average'] else '–']
        data.append(row)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), topMargin=10 * mm, bottomMargin=10 * mm,
                            leftMargin=8 * mm, rightMargin=8 * mm,
                            title=f"Broadsheet — {m['assignment'].display_name}")
    avail = landscape(A4)[0] - 16 * mm
    name_w, pos_w, tail_w = 44 * mm, 9 * mm, 12 * mm
    subj_w = max(9 * mm, (avail - pos_w - name_w - 3 * tail_w) / max(len(subjects), 1))
    widths = [pos_w, name_w] + [subj_w] * len(subjects) + [tail_w, tail_w, tail_w]

    t = Table(data, colWidths=widths, repeatRows=1)
    ts = TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 7.5),
        ('ALIGN', (2, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (-3, 1), (-1, -1), 'Helvetica-Bold'),
        ('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ])
    if mono:
        # Plain black on white — no fills or shading.
        ts.add('TEXTCOLOR', (0, 0), (-1, -1), colors.black)
        ts.add('GRID', (0, 0), (-1, -1), 0.4, colors.black)
        ts.add('LINEBELOW', (0, 0), (-1, 0), 1.2, colors.black)
    else:
        # Branded design (used by the HD-image export).
        ts.add('BACKGROUND', (0, 0), (-1, 0), primary)
        ts.add('TEXTCOLOR', (0, 0), (-1, 0), colors.white)
        ts.add('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#D5DED9'))
        ts.add('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light])
    t.setStyle(ts)

    elems = [Paragraph(pdf_escape(_school_name()), h),
             Paragraph(f"Broadsheet · {pdf_escape(m['assignment'].display_name)} · "
                       f"{pdf_escape(m['term'].full_name)}", sub),
             Spacer(1, 6), t]
    doc.build(elems)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Filled broadsheet — Word (.docx)
# --------------------------------------------------------------------------- #

def broadsheet_docx(term_id, assignment_id, min_score=None, filter_field=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    m = build_model(term_id, assignment_id, min_score=min_score, filter_field=filter_field)
    if not m:
        return None
    subjects = m['subjects']
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for attr in ('left_margin', 'right_margin', 'top_margin', 'bottom_margin'):
        setattr(sec, attr, Inches(0.4))

    title = doc.add_paragraph()
    run = title.add_run(_school_name()); run.bold = True; run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0D, 0x6A, 0x4E)
    meta = doc.add_paragraph(f"Broadsheet · {m['assignment'].display_name} · {m['term'].full_name}")
    meta.runs[0].font.size = Pt(10)

    head = ['Pos', 'Student'] + [(cs.subject.short_name or cs.subject.name[:6]) for cs in subjects] \
        + ['Total', 'Avg', 'Grade']
    table = doc.add_table(rows=1, cols=len(head))
    table.style = 'Light Grid Accent 1'
    for i, htext in enumerate(head):
        c = table.rows[0].cells[i]
        c.text = htext
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in p.runs:
                rn.bold = True; rn.font.size = Pt(8)
    for r in m['rows']:
        cells = table.add_row().cells
        cells[0].text = str(r['position'])
        cells[1].text = r['student'].full_name
        idx = 2
        for cs in subjects:
            v = r['subjects'].get(cs.id, 0)
            cells[idx].text = _n(v) if v else '–'; idx += 1
        cells[idx].text = _n(r['total'])
        cells[idx + 1].text = _n(r['average'])
        cells[idx + 2].text = GradeScale.get_grade(r['average']) if r['average'] else '–'
        for c in cells:
            for p in c.paragraphs:
                for rn in p.runs:
                    rn.font.size = Pt(8)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Filled broadsheet — high-resolution PNG (render the PDF at high DPI)
# --------------------------------------------------------------------------- #

def broadsheet_png(term_id, assignment_id, dpi=200, min_score=None, filter_field=None):
    """Render the broadsheet PDF to a single high-resolution PNG.

    Every page is rendered and the pages are stacked vertically into one image,
    so a broadsheet whose roster spills onto several pages exports in full
    (previously only the first page was captured). Keeps the branded colour
    design — only the downloadable PDF is monochrome."""
    pdf = broadsheet_pdf(term_id, assignment_id, mono=False, min_score=min_score, filter_field=filter_field)
    if not pdf:
        return None
    import fitz
    doc = fitz.open(stream=pdf, filetype='pdf')
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    pixmaps = [page.get_pixmap(matrix=matrix) for page in doc]
    if len(pixmaps) == 1:
        return pixmaps[0].tobytes('png')

    # Stitch all pages into one tall image (white gutter between pages).
    from PIL import Image
    imgs = [Image.frombytes('RGB', (p.width, p.height), p.samples) for p in pixmaps]
    gap = max(1, dpi // 20)
    width = max(im.width for im in imgs)
    height = sum(im.height for im in imgs) + gap * (len(imgs) - 1)
    canvas = Image.new('RGB', (width, height), 'white')
    y = 0
    for im in imgs:
        canvas.paste(im, ((width - im.width) // 2, y))
        y += im.height + gap
    out = io.BytesIO()
    canvas.save(out, format='PNG')
    return out.getvalue()


# --------------------------------------------------------------------------- #
# Blank score-entry sheet (A4, one subject) — reportlab
# --------------------------------------------------------------------------- #

def blank_sheet_pdf(term_id, assignment_id, subject_name=''):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer)
    from utils.web_exports import pdf_escape

    term = db.session.get(Term, term_id)
    asg = db.session.get(ClassArmAssignment, assignment_id)
    if not (term and asg):
        return None
    students = [e.student for e in (StudentEnrollment.query
                .filter_by(class_arm_assignment_id=assignment_id, is_active=True)
                .join(Student).order_by(Student.surname, Student.first_name).all())]
    cols = score_columns()
    # Plain black-and-white — no fills or shading (clean to photocopy).
    styles = getSampleStyleSheet()
    h = ParagraphStyle('h', parent=styles['Title'], fontSize=15, textColor=colors.black, spaceAfter=1)
    meta = ParagraphStyle('m', parent=styles['Normal'], fontSize=10, textColor=colors.black, leading=15)
    hd = ParagraphStyle('hd', parent=styles['Normal'], fontSize=6.8, leading=7.8,
                        alignment=1, textColor=colors.black, fontName='Helvetica-Bold')
    nm = ParagraphStyle('nm', parent=styles['Normal'], fontSize=8.5, leading=10)

    # Header row: S/N, First, Middle, Surname, <score cols…>, Exam Total, General Total.
    # No per-column max is printed — the max differs by subject, so the sheet stays generic.
    def col_head(label, mx=None):
        return Paragraph(label, hd)
    header = [Paragraph('S/N', hd), Paragraph('First Name', hd),
              Paragraph('Middle Name', hd), Paragraph('Surname', hd)]
    for c in cols:
        header.append(col_head(c['label'], c['max']))
    header += [col_head('Exam Total'), col_head('General Total')]
    data = [header]
    for i, st in enumerate(students, 1):
        row = [str(i), Paragraph(pdf_escape(st.first_name or ''), nm),
               Paragraph(pdf_escape(st.middle_name or ''), nm),
               Paragraph(pdf_escape(st.surname or ''), nm)]
        row += [''] * (len(cols) + 2)          # blank score cells
        data.append(row)
    # A couple of spare blank rows for late entrants.
    for _ in range(2):
        data.append([''] * (4 + len(cols) + 2))

    buf = io.BytesIO()
    # Portrait A4 gives the vertical room to keep a full class on a single page.
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=8 * mm, bottomMargin=8 * mm,
                            leftMargin=8 * mm, rightMargin=8 * mm,
                            title=f"Score sheet — {asg.display_name}")
    from reportlab.pdfbase.pdfmetrics import stringWidth
    avail = A4[0] - 16 * mm
    sn_w = 9 * mm

    # Name columns fit their contents (roster + header label), within sane bounds.
    def _fit(label, values, mn=18 * mm, mx=34 * mm):
        w = stringWidth(label, 'Helvetica-Bold', 7)
        for v in values:
            w = max(w, stringWidth(v or '', 'Helvetica', 8.5))
        return min(max(w + 6 * mm, mn), mx)
    first_w = _fit('First Name', [s.first_name for s in students])
    mid_w = _fit('Middle Name', [s.middle_name for s in students])
    sur_w = _fit('Surname', [s.surname for s in students])

    # Score columns share the remaining width (min 9mm each) so the grid still
    # spans the page even when the names are short.
    n_score = len(cols) + 2
    names_w = sn_w + first_w + mid_w + sur_w
    score_w = max(9 * mm, (avail - names_w) / n_score)
    widths = [sn_w, first_w, mid_w, sur_w] + [score_w] * n_score
    # If long names + many columns overspill the page width, shrink names to fit.
    overflow = sum(widths) - avail
    if overflow > 0:
        shrink = min(overflow, first_w + mid_w + sur_w - 3 * (16 * mm))
        if shrink > 0:
            scale = 1 - shrink / (first_w + mid_w + sur_w)
            first_w, mid_w, sur_w = first_w * scale, mid_w * scale, sur_w * scale
            widths = [sn_w, first_w, mid_w, sur_w] + [score_w] * n_score

    # Dynamic row height: fill the page's vertical space so a class up to ~45
    # students stays on one page, but each row keeps a comfortable writing floor
    # (~5.2mm) so larger classes spill onto a second page rather than cramming.
    body_rows = len(data) - 1
    header_h = 13 * mm
    # Vertical room left for the table after the title block and margins.
    body_avail = A4[1] - 16 * mm - 26 * mm - header_h
    row_h = max(5.2 * mm, min(10 * mm, body_avail / max(body_rows, 1)))
    heights = [header_h] + [row_h] * body_rows

    t = Table(data, colWidths=widths, rowHeights=heights, repeatRows=1)
    t.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ALIGN', (4, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTSIZE', (0, 1), (0, -1), 8),
        ('LINEBELOW', (0, 0), (-1, 0), 1.2, colors.black),
    ]))

    subj = subject_name or '__________________________'
    metaline = (f"<b>Class:</b> {pdf_escape(asg.display_name)} &nbsp;&nbsp; "
                f"<b>Term:</b> {pdf_escape(term.name)} &nbsp;&nbsp; "
                f"<b>Session:</b> {pdf_escape(term.session.name if term.session else '')} &nbsp;&nbsp; "
                f"<b>Subject:</b> {pdf_escape(subj)}")
    elems = [Paragraph('Continuous Assessment / Examination Score Sheet', h),
             Paragraph(metaline, meta), Spacer(1, 5), t]
    doc.build(elems)
    return buf.getvalue()


def blank_sheet_filename(assignment, term):
    base = (assignment.display_name or 'class').replace(' ', '_')
    return f"score_sheet_{base}_{term.name.replace(' ', '_')}.pdf"


# --------------------------------------------------------------------------- #
# Class analytics report (PDF) — KPIs + distributions + subject difficulty
# --------------------------------------------------------------------------- #

def analytics_pdf(term_id, assignment_id):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer)
    from utils.web_exports import pdf_escape
    from utils.results_analytics import class_analytics

    term = db.session.get(Term, term_id)
    asg = db.session.get(ClassArmAssignment, assignment_id)
    if not (term and asg):
        return None
    a = class_analytics(term_id, assignment_id, use_cache=False)
    s = (a or {}).get('summary') or {}
    if not s.get('assessed'):
        return None
    primary, accent, light, ink = _theme()
    styles = getSampleStyleSheet()
    h = ParagraphStyle('h', parent=styles['Title'], fontSize=16, textColor=primary, spaceAfter=1)
    sub = ParagraphStyle('sub', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#6B7A74'))
    hh = ParagraphStyle('hh', parent=styles['Heading2'], fontSize=12, textColor=primary, spaceBefore=8, spaceAfter=4)
    muted = ParagraphStyle('mu', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#6B7A74'))

    def kpi_table():
        cells = [
            ('Class average', _n(s.get('class_average'))), ('Pass rate', f"{_n(s.get('pass_rate'))}%"),
            ('Highest avg', _n(s.get('highest'))), ('Lowest avg', _n(s.get('lowest'))),
            ('Students assessed', f"{s.get('assessed')}/{s.get('students')}"),
            ('Entry completion', f"{_n(s.get('completion'))}%"),
        ]
        if s.get('trend') is not None:
            cells.append(('vs last term', f"{'+' if s['trend'] > 0 else ''}{_n(s['trend'])}"))
        rows, row = [], []
        for i, (lbl, val) in enumerate(cells):
            row.append(Paragraph(f"<b><font size=15 color='#0D6A4E'>{pdf_escape(val)}</font></b>"
                                 f"<br/><font size=8 color='#6B7A74'>{pdf_escape(lbl)}</font>", muted))
            if len(row) == 3:
                rows.append(row); row = []
        if row:
            row += [''] * (3 - len(row)); rows.append(row)
        t = Table(rows, colWidths=[(A4[0] - 30 * mm) / 3] * 3)
        t.setStyle(TableStyle([('BOX', (0, 0), (-1, -1), 0, colors.white),
                               ('BACKGROUND', (0, 0), (-1, -1), light),
                               ('INNERGRID', (0, 0), (-1, -1), 3, colors.white),
                               ('TOPPADDING', (0, 0), (-1, -1), 8), ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                               ('LEFTPADDING', (0, 0), (-1, -1), 10)]))
        return t

    def _bar_table(pairs, unit=''):
        mx = max([p[1] for p in pairs] + [1])
        data = []
        for label, val in pairs:
            w = int(round(38 * (val / mx))) if mx else 0
            bar = '█' * w
            data.append([Paragraph(pdf_escape(label), muted),
                         Paragraph(f"<font color='#0D6A4E'>{bar}</font> {pdf_escape(val)}{unit}", muted)])
        t = Table(data, colWidths=[45 * mm, A4[0] - 75 * mm])
        t.setStyle(TableStyle([('FONTSIZE', (0, 0), (-1, -1), 8),
                               ('TOPPADDING', (0, 0), (-1, -1), 1), ('BOTTOMPADDING', (0, 0), (-1, -1), 1)]))
        return t

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=14 * mm, bottomMargin=14 * mm,
                            leftMargin=15 * mm, rightMargin=15 * mm,
                            title=f"Analytics — {asg.display_name}")
    elems = [Paragraph(pdf_escape(_school_name()), h),
             Paragraph(f"Class Performance Report · {pdf_escape(asg.display_name)} · "
                       f"{pdf_escape(term.full_name)}", sub),
             Spacer(1, 8), kpi_table(), Spacer(1, 4)]

    if a.get('score_bands'):
        elems += [Paragraph('Distribution of student averages', hh),
                  _bar_table([(b['band'], b['count']) for b in a['score_bands']])]
    if a.get('grade_distribution'):
        elems += [Paragraph('Grade distribution', hh),
                  _bar_table([(f"Grade {g['grade']}", g['count']) for g in a['grade_distribution']])]
    if a.get('gender'):
        elems += [Paragraph('By gender', hh),
                  _bar_table([(f"{g['group']} (avg {_n(g['average'])}, {_n(g['pass_rate'])}% pass)", g['count'])
                              for g in a['gender']])]
    if a.get('subjects'):
        elems += [Paragraph('Subject difficulty (hardest first)', hh)]
        srows = [['Subject', 'Average', 'Pass %', 'Assessed']]
        for sub_ in a['subjects']:
            srows.append([sub_['name'], _n(sub_['average']) if sub_['assessed'] else '—',
                          (_n(sub_['pass_rate']) + '%') if sub_['assessed'] else '—',
                          str(sub_['assessed'])])
        st = Table(srows, colWidths=[(A4[0] - 30 * mm) * f for f in (0.46, 0.18, 0.18, 0.18)], repeatRows=1)
        st.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (1, 0), (-1, -1), 'CENTER'), ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#D5DED9')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light]),
            ('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 3)]))
        elems.append(st)
    if a.get('intervention'):
        elems += [Paragraph(f"Needs attention ({len(a['intervention'])})", hh)]
        irows = [['Student', 'Average', 'Failing']]
        for st_ in a['intervention']:
            irows.append([st_['name'], str(st_['average']), str(st_['failed'])])
        it = Table(irows, colWidths=[(A4[0] - 30 * mm) * f for f in (0.6, 0.2, 0.2)], repeatRows=1)
        it.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#B43A2E')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white), ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9), ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#D5DED9')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light]),
            ('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 3)]))
        elems.append(it)

    doc.build(elems)
    return buf.getvalue()


def analytics_filename(assignment, term):
    base = (assignment.display_name or 'class').replace(' ', '_')
    return f"analytics_{base}_{term.name.replace(' ', '_')}.pdf"


# --------------------------------------------------------------------------- #
# Results Explorer — combined cross-class / cross-arm exports
# --------------------------------------------------------------------------- #
# These take the already-built Explorer dataset (subjects union, per-scope meta
# and the filtered rows) rather than a single assignment, so a filtered view
# spanning several classes/arms exports as one Excel workbook or one print-ready
# PDF — matching what the user sees on screen.

def _explore_scope_stats(rows, scope_meta, field=None, pass_mark=50):
    """Per-scope summary over ``rows`` (already filtered): count, mean average,
    pass rate. Grouped by assignment; ordered to match ``scope_meta``."""
    order = [m['assignment_id'] for m in scope_meta]
    label = {m['assignment_id']: m['label'] for m in scope_meta}
    buckets = {aid: [] for aid in order}
    for r in rows:
        buckets.setdefault(r['assignment_id'], []).append(r)
    out = []
    for aid in order:
        rs = buckets.get(aid, [])
        n = len(rs)
        mean_avg = round(sum(r['average'] for r in rs) / n, 2) if n else None
        passed = sum(1 for r in rs if r['average'] >= pass_mark)
        out.append({'label': label.get(aid, ''), 'n': n,
                    'mean_avg': mean_avg,
                    'pass_rate': round(passed / n * 100) if n else None})
    return out


def explore_xlsx(subjects, scope_meta, rows, term_name, pass_mark=50, filter_label=''):
    """Filtered cross-class Explorer view as a two-sheet Excel workbook
    (Students + Comparison)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook()
    head_font = Font(bold=True, color='FFFFFF')
    head_fill = PatternFill('solid', fgColor='0D6A4E')
    centre = Alignment(horizontal='center')

    ws = wb.active
    ws.title = 'Students'
    if filter_label:
        ws.append([f'Filter: {filter_label}'])
        ws.append([])
    header = ['Position', 'Student', 'Class', 'Arm'] + [s['name'] for s in subjects] \
        + ['Total', 'Average', 'Grade', 'Passed', 'Failed']
    ws.append(header)
    hrow = ws.max_row
    for c in range(1, len(header) + 1):
        cell = ws.cell(row=hrow, column=c)
        cell.font = head_font; cell.fill = head_fill; cell.alignment = centre
    for i, r in enumerate(rows, 1):
        grade = GradeScale.get_grade(r['average']) if r['average'] else '-'
        line = [i, r['student'], r['class_name'], r['arm_name']]
        line += [r['subjects'].get(str(s['id']), '') for s in subjects]
        line += [r['total'], r['average'], grade, r['passed'], r['failed']]
        ws.append(line)
    ws.freeze_panes = ws.cell(row=hrow + 1, column=3)

    cmp = wb.create_sheet('Comparison')
    cmp.append(['Class arm', 'Students', 'Mean average', 'Pass rate %'])
    for c in range(1, 5):
        cell = cmp.cell(row=1, column=c)
        cell.font = head_font; cell.fill = head_fill; cell.alignment = centre
    for st in _explore_scope_stats(rows, scope_meta, pass_mark=pass_mark):
        cmp.append([st['label'], st['n'],
                    '' if st['mean_avg'] is None else st['mean_avg'],
                    '' if st['pass_rate'] is None else st['pass_rate']])
    # crude column sizing
    for sheet in (ws, cmp):
        for col in sheet.columns:
            width = max((len(str(c.value)) for c in col if c.value is not None), default=8)
            sheet.column_dimensions[col[0].column_letter].width = min(max(width + 2, 8), 32)
    return wb


def explore_pdf(subjects, scope_meta, rows, term_name, pass_mark=50, filter_label=''):
    """Filtered cross-class Explorer view as one print-ready PDF (landscape A4):
    a scope-comparison table followed by the unified student table."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer)
    from utils.web_exports import pdf_escape

    primary, accent, light, ink = _theme()
    styles = getSampleStyleSheet()
    h = ParagraphStyle('h', parent=styles['Title'], fontSize=15, textColor=primary, spaceAfter=2)
    sub = ParagraphStyle('sub', parent=styles['Normal'], fontSize=9.5, textColor=colors.HexColor('#6B7A74'))
    cell = ParagraphStyle('c', parent=styles['Normal'], fontSize=7.5, leading=9)
    h3 = ParagraphStyle('h3', parent=styles['Heading3'], fontSize=10.5, textColor=ink, spaceBefore=6, spaceAfter=3)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), topMargin=10 * mm, bottomMargin=10 * mm,
                            leftMargin=8 * mm, rightMargin=8 * mm, title='Results Explorer')
    avail = landscape(A4)[0] - 16 * mm

    scope_names = ', '.join(m['label'] for m in scope_meta) or '—'
    meta_bits = [pdf_escape(scope_names)]
    if term_name:
        meta_bits.append(pdf_escape(term_name))
    if filter_label:
        meta_bits.append('Filter: ' + pdf_escape(filter_label))
    elems = [Paragraph(pdf_escape(_school_name()), h),
             Paragraph('Results Explorer · ' + ' · '.join(meta_bits), sub), Spacer(1, 4)]

    # --- Comparison table ---
    stats = _explore_scope_stats(rows, scope_meta, pass_mark=pass_mark)
    cdata = [['Class arm', 'Students', 'Mean avg', 'Pass rate']]
    for st in stats:
        cdata.append([Paragraph(pdf_escape(st['label']), cell), str(st['n']),
                      '–' if st['mean_avg'] is None else _n(st['mean_avg']),
                      '–' if st['pass_rate'] is None else f"{st['pass_rate']}%"])
    if len(stats) > 1:
        n = len(rows)
        mean_all = round(sum(r['average'] for r in rows) / n, 2) if n else None
        pass_all = sum(1 for r in rows if r['average'] >= pass_mark)
        cdata.append([Paragraph('<b>All selected</b>', cell), str(n),
                      '–' if mean_all is None else _n(mean_all),
                      '–' if not n else f'{round(pass_all / n * 100)}%'])
    ct = Table(cdata, colWidths=[avail * 0.4, avail * 0.2, avail * 0.2, avail * 0.2], repeatRows=1)
    ct.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), primary), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (1, 0), (-1, -1), 'CENTER'), ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#D5DED9')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light]),
        ('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 3)]))
    elems += [Paragraph('Compare scopes', h3), ct, Spacer(1, 6)]

    # --- Student table ---
    head = ['#', 'Student', 'Class', 'Arm'] + [(s['short'] or s['name'][:6]) for s in subjects] \
        + ['Total', 'Avg', 'Grade']
    data = [head]
    for i, r in enumerate(rows, 1):
        line = [str(i), Paragraph(pdf_escape(r['student']), cell),
                Paragraph(pdf_escape(r['class_name']), cell), Paragraph(pdf_escape(r['arm_name']), cell)]
        for s in subjects:
            v = r['subjects'].get(str(s['id']))
            line.append(_n(v) if v else '–')
        line += [_n(r['total']), _n(r['average']),
                 GradeScale.get_grade(r['average']) if r['average'] else '–']
        data.append(line)

    pos_w, name_w, cls_w, arm_w, tail_w = 8 * mm, 40 * mm, 20 * mm, 16 * mm, 12 * mm
    fixed = pos_w + name_w + cls_w + arm_w + 3 * tail_w
    subj_w = max(8 * mm, (avail - fixed) / max(len(subjects), 1))
    widths = [pos_w, name_w, cls_w, arm_w] + [subj_w] * len(subjects) + [tail_w, tail_w, tail_w]
    st = Table(data, colWidths=widths, repeatRows=1)
    st.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), primary), ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, -1), 7.5),
        ('ALIGN', (4, 0), (-1, -1), 'CENTER'), ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('FONTNAME', (-3, 1), (-1, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#D5DED9')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light]),
        ('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 3)]))
    elems += [Paragraph(f'Students ({len(rows)})', h3), st]
    doc.build(elems)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Generic, column-selectable table export (Combination Explorer)
# Used by the subject-combination results tool: the caller decides the exact
# columns (student, per-subject scores, combined total/average, …) and rows,
# and picks the format. One code path → PDF, PNG and Excel stay identical.
# --------------------------------------------------------------------------- #

# Short codes for subjects so wide broadsheets fit A4 at a legible font size.
_SUBJECT_ABBR = {
    'Mathematics': 'MTH', 'English Language': 'ENG', 'Civic Education': 'CIV',
    'Biology': 'BIO', 'Chemistry': 'CHE', 'Physics': 'PHY', 'Economics': 'ECO',
    'Government': 'GOV', 'Literature in English': 'LIT', 'Commerce': 'COM',
    'Christian Religious Studies': 'CRS', 'Islamic Religious Studies': 'IRS',
    'Financial Accounting': 'ACC', 'Accounting': 'ACC', 'Geography': 'GEO',
    'Agricultural Science': 'AGR', 'Further Mathematics': 'FMT', 'History': 'HIS',
    'Data Processing': 'DPR', 'Computer Studies': 'CMP', 'Technical Drawing': 'TDR',
    'Food and Nutrition': 'FDN', 'Home Management': 'HMG', 'Marketing': 'MKT',
    'Yoruba': 'YOR', 'Hausa': 'HAU', 'Igbo': 'IGB', 'French': 'FRN', 'Fine Art': 'ART',
    'Music': 'MUS', 'Livestock Farming': 'LVF', 'Digital Technologies': 'DGT',
    'Health Education': 'HED', 'Physical Education': 'PHE', 'Book Keeping': 'BKP',
    'Catering Craft Practice': 'CCP', 'Animal Husbandry': 'ANH', 'Insurance': 'INS',
}


def _abbr_one(name):
    """A short code for a subject name: a known code, else an acronym of the
    significant words, else the first letters — always 2–4 upper-case chars."""
    name = (name or '').strip()
    if not name:
        return '?'
    if name in _SUBJECT_ABBR:
        return _SUBJECT_ABBR[name]
    stop = {'in', 'of', 'and', 'the', '&'}
    words = [w for w in name.replace('/', ' ').split() if w.lower() not in stop]
    if len(words) >= 2:
        code = ''.join(w[0] for w in words[:4]).upper()
    else:
        code = name[:3].upper()
    return code


def abbreviate_subjects(names):
    """Map a list of subject names to unique short codes, returning
    ``(codes, legend)`` where legend is a list of ``(code, full_name)`` for a key
    printed beneath the table. Collisions get a numeric suffix so codes stay
    unique."""
    codes, legend, seen = [], [], {}
    for n in names:
        base = _abbr_one(n)
        code = base
        k = 2
        while code in seen and seen[code] != n:
            code = f'{base}{k}'
            k += 1
        seen[code] = n
        codes.append(code)
        legend.append((code, n))
    return codes, legend


def combo_pdf(headers, data_rows, title, subtitle='', numeric_from=1, legend=None,
              logo_path=None, school_name=None, sections=None, body_fs=None):
    """Print-ready PDF of an arbitrary table (landscape A4). ``headers`` is a
    list of column titles; ``data_rows`` a list of equal-length string rows;
    ``numeric_from`` is the first column index to centre (names stay left).

    ``sections`` (optional) is a list of ``(section_title, rows)`` — each renders
    on its own fresh page(s) with the masthead + a "… — TITLE" heading, so
    groups land on separate sheets. The masthead shows on the first page of each
    section only. Column widths adapt to content and fill the A4 width.
    Colours are light/toner-friendly for printing."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph,
                                    Spacer, Image, PageBreak, KeepTogether)
    from reportlab.lib.enums import TA_LEFT
    from utils.web_exports import pdf_escape
    from reportlab.pdfbase.pdfmetrics import stringWidth

    if sections is None:
        sections = [(None, data_rows)]
    all_rows = [r for _t, rows in sections for r in rows]

    primary, accent, light, ink = _neutral()
    # Print-friendly header: a light fill with dark bold text + a rule, instead
    # of a heavy dark band (saves toner and reads cleanly in black & white).
    head_bg = colors.HexColor('#E9EDF2')
    head_fg = colors.HexColor('#0F172A')
    rule = colors.HexColor('#94A3B8')
    grid = colors.HexColor('#D5DBE3')
    zebra = colors.HexColor('#F5F7FA')
    if school_name is None:
        school_name = _school_name()
    styles = getSampleStyleSheet()
    ncol = len(headers)
    if body_fs:
        fs = body_fs
    elif ncol <= 8:
        fs = 13
    elif ncol <= 10:
        fs = 12
    elif ncol <= 14:
        fs = 11
    elif ncol <= 18:
        fs = 9.5
    else:
        fs = 8.5
    schoolst = ParagraphStyle('sn', parent=styles['Normal'], fontSize=11, textColor=colors.HexColor('#64748B'),
                              fontName='Helvetica-Bold', alignment=TA_LEFT, spaceAfter=1)
    h = ParagraphStyle('h', parent=styles['Title'], fontSize=17, textColor=primary, spaceAfter=2, alignment=TA_LEFT)
    sub = ParagraphStyle('sub', parent=styles['Normal'], fontSize=9.5, textColor=colors.HexColor('#6B7A74'))
    cell = ParagraphStyle('c', parent=styles['Normal'], fontSize=fs, leading=fs + 1.5)
    cellc = ParagraphStyle('cc', parent=cell, alignment=1)
    headp = ParagraphStyle('hp', parent=styles['Normal'], fontSize=fs, leading=fs + 1.5,
                           textColor=head_fg, fontName='Helvetica-Bold', alignment=1)
    keyst = ParagraphStyle('key', parent=styles['Normal'], fontSize=8.5, leading=11,
                           textColor=colors.HexColor('#44524C'))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), topMargin=10 * mm, bottomMargin=12 * mm,
                            leftMargin=8 * mm, rightMargin=8 * mm, title=title or 'Results')
    avail = landscape(A4)[0] - 16 * mm

    # Column widths: each column tight to its content; the leftover width is
    # given to the flexible name column so the table fills A4 without padding
    # out the numeric columns (bootstrap-grid style). Shrinks to fit if wide.
    nat = []
    for j in range(ncol):
        w = stringWidth(str(headers[j]), 'Helvetica-Bold', fs)
        for r in all_rows:
            if j < len(r):
                w = max(w, stringWidth(str(r[j]), 'Helvetica', fs))
        nat.append(w + 6 * mm)
    tot = sum(nat) or 1
    flex = 1 if ncol > 1 else 0            # the Student column absorbs slack
    if tot < avail:
        widths = list(nat); widths[flex] += (avail - tot)
    else:
        widths = [w * (avail / tot) for w in nat]

    def logo_flowable():
        if not logo_path:
            return None
        try:
            from PIL import Image as _PILImage
            iw, ih = _PILImage.open(logo_path).size
            lw = 20 * mm
            lh = lw * (ih / iw) if iw else 20 * mm
            return Image(logo_path, width=lw, height=min(lh, 22 * mm))
        except Exception:
            return None

    def masthead(sec_title):
        heading_text = (title or 'Results') + (' — ' + sec_title if sec_title else '')
        cells = []
        if school_name:
            cells.append(Paragraph(pdf_escape(school_name), schoolst))
        cells.append(Paragraph(pdf_escape(heading_text), h))
        if subtitle:
            cells.append(Paragraph(pdf_escape(subtitle), sub))
        lg = logo_flowable()
        if lg is not None:
            mast = Table([[lg, cells]], colWidths=[24 * mm, avail - 24 * mm])
            mast.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                      ('LEFTPADDING', (0, 0), (-1, -1), 0),
                                      ('RIGHTPADDING', (0, 0), (-1, -1), 0)]))
            return [mast, Spacer(1, 6)]
        return list(cells) + [Spacer(1, 5)]

    def make_table(rows):
        data = [[Paragraph(pdf_escape(str(x)), headp) for x in headers]]
        for r in rows:
            line = []
            for j, v in enumerate(r):
                if j == 0 or (numeric_from is not None and j < numeric_from):
                    line.append(Paragraph(pdf_escape(str(v)), cell))
                else:
                    line.append(Paragraph(pdf_escape(str(v)), cellc))
            data.append(line)
        t = Table(data, colWidths=widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), head_bg),
            ('LINEBELOW', (0, 0), (-1, 0), 1.1, rule),
            ('FONTSIZE', (0, 0), (-1, -1), fs),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LINEBELOW', (0, 1), (-1, -1), 0.4, grid),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, zebra]),
            ('BOX', (0, 0), (-1, -1), 0.5, grid),
            ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 5), ('RIGHTPADDING', (0, 0), (-1, -1), 5)]))
        return t

    elems = []
    for si, (sec_title, rows) in enumerate(sections):
        if si:
            elems.append(PageBreak())
        # Keep the masthead with the header row so a section never starts with a
        # lone masthead at the foot of a page.
        elems.append(KeepTogether(masthead(sec_title)))
        elems.append(make_table(rows))
        if legend and si == len(sections) - 1:
            parts = ' &nbsp;·&nbsp; '.join(f'<b>{pdf_escape(code)}</b> = {pdf_escape(full)}'
                                           for code, full in legend)
            elems.append(Spacer(1, 6))
            elems.append(Paragraph('Key: ' + parts, keyst))
    doc.build(elems)
    return buf.getvalue()


def combo_png_pages(headers, data_rows, title, subtitle='', legend=None,
                    logo_path=None, school_name=None, sections=None):
    """Render an arbitrary table as one or more **HD landscape-A4** PNG pages
    (print-friendly light styling). Column widths adapt to their content and
    fill the page width. ``sections`` (optional) is a list of
    ``(section_title, rows)`` — each group starts on a fresh page with the
    masthead + a "… — TITLE" heading, and the masthead shows only on the first
    page of each group. Returns a list of PNG byte strings."""
    from PIL import Image, ImageDraw, ImageFont
    S = 2                                   # supersample, downscaled on save
    DPI = 200                               # HD
    PW = int(round(297 / 25.4 * DPI))       # landscape-A4 width  px (~2339)
    PH = int(round(210 / 25.4 * DPI))       # landscape-A4 height px (~1654)
    C = _NEUTRAL_RGB
    ncol = len(headers)
    if sections is None:
        sections = [(None, data_rows)]
    all_rows = [r for _t, rows in sections for r in rows]
    if school_name is None:
        school_name = _school_name()
    logo_im = None
    if logo_path:
        try:
            logo_im = Image.open(logo_path).convert('RGBA')
        except Exception:
            logo_im = None

    def fnt(size, bold=False):
        path = "/usr/share/fonts/truetype/dejavu/DejaVuSans%s.ttf" % ("-Bold" if bold else "")
        try:
            return ImageFont.truetype(path, int(size * S))
        except Exception:
            return ImageFont.load_default()

    if ncol <= 8:
        fs = 22
    elif ncol <= 10:
        fs = 20
    elif ncol <= 14:
        fs = 17
    elif ncol <= 18:
        fs = 15
    else:
        fs = 13
    body, body_b = fnt(fs), fnt(fs, True)
    title_f, sub_f, key_f = fnt(26, True), fnt(14), fnt(13)

    tmp = ImageDraw.Draw(Image.new('RGB', (1, 1)))

    def tw(text, f):
        b = tmp.textbbox((0, 0), str(text), font=f); return b[2] - b[0]

    def fit(text, f, maxw):
        text = str(text)
        if tw(text, f) <= maxw:
            return text
        while text and tw(text + '…', f) > maxw:
            text = text[:-1]
        return (text + '…') if text else ''

    margin = 34 * S
    avail = PW * S - 2 * margin
    cpx, cpy = int(9 * S), int(7 * S)

    # Each column tight to its content; the flexible name column absorbs the
    # leftover so the table fills the page width without padding out numbers.
    nat = []
    for j in range(ncol):
        w = tw(headers[j], body_b)
        for r in all_rows:
            w = max(w, tw(r[j] if j < len(r) else '', body))
        nat.append(w + 2 * cpx)
    tot = sum(nat) or 1
    flex = 1 if ncol > 1 else 0
    if tot < avail:
        col_w = list(nat); col_w[flex] += (avail - tot)
    else:
        col_w = [int(w * (avail / tot)) for w in nat]
    col_w = [int(w) for w in col_w]
    table_w = sum(col_w)

    line_h = tmp.textbbox((0, 0), "Ay", font=body)[3]
    row_h = line_h + 2 * cpy
    header_h = row_h + int(4 * S)

    logo_box = int(64 * S)
    logo_draw = None
    if logo_im is not None:
        lw, lh = logo_im.size
        scale = logo_box / max(lw, lh)
        logo_draw = logo_im.resize((max(1, int(lw * scale)), max(1, int(lh * scale))), Image.LANCZOS)
    text_x = margin + (logo_box + int(14 * S) if logo_draw is not None else 0)
    mast_lines = int(34 * S) + (int(20 * S) if school_name else 0) + (int(22 * S) if subtitle else 0)
    mast_h = max(mast_lines, logo_box if logo_draw is not None else 0)

    top_gap = int(12 * S)                          # top margin on masthead-free pages
    bottom_reserve = margin + int(24 * S)
    per_first = max(1, int((PH * S - (margin + mast_h) - header_h - bottom_reserve) // row_h))
    per_rest = max(1, int((PH * S - (margin + top_gap) - header_h - bottom_reserve) // row_h))

    # Build the page plan: each section starts fresh; masthead on its 1st page.
    plan = []      # (sec_title, chunk, draw_mast)
    for sec_title, rows in sections:
        if not rows:
            plan.append((sec_title, [], True)); continue
        i, first = 0, True
        while i < len(rows):
            cap = per_first if first else per_rest
            plan.append((sec_title, rows[i:i + cap], first))
            i += cap; first = False
    n_pages = len(plan)

    def draw_masthead(d, img, sec_title):
        if logo_draw is not None:
            img.paste(logo_draw, (margin, margin), logo_draw)
        ty = margin
        text_avail = PW * S - margin - text_x
        if school_name:
            d.text((text_x, ty), fit(school_name, sub_f, text_avail), fill=C['muted'], font=sub_f)
            ty += int(20 * S)
        heading = (title or 'Results') + (' — ' + sec_title if sec_title else '')
        d.text((text_x, ty), fit(heading, title_f, text_avail), fill=C['header'], font=title_f)
        ty += int(34 * S)
        if subtitle:
            d.text((text_x, ty), fit(subtitle, sub_f, text_avail), fill=C['muted'], font=sub_f)

    pages = []
    for pi, (sec_title, chunk, draw_mast) in enumerate(plan):
        img = Image.new('RGB', (PW * S, PH * S), C['white'])
        d = ImageDraw.Draw(img)
        if draw_mast:
            draw_masthead(d, img, sec_title)
            y0 = margin + mast_h
        else:
            y0 = margin + top_gap
        # header band — light fill, dark text, accent rule beneath
        d.rectangle([margin, y0, margin + table_w, y0 + header_h], fill=C['head_bg'])
        x = margin
        for j in range(ncol):
            d.text((x + cpx, y0 + (header_h - line_h) // 2), fit(headers[j], body_b, col_w[j] - 2 * cpx),
                   fill=C['head_fg'], font=body_b)
            x += col_w[j]
        d.rectangle([margin, y0 + header_h - max(2, S), margin + table_w, y0 + header_h], fill=C['rule'])
        y = y0 + header_h
        for i, r in enumerate(chunk):
            if i % 2:
                d.rectangle([margin, y, margin + table_w, y + row_h], fill=C['zebra'])
            x = margin
            for j in range(ncol):
                val = str(r[j]) if j < len(r) else ''
                align_left = (j <= 1)
                txt = fit(val, body, col_w[j] - 2 * cpx)
                tx = x + cpx if align_left else x + (col_w[j] - tw(txt, body)) / 2
                d.text((tx, y + cpy), txt, fill=C['text'], font=body)
                x += col_w[j]
            d.line([margin, y, margin + table_w, y], fill=C['line'], width=1)
            y += row_h
        d.rectangle([margin, y0, margin + table_w, y], outline=C['line'], width=1)
        x = margin
        for j in range(ncol - 1):
            x += col_w[j]
            d.line([x, y0, x, y], fill=C['line'], width=1)

        d.text((PW * S - margin - int(200 * S), PH * S - margin), 'Page %d of %d' % (pi + 1, n_pages),
               fill=C['muted'], font=key_f)
        if legend and pi == n_pages - 1:
            ky = y + int(12 * S)
            line = 'Key:  '
            for code, full in legend:
                line += '%s = %s     ' % (code, full)
            if ky + key_f.size < PH * S - margin:
                d.text((margin, ky), fit(line.rstrip(), key_f, table_w), fill=C['muted'], font=key_f)
        out = io.BytesIO()
        img.resize((PW, PH), Image.LANCZOS).save(out, format='PNG')
        pages.append(out.getvalue())
    return pages


def combo_png(headers, data_rows, title, subtitle='', legend=None):
    """Backward-compatible single-image export — the first A4 page."""
    return combo_png_pages(headers, data_rows, title, subtitle, legend=legend)[0]


def zip_pngs(pages, base='page'):
    """Bundle a list of PNG byte strings into a zip archive (bytes)."""
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for i, data in enumerate(pages, 1):
            zf.writestr('%s_page_%d.png' % (base, i), data)
    return buf.getvalue()


def combo_xlsx(headers, data_rows, title, subtitle='', sections=None):
    """Workbook of the combination table. When ``sections`` (list of
    ``(section_title, rows)``) is given, each group gets its own worksheet
    (Group A, Group B, …); otherwise a single sheet is used."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    if sections is None:
        sections = [(None, data_rows)]
    wb = Workbook()
    hf = Font(bold=True, color='0F172A')
    fill = PatternFill('solid', fgColor='E9EDF2')      # print-friendly light header
    ctr = Alignment(horizontal='center')

    def _sheet(ws, sec_title, rows):
        if subtitle:
            ws.append([subtitle]); ws.append([])
        ws.append(list(headers))
        hrow = ws.max_row
        for c in range(1, len(headers) + 1):
            cell = ws.cell(row=hrow, column=c); cell.font = hf; cell.fill = fill; cell.alignment = ctr
        for r in rows:
            ws.append(list(r))
        ws.freeze_panes = ws.cell(row=hrow + 1, column=1)
        for col in ws.columns:
            width = max((len(str(c.value)) for c in col if c.value is not None), default=8)
            ws.column_dimensions[col[0].column_letter].width = min(max(width + 2, 8), 34)

    first = True
    for si, (sec_title, rows) in enumerate(sections):
        ws = wb.active if first else wb.create_sheet()
        ws.title = (sec_title or 'Combination')[:31]
        _sheet(ws, sec_title, rows)
        first = False
    return wb


def combo_docx(headers, data_rows, title, subtitle='', legend=None,
               logo_path=None, school_name=None, sections=None, numeric_from=1):
    """Word (.docx) of the combination table. Each ``sections`` group starts on
    a new page with the masthead + a "… — TITLE" heading; the masthead shows
    once per group. Returns the raw .docx bytes."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.enum.section import WD_ORIENT
    import os as _os
    import io as _io

    if sections is None:
        sections = [(None, data_rows)]
    if school_name is None:
        school_name = _school_name()
    navy = RGBColor(0x33, 0x41, 0x55); muted = RGBColor(0x64, 0x74, 0x8B)

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Cm(1.0); sec.top_margin = sec.bottom_margin = Cm(1.0)

    def masthead(sec_title):
        if logo_path and _os.path.exists(logo_path):
            try:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run().add_picture(logo_path, height=Cm(1.4))
            except Exception:
                pass
        if school_name:
            sp = doc.add_paragraph(); r = sp.add_run(school_name)
            r.bold = True; r.font.size = Pt(11); r.font.color.rgb = muted
        hp = doc.add_paragraph()
        hr = hp.add_run((title or 'Results') + (' — ' + sec_title if sec_title else ''))
        hr.bold = True; hr.font.size = Pt(16); hr.font.color.rgb = navy
        if subtitle:
            sb = doc.add_paragraph(); sr = sb.add_run(subtitle)
            sr.font.size = Pt(9); sr.font.color.rgb = muted

    def build_table(rows):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, hlabel in enumerate(headers):
            cell = t.rows[0].cells[i]; cell.text = str(hlabel)
            pr = cell.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if pr.runs:
                pr.runs[0].bold = True; pr.runs[0].font.size = Pt(9)
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E9EDF2"/>'))
        for r in rows:
            cells = t.add_row().cells
            for i in range(len(headers)):
                v = '' if i >= len(r) else str(r[i])
                cells[i].text = v
                pr = cells[i].paragraphs[0]
                pr.alignment = (WD_ALIGN_PARAGRAPH.LEFT if (i == 0 or i < numeric_from)
                                else WD_ALIGN_PARAGRAPH.CENTER)
                if pr.runs:
                    pr.runs[0].font.size = Pt(9)

    for si, (sec_title, rows) in enumerate(sections):
        if si:
            doc.add_page_break()
        masthead(sec_title)
        build_table(rows)
    if legend:
        lp = doc.add_paragraph()
        lr = lp.add_run('Key:  ' + '   '.join('%s = %s' % (c, f) for c, f in legend))
        lr.font.size = Pt(8); lr.font.color.rgb = muted

    out = _io.BytesIO(); doc.save(out); return out.getvalue()
