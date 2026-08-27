"""Branded single-applicant record sheet — PDF (reportlab) and Word (python-docx).

Shares the visual language of the other exports: a masthead with the school
logo / name / address / motto, a title band, and labelled sections
(Applicant, Application, Parent/Guardian, Emergency Contact, Notes).

``record`` shape (built by the route)::

    {
      'title': 'Applicant Record',
      'name': 'Surname First Middle',
      'meta': [('Application No', 'APP2026-0001'), ('Status', 'Applied')],
      'sections': [('Applicant', [('Gender', 'Male'), ...]), ...],
    }
"""
import io
import os

from utils import timeutil
from utils.student_export import (NAVY, GOLD, INK, MUTED, LINE, PANEL,
                                  _FONT_REG, _FONT_BOLD, _FONT_ITAL)


def applicant_pdf(record, school):
    """A4-portrait applicant record sheet built with reportlab. Returns bytes."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (BaseDocTemplate, Frame, PageTemplate, Table,
                                    TableStyle, Paragraph, Spacer, KeepTogether)
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    base, boldf, obl = 'Helvetica', 'Helvetica-Bold', 'Helvetica-Oblique'
    try:
        pdfmetrics.registerFont(TTFont('ApBody', _FONT_REG)); base = 'ApBody'
        pdfmetrics.registerFont(TTFont('ApBody-Bold', _FONT_BOLD)); boldf = 'ApBody-Bold'
        pdfmetrics.registerFont(TTFont('ApBody-Obl', _FONT_ITAL)); obl = 'ApBody-Obl'
    except Exception:
        pass

    PW, PH = A4
    margin = 14 * mm
    mast_h = 40 * mm
    avail = PW - 2 * margin

    lab = ParagraphStyle('lab', fontName=boldf, fontSize=10, leading=13,
                         textColor=colors.HexColor(MUTED))
    val = ParagraphStyle('val', fontName=base, fontSize=11.5, leading=15,
                         textColor=colors.HexColor(INK))
    sec = ParagraphStyle('sec', fontName=boldf, fontSize=12.5, leading=16,
                         textColor=colors.HexColor(NAVY))

    def _draw_masthead(cv, doc):
        from reportlab.pdfbase.pdfmetrics import stringWidth
        cv.setStrokeColor(colors.HexColor(GOLD)); cv.setLineWidth(1.6)
        cv.roundRect(margin * 0.5, margin * 0.5, PW - margin, PH - margin, 9, stroke=1, fill=0)
        top = PH - margin - 3 * mm
        x = margin + 4 * mm
        logo = (school or {}).get('logo_path')
        lx = x
        if logo and os.path.exists(logo):
            try:
                from reportlab.lib.utils import ImageReader
                ir = ImageReader(logo); iw, ih = ir.getSize()
                h = 22 * mm; w = h * (iw / ih) if ih else 22 * mm
                cv.drawImage(ir, x, top - h, width=min(w, 24 * mm), height=h,
                             mask='auto', preserveAspectRatio=True)
                lx = x + min(w, 24 * mm) + 5 * mm
            except Exception:
                lx = x
        name = ((school or {}).get('name') or 'School').upper()
        size = 21
        while size > 12 and stringWidth(name, boldf, size) > (avail - (lx - x)):
            size -= 1
        cv.setFillColor(colors.HexColor(NAVY)); cv.setFont(boldf, size)
        cv.drawString(lx, top - 7 * mm, name)
        cv.setFillColor(colors.HexColor(INK)); cv.setFont(base, 9.5)
        addr = (school or {}).get('address') or ''
        if addr:
            cv.drawString(lx, top - 12.5 * mm, addr[:110])
        contact = '    '.join(p for p in [(school or {}).get('phone') or '',
                                          (school or {}).get('email') or ''] if p)
        if contact:
            cv.drawString(lx, top - 17 * mm, contact[:110])
        motto = (school or {}).get('motto') or ''
        if motto:
            cv.setFillColor(colors.HexColor(GOLD)); cv.setFont(obl, 10)
            cv.drawString(lx, top - 22 * mm, ('—  ' + motto + '  —')[:110])
        # passport photo (top-right corner of the first page)
        pb = record.get('photo_bytes')
        if pb and doc.page == 1:
            try:
                from reportlab.lib.utils import ImageReader
                pw_, ph_ = 24 * mm, 30 * mm
                px = PW - margin - 4 * mm - pw_
                cv.drawImage(ImageReader(io.BytesIO(pb)), px, top - ph_, width=pw_, height=ph_,
                             preserveAspectRatio=True, mask='auto')
                cv.setStrokeColor(colors.HexColor(LINE)); cv.setLineWidth(0.6)
                cv.rect(px, top - ph_, pw_, ph_, stroke=1, fill=0)
            except Exception:
                pass
        # gold divider under the masthead
        cv.setStrokeColor(colors.HexColor(GOLD)); cv.setLineWidth(1)
        cv.line(margin, PH - margin - mast_h + 4 * mm, PW - margin, PH - margin - mast_h + 4 * mm)
        # footer
        cv.setStrokeColor(colors.HexColor(LINE)); cv.setLineWidth(0.6)
        cv.line(margin, margin + 8 * mm, PW - margin, margin + 8 * mm)
        cv.setFillColor(colors.HexColor(MUTED)); cv.setFont(obl, 8)
        cv.drawString(margin, margin + 4 * mm, 'Generated %s · system-generated and confidential.'
                      % timeutil.today().strftime('%d %b %Y'))
        cv.drawRightString(PW - margin, margin + 4 * mm, 'Page %d' % doc.page)

    story = []
    # Title band + name + meta
    title = record.get('title') or 'Applicant Record'
    story.append(Paragraph('<font color="%s"><b>%s</b></font>' % (NAVY, title),
                           ParagraphStyle('t', fontName=boldf, fontSize=16, leading=20)))
    story.append(Paragraph('<font color="%s">%s</font>' % (INK, record.get('name') or ''),
                           ParagraphStyle('n', fontName=base, fontSize=13, leading=17)))
    meta = record.get('meta') or []
    if meta:
        cells = [[Paragraph('<font color="%s"><b>%s</b></font>&nbsp;&nbsp;%s'
                            % (MUTED, k, v), val) for k, v in meta]]
        mt = Table(cells, colWidths=[avail / max(1, len(meta))] * len(meta))
        mt.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(PANEL)),
            ('BOX', (0, 0), (-1, -1), 0.6, colors.HexColor(LINE)),
            ('INNERGRID', (0, 0), (-1, -1), 0.6, colors.HexColor(LINE)),
            ('TOPPADDING', (0, 0), (-1, -1), 7), ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
            ('LEFTPADDING', (0, 0), (-1, -1), 9), ('RIGHTPADDING', (0, 0), (-1, -1), 9),
        ]))
        story.append(Spacer(1, 5 * mm)); story.append(mt)

    for heading, rows in record.get('sections') or []:
        rows = [r for r in rows if r]
        if not rows:
            continue
        block = [Spacer(1, 6 * mm), Paragraph(heading, sec), Spacer(1, 2 * mm)]
        data = [[Paragraph(str(k), lab), Paragraph('' if v is None else str(v), val)] for k, v in rows]
        t = Table(data, colWidths=[46 * mm, avail - 46 * mm])
        t.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 2), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('LINEBELOW', (0, 0), (-1, -2), 0.4, colors.HexColor(LINE)),
        ]))
        block.append(t)
        story.append(KeepTogether(block))

    buf = io.BytesIO()
    frame = Frame(margin, margin + 10 * mm, avail, PH - mast_h - margin - 12 * mm,
                  leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    doc = BaseDocTemplate(buf, pagesize=A4, leftMargin=margin, rightMargin=margin,
                          topMargin=margin, bottomMargin=margin)
    doc.addPageTemplates([PageTemplate(id='main', frames=[frame], onPage=_draw_masthead)])
    doc.build(story)
    return buf.getvalue()


def applicant_docx(record, school, filename='applicant.docx'):
    """A branded applicant record sheet as .docx. Returns a Flask Response."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from io import BytesIO
    from flask import Response

    navy = RGBColor(0x1E, 0x2A, 0x4A); gold = RGBColor(0xB8, 0x86, 0x2B)
    muted = RGBColor(0x6B, 0x72, 0x80)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Cm(1.6)
    section.top_margin = section.bottom_margin = Cm(1.3)

    logo = (school or {}).get('logo_path')
    if logo and os.path.exists(logo):
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(logo, height=Cm(1.5))
        except Exception:
            pass
    nm = doc.add_paragraph(); nm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = nm.add_run(((school or {}).get('name') or 'School').upper())
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = navy
    line2 = '    '.join(x for x in [(school or {}).get('address') or '',
                                    (school or {}).get('phone') or '',
                                    (school or {}).get('email') or ''] if x)
    if line2:
        sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.add_run(line2).font.size = Pt(9.5)
    if (school or {}).get('motto'):
        mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rm = mp.add_run('—  ' + school['motto'] + '  —'); rm.italic = True
        rm.font.size = Pt(10); rm.font.color.rgb = gold

    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(record.get('title') or 'Applicant Record')
    tr.bold = True; tr.font.size = Pt(15); tr.font.color.rgb = navy
    if record.get('name'):
        npar = doc.add_paragraph(); npar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        npar.add_run(record['name']).font.size = Pt(12)
    meta = record.get('meta') or []
    if meta:
        mp2 = doc.add_paragraph(); mp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = mp2.add_run('    '.join('%s: %s' % (k, v) for k, v in meta))
        mr.font.size = Pt(10); mr.font.color.rgb = muted

    for heading, rows in record.get('sections') or []:
        rows = [r for r in rows if r]
        if not rows:
            continue
        hp = doc.add_paragraph(); hr = hp.add_run(heading)
        hr.bold = True; hr.font.size = Pt(12); hr.font.color.rgb = navy
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for k, v in rows:
            row = table.add_row().cells
            row[0].width = Cm(5.5); row[1].width = Cm(12.5)
            kp = row[0].paragraphs[0]; kr = kp.add_run(str(k))
            kr.bold = True; kr.font.size = Pt(9.5); kr.font.color.rgb = muted
            kp2 = row[1].paragraphs[0]; kr2 = kp2.add_run('' if v is None else str(v))
            kr2.font.size = Pt(10.5)
            row[0]._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7F8FA"/>'))
        doc.add_paragraph()

    foot = doc.add_paragraph('Generated %s · system-generated and confidential.'
                             % timeutil.today().strftime('%d %b %Y'))
    if foot.runs:
        foot.runs[0].italic = True; foot.runs[0].font.size = Pt(8)
        foot.runs[0].font.color.rgb = muted

    out = BytesIO(); doc.save(out); out.seek(0)
    return Response(out.getvalue(),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': 'attachment; filename=%s' % filename})


def applicant_blank_pdf(school, bw=False):
    """A branded, *fillable* blank application form (interactive AcroForm PDF).

    Every field is a real PDF form field the applicant can type into (or print
    and fill by hand). Built with reportlab. Returns bytes.

    ``bw=True`` produces a black-and-white print-friendly variant: black text,
    strong dark field borders, white fills and no colour accents, so it stays
    crisp and legible on a monochrome printer."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas as _canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import stringWidth

    base, boldf, obl = 'Helvetica', 'Helvetica-Bold', 'Helvetica-Oblique'
    try:
        pdfmetrics.registerFont(TTFont('ApBody', _FONT_REG)); base = 'ApBody'
        pdfmetrics.registerFont(TTFont('ApBody-Bold', _FONT_BOLD)); boldf = 'ApBody-Bold'
        pdfmetrics.registerFont(TTFont('ApBody-Obl', _FONT_ITAL)); obl = 'ApBody-Obl'
    except Exception:
        pass

    PW, PH = A4
    margin = 14 * mm
    avail = PW - 2 * margin
    if bw:
        # Monochrome, high-contrast: black ink, strong grey borders, white
        # field fills, and no colour accents — prints cleanly on a mono printer.
        navy = colors.black
        gold = colors.HexColor('#333333')
        ink = colors.black
        muted = colors.HexColor('#333333')
        line = colors.HexColor('#555555')
        panel = colors.white
        field_border = colors.HexColor('#333333')
        field_bw = 1.0
    else:
        navy = colors.HexColor(NAVY); gold = colors.HexColor(GOLD)
        ink = colors.HexColor(INK); muted = colors.HexColor(MUTED)
        line = colors.HexColor(LINE); panel = colors.HexColor(PANEL)
        field_border = line; field_bw = 0.7

    buf = io.BytesIO()
    c = _canvas.Canvas(buf, pagesize=A4)
    c.setTitle('Application Form')
    form = c.acroForm
    state = {'y': 0, 'n': 0, 'page': 0}

    def new_page(first=False):
        state['page'] += 1
        c.setStrokeColor(gold); c.setLineWidth(1.6)
        c.roundRect(margin * 0.5, margin * 0.5, PW - margin, PH - margin, 9, stroke=1, fill=0)
        top = PH - margin - 3 * mm
        if first:
            x = margin + 4 * mm
            lx = x
            logo = (school or {}).get('logo_path')
            if logo and os.path.exists(logo):
                try:
                    from reportlab.lib.utils import ImageReader
                    ir = ImageReader(logo); iw, ih = ir.getSize()
                    h = 20 * mm; w = h * (iw / ih) if ih else 20 * mm
                    c.drawImage(ir, x, top - h, width=min(w, 22 * mm), height=h,
                                mask='auto', preserveAspectRatio=True)
                    lx = x + min(w, 22 * mm) + 5 * mm
                except Exception:
                    lx = x
            name = ((school or {}).get('name') or 'School').upper()
            size = 20
            while size > 12 and stringWidth(name, boldf, size) > (avail - (lx - x)):
                size -= 1
            c.setFillColor(navy); c.setFont(boldf, size)
            c.drawString(lx, top - 6.5 * mm, name)
            c.setFillColor(ink); c.setFont(base, 9)
            addr = (school or {}).get('address') or ''
            if addr:
                c.drawString(lx, top - 11.5 * mm, addr[:110])
            contact = '    '.join(p for p in [(school or {}).get('phone') or '',
                                              (school or {}).get('email') or ''] if p)
            if contact:
                c.drawString(lx, top - 15.5 * mm, contact[:110])
            motto = (school or {}).get('motto') or ''
            if motto:
                c.setFillColor(gold); c.setFont(obl, 9.5)
                c.drawString(lx, top - 20 * mm, ('—  ' + motto + '  —')[:110])
            c.setFillColor(navy); c.setFont(boldf, 16)
            c.drawString(margin, top - 30 * mm, 'APPLICATION FORM')
            c.setFillColor(muted); c.setFont(obl, 9)
            c.drawString(margin, top - 35 * mm, 'Please complete in BLOCK letters.')
            c.setStrokeColor(gold); c.setLineWidth(1)
            c.line(margin, top - 38 * mm, PW - margin, top - 38 * mm)
            state['y'] = top - 46 * mm
        else:
            state['y'] = top
        # footer
        c.setStrokeColor(line); c.setLineWidth(0.6)
        c.line(margin, margin + 7 * mm, PW - margin, margin + 7 * mm)
        c.setFillColor(muted); c.setFont(obl, 8)
        c.drawString(margin, margin + 3.5 * mm, 'Blank application form · %s'
                     % ((school or {}).get('name') or 'School'))
        c.drawRightString(PW - margin, margin + 3.5 * mm, 'Page %d' % state['page'])

    def ensure(space):
        if state['y'] - space < margin + 12 * mm:
            c.showPage(); new_page(first=False)

    def heading(text):
        ensure(13 * mm)
        state['y'] -= 2 * mm
        c.setFillColor(navy); c.setFont(boldf, 11.5)
        c.drawString(margin, state['y'] - 3.5 * mm, text)
        c.setStrokeColor(line); c.setLineWidth(0.5)
        c.line(margin, state['y'] - 5.5 * mm, PW - margin, state['y'] - 5.5 * mm)
        state['y'] -= 8 * mm

    fh = 7 * mm   # field height

    def field(label, x, w, big=False):
        h = (fh + 6 * mm) if big else fh
        c.setFillColor(muted); c.setFont(boldf, 8)
        c.drawString(x, state['y'] - 2.6 * mm, label.upper())
        state['n'] += 1
        form.textfield(name='f%d' % state['n'], tooltip=label,
                       x=x, y=state['y'] - 4 * mm - h, width=w, height=h,
                       borderColor=field_border, fillColor=panel, textColor=ink,
                       borderWidth=field_bw, forceBorder=True, fontSize=11,
                       fieldFlags=('multiline' if big else ''))
        return h

    def row(fields, maxw=None):
        # fields: list of (label, weight, big?)
        big = any(len(f) > 2 and f[2] for f in fields)
        gap = 5 * mm
        wsum = sum(f[1] for f in fields)
        span = (maxw if maxw is not None else avail) - gap * (len(fields) - 1)
        ensure((fh + 6 * mm if big else fh) + 8 * mm)
        x = margin
        maxh = 0
        for f in fields:
            w = span * (f[1] / wsum)
            maxh = max(maxh, field(f[0], x, w, big=(len(f) > 2 and f[2])))
            x += w + gap
        state['y'] -= (maxh + 8 * mm)

    new_page(first=True)

    # Photo box (top-right of the applicant section)
    heading('Applicant')
    box_w, box_h = 32 * mm, 38 * mm
    bx = PW - margin - box_w
    c.setStrokeColor(line); c.setDash(2, 2); c.setLineWidth(0.8)
    c.rect(bx, state['y'] - box_h + 6 * mm, box_w, box_h, stroke=1, fill=0)
    c.setDash(); c.setFillColor(muted); c.setFont(base, 8)
    c.drawCentredString(bx + box_w / 2, state['y'] - box_h + box_h / 2 + 4 * mm, 'Passport')
    c.drawCentredString(bx + box_w / 2, state['y'] - box_h + box_h / 2, 'photograph')
    # the first two rows sit alongside the photo box → keep them to its left
    clear = avail - box_w - 12 * mm
    row([('First name', 1), ('Surname', 1)], maxw=clear)
    row([('Middle name', 1), ('Gender', 1)], maxw=clear)
    row([('Date of birth', 1), ('Previous school', 1.4)])

    heading('Origin & Health')
    row([('Country', 1), ('State of origin', 1), ('L.G.A. of origin', 1.1)])
    row([("Father's occupation", 1.4), ('Blood group', 0.8), ('Genotype', 0.8)])

    heading('Application')
    row([('Intended class', 1), ('Session / year', 1), ('Entrance score', 0.7)])

    heading('Parent / Guardian')
    row([('Full name', 1.3), ('Relationship', 1), ('Phone', 1)])
    row([('Email', 1), ('Residential address', 1.6)])

    heading('Emergency Contact')
    row([('Full name', 1.3), ('Relationship', 1), ('Phone', 1), ('Address', 1.4)])

    # signature line
    ensure(20 * mm)
    state['y'] -= 4 * mm
    c.setStrokeColor(ink); c.setLineWidth(0.7)
    c.line(margin, state['y'] - 6 * mm, margin + 70 * mm, state['y'] - 6 * mm)
    c.line(PW - margin - 55 * mm, state['y'] - 6 * mm, PW - margin, state['y'] - 6 * mm)
    c.setFillColor(muted); c.setFont(base, 9)
    c.drawString(margin, state['y'] - 11 * mm, 'Parent / Guardian signature')
    c.drawString(PW - margin - 55 * mm, state['y'] - 11 * mm, 'Date')

    c.showPage(); c.save()
    return buf.getvalue()
